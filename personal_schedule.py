import datetime
import json
import mimetypes
import os
import re
import sys

import dateutil.rrule
import requests
from cryptography.fernet import Fernet
from flask import Flask
from flask import render_template, request, redirect, url_for, session
from flask import send_file
from flask_mail import Mail, Message, Attachment

import mail_data
from models import *

from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import vobject

import pandas as pd
import random
import string
from bs4 import BeautifulSoup

app = Flask(__name__, instance_relative_config=False)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///team_db.db'
# app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///team_db_arch.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.app_context().push()
db.app = app
db.init_app(app)
db.create_all()

app.config['SECRET_KEY'] = mail_data.mail['SECRET_KEY']

app.config['MAIL_SERVER'] = mail_data.mail['MAIL_SERVER']
app.config['MAIL_PORT'] = mail_data.mail['MAIL_PORT']
app.config['MAIL_USERNAME'] = mail_data.mail['MAIL_USERNAME']
app.config['MAIL_PASSWORD'] = mail_data.mail['MAIL_PASSWORD']
app.config['MAIL_USE_TLS'] = mail_data.mail['MAIL_USE_TLS']
app.config['MAIL_USE_SSL'] = mail_data.mail['MAIL_USE_SSL']

mail = Mail(app)

tel_unneeded = '-() '
curr_year = 2024
# curr_year = 2023
fee = 5000
tour_fee = 3900

days = {'1': 'Пн', '2': 'Вт', '3': 'Ср', '4': 'Чт', '5': 'Пт', '6': 'Сб', '0': 'Вс'}
days_full = {'1': 'Понедельник',
             '2': 'Вторник',
             '3': 'Среда',
             '4': 'Четверг',
             '5': 'Пятница',
             '6': 'Суббота',
             '0': 'Воскресенье'}
months = {'01': 'янв',
          '02': 'фев',
          '03': 'мар',
          '04': 'апр',
          '05': 'май',
          '06': 'июн',
          '07': 'июл',
          '08': 'авг',
          '09': 'сен',
          '10': 'окт',
          '11': 'ноя',
          '12': 'дек'}
months_full = {'01': 'января',
               '02': 'февраля',
               '03': 'марта',
               '04': 'апреля',
               '05': 'мая',
               '06': 'июня',
               '07': 'июля',
               '08': 'августа',
               '09': 'сентября',
               '10': 'октября',
               '11': 'ноября',
               '12': 'декабря'}

access_types = {'guest': 0,
                'user': 1,
                'approved_user': 2,
                'team': 3,
                'secretary': 5,
                'supervisor': 6,
                'other-org': 7,
                'org': 8,
                'manager': 9,
                'admin': 10}

tasks = [{'id': t.task_id,
          'task_name': t.task_name,
          'location': t.location,
          'address': t.address,
          'description': t.description,
          'real_date': t.start_time,
          'task_date': days[t.start_time.strftime('%w')] + ', ' + t.start_time.strftime('%d') + ' ' +
                       months_full[t.start_time.strftime('%m')],
          'start_time': datetime.datetime.strftime(t.start_time, '%H:%M'),
          'end_time': datetime.datetime.strftime(t.end_time, '%H:%M'),
          'volunteers_required': t.volunteers_required}
         for t in VolunteerTasks.query.filter(VolunteerTasks.year == curr_year)
         .order_by(VolunteerTasks.start_time).all()]

# task_days = [task['real_date'] for task in tasks]
# day_tasks = {day.date(): [] for day in sorted(list(set(task_days)))}
# t = [d for d in day_tasks.keys()]
#
# for task in tasks:
#     day_tasks[task['real_date'].date()].append(task)
#
# name = 'Бурмистрова Дарья Ильинична'
#
# cl = '11ПРАВ'
# eng = 'Английский язык (40)'
# l2 = 'Китайский язык (6)'
#
# document = Document()
# sections = document.sections
# for section in sections:
#     section.top_margin = Cm(0.6)
#     section.bottom_margin = Cm(0.6)
#     section.left_margin = Cm(0.6)
#     section.right_margin = Cm(0.6)
# style = document.styles['Normal']
# font = style.font
# font.name = 'Calibri'
# font.size = Pt(12)
#
# i = 0
# while i < len(t):
#     d_tasks = day_tasks[t[i]]
#     a = document.add_paragraph()
#     a.alignment = 1
#     p = a.add_run('Пропуск уроков обучающимся: ' + name + ' (' + cl + ')' + '\n')
#     p.bold = True
#     pp = a.add_run('(волонтер Чтений им. В. И. Вернадского)')
#     a.paragraph_format.space_after = Pt(6)
#
#     b = document.add_paragraph()
#     b.alignment = 1
#     ppp = b.add_run(d_tasks[0]['task_date'])
#     ppp.bold = True
#
#     font = ppp.font
#     font.name = 'Calibri'
#     font.size = Pt(14)
#
#     font = p.font
#     font.name = 'Calibri'
#     font.size = Pt(16)
#
#     table = document.add_table(cols=4, rows=1)
#     table.style = 'Table Grid'
#     hdr_cells = table.rows[0].cells
#
#     hdr_cells[0].paragraphs[0].add_run('Место').bold = True
#     hdr_cells[0].width = Cm(7.2)
#     hdr_cells[1].paragraphs[0].add_run('Задача волонтера').bold = True
#     hdr_cells[1].width = Cm(7.4)
#     hdr_cells[2].paragraphs[0].add_run('Время работы/ уроки').bold = True
#     hdr_cells[2].width = Cm(3.4)
#     hdr_cells[3].paragraphs[0].add_run('Выбор задачи').bold = True
#     hdr_cells[3].width = Cm(1.8)
#
#     for task in d_tasks:
#         row_cells = table.add_row().cells
#         row_cells[0].text = task['location'] + '\n(' + task['address'] + ')'
#         row_cells[1].text = task['task_name']
#         row_cells[2].text = task['start_time'] + '-' + task['end_time']
#
#     for row in table.rows:
#         for cell in row.cells:
#             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
#             cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#         row.height = Cm(1.4)
#
#     # document.add_paragraph()
#     document.add_paragraph()
#
#     table = document.add_table(cols=3, rows=1)
#     table.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     table.style = 'Table Grid'
#     hdr_cells = table.rows[0].cells
#     a = table.cell(0, 0)
#     b = table.cell(0, 1)
#     A = a.merge(b)
#
#     hdr_cells[0].paragraphs[0].add_run('Урок').bold = True
#     hdr_cells[0].width = Cm(9.9)
#     # hdr_cells[1].paragraphs[0].add_run('Задача волонтера').bold = True
#     hdr_cells[1].width = Cm(1)
#     hdr_cells[2].paragraphs[0].add_run('Пропуск согласован' + '\n' + '(подпись, задание при наличии)').bold = True
#     hdr_cells[2].width = Cm(8.9)
#
#     for lesson in range(0, 8):
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(lesson + 1)
#         row_cells[1].text = 'one'
#         row_cells[2].text = 'two'
#
#     for row in table.rows:
#         for cell in row.cells:
#             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
#             cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#         row.height = Cm(1.4)
#
#     if i < len(day_tasks) - 1:
#         document.add_page_break()
#     i += 1
#
# document.save('demo.docx')


# for day in c_dates:
#     document.add_heading(day['day_full'], level=1)
#     for work in day['works']:
#         document.add_paragraph(str(work['report_order']) + '. ' + str(work['work_id']) + ' – ' + work['work_name'] +
#                                ' – ' + work['authors'], style='Normal')
#
# document.save(path)

a = datetime.datetime.strptime('1900-01-01 09:15:00', '%Y-%m-%d %H:%M:%S').date()
b = datetime.datetime.strptime('2024-04-14 09:00:00', '%Y-%m-%d %H:%M:%S').date()

c = b - a

print(c)
print(type(c))
