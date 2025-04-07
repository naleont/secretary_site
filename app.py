import datetime
import json
import mimetypes
import os
import re
import sys

# import dateutil.rrule
import requests
from cryptography.fernet import Fernet
from flask import Flask
from flask import render_template, request, redirect, url_for, session
from flask import send_file
from flask_mail import Mail, Message, Attachment

import mail_data
from models import *

from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

import vobject

import pandas as pd
import random
import string
from bs4 import BeautifulSoup

from gmail_sender import *

from urllib.parse import unquote

app = Flask(__name__, instance_relative_config=False)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///team_db.db'
# app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///team_db_arch_2024.db'
# app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///team_db_arch.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.app_context().push()
db.app = app
db.init_app(app)
db.create_all()

app.config['SECRET_KEY'] = mail_data.mail['SECRET_KEY']

app.config['MAIL_SERVER'] = mail_data.mail['MAIL_SERVER']
app.config['MAIL_PORT'] = mail_data.mail['MAIL_PORT']
app.config['MAIL_USERNAME'] = mail_data.mail['MAIL_USERNAME']
app.config['MAIL_PASSWORD'] = mail_data.mail['MAIL_PASSWORD']
app.config['MAIL_USE_TLS'] = mail_data.mail['MAIL_USE_TLS']
app.config['MAIL_USE_SSL'] = mail_data.mail['MAIL_USE_SSL']

mail = Mail(app)

tel_unneeded = '-() '
curr_year = 2025
# curr_year = 2024
# curr_year = 2023
fee = 5500
tour_fee = 4000

days = {'1': 'Пн', '2': 'Вт', '3': 'Ср', '4': 'Чт', '5': 'Пт', '6': 'Сб', '0': 'Вс'}
days_full = {'1': 'Понедельник',
             '2': 'Вторник',
             '3': 'Среда',
             '4': 'Четверг',
             '5': 'Пятница',
             '6': 'Суббота',
             '0': 'Воскресенье'}
months = {'01': 'янв',
          '02': 'фев',
          '03': 'мар',
          '04': 'апр',
          '05': 'май',
          '06': 'июн',
          '07': 'июл',
          '08': 'авг',
          '09': 'сен',
          '10': 'окт',
          '11': 'ноя',
          '12': 'дек'}
months_full = {'01': 'января',
               '02': 'февраля',
               '03': 'марта',
               '04': 'апреля',
               '05': 'мая',
               '06': 'июня',
               '07': 'июля',
               '08': 'августа',
               '09': 'сентября',
               '10': 'октября',
               '11': 'ноября',
               '12': 'декабря'}

access_types = {'guest': 0,
                'user': 1,
                'approved_user': 2,
                'team': 3,
                'secretary': 5,
                'supervisor': 6,
                'other-org': 7,
                'org': 8,
                'manager': 9,
                'admin': 10}

MSU_lessons = {1: {'start': datetime.datetime.strptime('08:30', '%H:%M'),
                   'end': datetime.datetime.strptime('09:15', '%H:%M')},
               2: {'start': datetime.datetime.strptime('09:25', '%H:%M'),
                   'end': datetime.datetime.strptime('10:10', '%H:%M')},
               3: {'start': datetime.datetime.strptime('10:30', '%H:%M'),
                   'end': datetime.datetime.strptime('11:15', '%H:%M')},
               4: {'start': datetime.datetime.strptime('11:25', '%H:%M'),
                   'end': datetime.datetime.strptime('12:10', '%H:%M')},
               5: {'start': datetime.datetime.strptime('13:10', '%H:%M'),
                   'end': datetime.datetime.strptime('13:55', '%H:%M')},
               6: {'start': datetime.datetime.strptime('14:05', '%H:%M'),
                   'end': datetime.datetime.strptime('14:50', '%H:%M')},
               7: {'start': datetime.datetime.strptime('15:05', '%H:%M'),
                   'end': datetime.datetime.strptime('15:50', '%H:%M')},
               8: {'start': datetime.datetime.strptime('16:05', '%H:%M'),
                   'end': datetime.datetime.strptime('16:50', '%H:%M')}}

MSU_lessons_10 = {1: {'start': datetime.datetime.strptime('08:30', '%H:%M'),
                      'end': datetime.datetime.strptime('09:15', '%H:%M')},
                  2: {'start': datetime.datetime.strptime('09:25', '%H:%M'),
                      'end': datetime.datetime.strptime('10:10', '%H:%M')},
                  3: {'start': datetime.datetime.strptime('10:30', '%H:%M'),
                      'end': datetime.datetime.strptime('11:15', '%H:%M')},
                  4: {'start': datetime.datetime.strptime('11:25', '%H:%M'),
                      'end': datetime.datetime.strptime('12:10', '%H:%M')},
                  5: {'start': datetime.datetime.strptime('12:20', '%H:%M'),
                      'end': datetime.datetime.strptime('13:05', '%H:%M')},
                  6: {'start': datetime.datetime.strptime('14:05', '%H:%M'),
                      'end': datetime.datetime.strptime('14:50', '%H:%M')},
                  7: {'start': datetime.datetime.strptime('15:05', '%H:%M'),
                      'end': datetime.datetime.strptime('15:50', '%H:%M')},
                  8: {'start': datetime.datetime.strptime('16:05', '%H:%M'),
                      'end': datetime.datetime.strptime('16:50', '%H:%M')}}


def renew_session():
    if 'user_id' in session.keys():
        user_db = db.session.query(Users).filter(Users.user_id == session['user_id']).first()
        if user_db is not None:
            cat_sec = db.session.query(CatSecretaries).filter(CatSecretaries.secretary_id == session['user_id']).all()
            cat_online_sec = db.session.query(OnlineSecretaries).filter(OnlineSecretaries.secretary_id ==
                                                                        session['user_id']).all()
            user = session['user_id']
            session['type'] = user_db.user_type
            session['approved'] = user_db.approved
            if session['approved'] is True:
                session['access'] = 2
            else:
                session['access'] = 1
            if user in [u.secretary_id for u in CatSecretaries.query.all()]:
                session['secretary'] = True
                session['access'] = 5
                if 'cat_id' not in session.keys():
                    session['cat_id'] = [c.cat_id for c in cat_sec]
                else:
                    session['cat_id'].extend([c.cat_id for c in cat_sec])
            if user in [u.secretary_id for u in OnlineSecretaries.query.all()]:
                session['secretary'] = True
                session['access'] = 5
                if 'cat_id' not in session.keys():
                    session['cat_id'] = [c.cat_id for c in cat_online_sec]
                else:
                    session['cat_id'].extend([c.cat_id for c in cat_online_sec])
            if user in [u.user_id for u in SupervisorUser.query.all()]:
                session['supervisor'] = True
                session['access'] = 6
                supervisor = SupervisorUser.query.filter(SupervisorUser.user_id == user).first()
                if supervisor.supervisor_id in [s.supervisor_id for s in CatSupervisors.query.all()]:
                    cat_sup = CatSupervisors.query.filter(CatSupervisors.supervisor_id == supervisor.supervisor_id
                                                          ).all()
                    session['cat_id'] = [c.cat_id for c in cat_sup]
            else:
                session['supervisor'] = False
            if user in [p.user_id for p in Profile.query.all()]:
                session['profile'] = True
            if user in [a.user_id for a in Application.query.filter(Application.year == curr_year)]:
                session['application'] = True
            else:
                session['application'] = False
            for a, t in access_types.items():
                if a == session['type']:
                    if 'access' in session.keys():
                        if session['access'] < t:
                            session['access'] = t
                    else:
                        session['access'] = t
            if user in [u.user_id for u in TutorUser.query
                    .join(SchoolClasses, TutorUser.class_id == SchoolClasses.class_id)
                    .filter(SchoolClasses.year == curr_year).all()]:
                session['tutor'] = True
                if session['access'] < 7:
                    session['access'] = 7
                session['class_id'] = TutorUser.query.join(SchoolClasses, TutorUser.class_id == SchoolClasses.class_id) \
                    .filter(SchoolClasses.year == curr_year).filter(TutorUser.user_id == user).first().class_id
            else:
                session['tutor'] = False
    return session


def check_access(access):
    if not request.url:
        url = ''
    else:
        url = request.url.replace(request.url_root, '').strip('/').split('/')
    renew_session()
    if 'access' not in session.keys():
        return redirect(url_for('.no_access', url=url, message='login_first'))
    else:
        if session['access'] < access:
            return redirect(url_for('.no_access', url=url, message='ok'))
        else:
            return True


def create_key():
    key = Fernet.generate_key()
    file = open('secret.key', 'wb')
    file.write(key)
    file.close()
    return key


# Загрузка ключа шифрования
def load_key():
    if not os.path.isfile('secret.key'):
        create_key()
    return open("secret.key", "rb").read()


# Шифрование текста в переменной message
def encrypt(message):
    key = load_key()
    encoded_message = message.encode()
    f = Fernet(key)
    encrypted = f.encrypt(encoded_message)
    return encrypted


# Расшифровка текста в переменной encrypted_message
def decrypt(encrypted_message):
    key = load_key()
    f = Fernet(key)
    decrypted = f.decrypt(encrypted_message)
    return decrypted.decode()


# Отправка письма для подтверждения регистрации на адрес email
def send_email(email):
    user_id = db.session.query(Users).filter(Users.email == email).first().user_id
    link = request.url_root + 'approve/' + str(user_id)

    sender = "info@vernadsky.info"  # Здесь можно указать просто адрес.
    # Или "Команда Конкурса им. В. И. Вернадского <team@vernadsky.info>".
    subject = "Подтверждение e-mail"
    body = (
        "Это подтверждение вашей регистрации на сайте для секретарей Конкурса им. В. И. Вернадского.\n"
        f"Перейдите по ссылке для подтверждения email: {link}"
    )

    service = get_service()
    message = create_message_text(sender, email, subject, body)
    send_message(service, "me", message)

def find_user(user_got):
    tel = re.sub(
        r'(^\+7|^8|^7|^9)(-|\(|\)|\s)*(?P<a>\d+)(-|\(|\)|\s)*(?P<b>\d+)(-|\(|\)|\s)*(?P<c>\d+)(-|\(|\)|\s)*(?P<d>\d+)',
        r'+7\g<a>\g<b>\g<c>\g<d>', user_got)
    if user_got in [user.email for user in Users.query.all()]:
        user = db.session.query(Users).filter(Users.email == user_got).first()
    elif tel in [user.tel for user in Users.query.all()]:
        user = db.session.query(Users).filter(Users.tel == tel).first()
    else:
        return None
    return user


def personal_info_form():
    info = dict()
    if 'user_id' in request.form.keys():
        info['user_id'] = request.form['user_id']
    info['email'] = request.form['email']
    tel_n = request.form['tel']
    info['tel'] = re.sub(r'^8|^7|^(?=9)', '+7', ''.join([n for n in tel_n if n not in tel_unneeded]))
    info['last_name'] = request.form['last_name']
    info['first_name'] = request.form['first_name']
    info['patronymic'] = request.form['patronymic']
    return info


# Загрузка информации пользователя из БД
def get_user_info(user):
    user = int(user)
    user_info = dict()
    user_db = db.session.query(Users).filter(Users.user_id == user).first()
    user_info['user_id'] = user
    user_info['email'] = user_db.email
    user_info['tel'] = user_db.tel
    user_info['last_name'] = user_db.last_name
    user_info['first_name'] = user_db.first_name
    user_info['patronymic'] = user_db.patronymic
    user_info['name_initials'] = user_info['last_name'] + ' ' + user_info['first_name'][0] + '. ' + \
        user_info['patronymic'][0] + '.'
    user_info['type'] = user_db.user_type
    user_info['approved'] = user_db.approved
    user_info['created_on'] = user_db.created_on.strftime('%d.%m.%Y %H:%M:%S')
    year_cats = [c.cat_id for c in Categories.query.filter(Categories.year == curr_year).all()]
    if user_db.last_login:
        user_info['last_login'] = user_db.last_login.strftime('%d.%m.%Y %H:%M:%S')
    if user in [u.secretary_id for u in CatSecretaries.query.all()]:
        user_info['secretary'] = True
        if 'cat_id' not in user_info.keys():
            user_info['cat_id'] = [c.cat_id for c in db.session.query(CatSecretaries).filter(
                CatSecretaries.secretary_id == user).all() if c.cat_id in year_cats]
        else:
            user_info['cat_id'].extend([c.cat_id for c in db.session.query(CatSecretaries).filter(
                CatSecretaries.secretary_id == user).all() if c.cat_id in year_cats])
    else:
        user_info['cat_id'] = []
    if user in [u.secretary_id for u in OnlineSecretaries.query.all()]:
        user_info['online_secretary'] = True
        if 'online_cat_id' not in user_info.keys():
            user_info['online_cat_id'] = [c.cat_id for c in db.session.query(OnlineSecretaries).filter(
                OnlineSecretaries.secretary_id == user).all() if c.cat_id in year_cats]
        else:
            user_info['online_cat_id'].extend([c.cat_id for c in db.session.query(OnlineSecretaries).filter(
                OnlineSecretaries.secretary_id == user).all() if c.cat_id in year_cats])
    else:
        user_info['online_cat_id'] = []
    if user in [s.user_id for s in SupervisorUser.query.all()]:
        user_info['supervisor_id'] = SupervisorUser.query.filter(SupervisorUser.user_id == user).first().supervisor_id

    if user in [u.user_id for u in TutorUser.query
            .join(SchoolClasses, TutorUser.class_id == SchoolClasses.class_id)
            .filter(SchoolClasses.year == curr_year).all()]:
        user_info['tutor'] = True
        user_info['class_id'] = TutorUser.query.join(SchoolClasses, TutorUser.class_id == SchoolClasses.class_id) \
            .filter(SchoolClasses.year == curr_year).filter(TutorUser.user_id == user).first().class_id
    return user_info


def get_org_info(user_id):
    org = get_user_info(user_id)
    resps = [r.responsibility_id for r
             in ResponsibilityAssignment.query.join(Responsibilities, ResponsibilityAssignment.responsibility_id ==
                                                    Responsibilities.responsibility_id)
             .filter(ResponsibilityAssignment.user_id == org['user_id']).filter(Responsibilities.year == curr_year)
             .all()]
    responsibilities = []
    for resp in resps:
        resp_db = db.session.query(Responsibilities).filter(Responsibilities.responsibility_id == resp).first()
        responsibility = {'id': resp_db.responsibility_id, 'name': resp_db.name, 'description': resp_db.description}
        responsibilities.append(responsibility)
    responsibs = sorted(responsibilities, key=lambda u: u['name'])
    org['responsibilities'] = responsibs
    return org


def all_users():
    users = dict()
    for u in Users.query.order_by(Users.user_id.desc()).all():
        users[u.user_id] = get_user_info(u.user_id)
    return users


# Загрузка информации профиля из БД
def get_profile_info(user):
    user_id = int(user)
    profile = dict()
    profile['user_id'] = user_id
    if db.session.query(Profile).filter(Profile.user_id == user_id).first():
        prof_info = db.session.query(Profile).filter(Profile.user_id == user).first()
        profile['vk'] = prof_info.vk
        profile['tg'] = prof_info.telegram
        profile['username'] = prof_info.vernadsky_username
        profile['filled'] = True
        profile['occupation'] = prof_info.occupation
        profile['involved'] = prof_info.involved
        profile['place_of_work'] = prof_info.place_of_work
        profile['grade'] = prof_info.grade
        profile['year'] = prof_info.year
        profile['born'] = prof_info.born
    else:
        profile = {'user_id': profile['user_id'], 'filled': False, 'vk': None, 'tg': None, 'username': None,
                   'occupation': None, 'involved': None, 'place_of_work': None, 'grade': None, 'year': None,
                   'born': None}
    return profile


# Запись исправленной информации пользователя в БД
def write_user(user_info):
    if 'user_id' not in user_info.keys() and 'user_id' in session.keys():
        user_id = int(session['user_id'])
    elif 'user_id' in session.keys():
        user_id = int(user_info['user_id'])
    else:
        user_id = None
    if user_id:
        # Загрузка информации пользователя из БД
        user_db = db.session.query(Users).filter(Users.user_id == user_id).first()
        # Проверка существования другого пользователя с новым введенным email
        same_email = [user.user_id for user in db.session.query(Users).filter(Users.email == user_info['email']).all()]
        if not same_email:
            user_db.email = user_info['email']
        elif user_id in same_email:
            if not same_email.remove(user_id):
                user_db.email = user_info['email']
        else:
            return 'email'
        # Проверка существования другого пользователя с новым введенным телефоном
        same_tel = [user.user_id for user in db.session.query(Users).filter(Users.email == user_info['tel']).all()]
        if not same_tel:
            user_db.tel = user_info['tel']
        elif user_id in same_tel:
            if same_tel.remove(user_id) is None:
                user_db.tel = user_info['tel']
            else:
                return 'tel'

        db.session.query(Users).filter(Users.user_id == user_id).update(
            {Users.last_name: user_info['last_name'], Users.first_name: user_info['first_name'],
             Users.patronymic: user_info['patronymic']})
    else:
        user = Users(user_info['email'], user_info['tel'], user_info['password'], user_info['last_name'],
                     user_info['first_name'], user_info['patronymic'], user_info['user_type'],
                     user_info['approved'], None)
        db.session.add(user)
    db.session.commit()
    return 'ok'


def write_category(cat_info):
    if cat_info['cat_id'] is None:
        cat_info['cat_id'] = max([c.cat_id for c in Categories.query.all()]) + 1
    if cat_info['cat_id'] in [cat.cat_id for cat in Categories.query.all()]:
        db.session.query(Categories).filter(Categories.cat_id == cat_info['cat_id']).update(
            {Categories.year: curr_year, Categories.cat_name: cat_info['cat_name'],
             Categories.short_name: cat_info['short_name'], Categories.tg_channel: cat_info['tg_channel'],
             Categories.cat_site_id: cat_info['cat_site_id'], Categories.drive_link: cat_info['drive_link']})
        if cat_info['cat_id'] in [cat_dir.cat_id for cat_dir in CatDirs.query.all()]:
            db.session.query(CatDirs).filter(CatDirs.cat_id == cat_info['cat_id']).update(
                {CatDirs.cat_id: cat_info['cat_id'], CatDirs.dir_id: cat_info['direction'],
                 CatDirs.contest_id: cat_info['contest']})
        else:
            cat_dir = CatDirs(cat_info['cat_id'], cat_info['direction'], cat_info['contest'])
            db.session.add(cat_dir)
        if cat_info['supervisor'] is not None and cat_info['supervisor'] != '':
            if cat_info['cat_id'] in [sup.cat_id for sup in CatSupervisors.query.all()]:
                db.session.query(CatSupervisors).filter(CatSupervisors.cat_id == cat_info['cat_id']).update(
                    {CatSupervisors.supervisor_id: cat_info['supervisor']})
            else:
                sup = db.session.query(Supervisors).filter(Supervisors.supervisor_id == cat_info['supervisor']).first()
                db_cat = db.session.query(Categories).filter(Categories.cat_id == cat_info['cat_id']).first()
                cat = CatSupervisors(db_cat.cat_id, sup.supervisor_id)
                db.session.add(cat)
        else:
            db.session.query(CatSupervisors).filter(CatSupervisors.cat_id == cat_info['cat_id']).delete()
    else:
        cat = Categories(curr_year, cat_info['cat_name'], cat_info['short_name'], cat_info['tg_channel'],
                         cat_info['cat_site_id'], cat_info['drive_link'])
        db.session.add(cat)
        db.session.commit()
        categ = db.session.query(Categories).filter(Categories.cat_name == cat_info['cat_name']
                                                    ).filter(Categories.year == curr_year).first()
        if type(cat_info['direction']) is int:
            direct = db.session.query(Directions).filter(Directions.direction_id == cat_info['direction']).first()
        else:
            direct = db.session.query(Directions).filter(Directions.dir_name == cat_info['direction']).first()
        if type(cat_info['contest']) is int:
            cont = db.session.query(Contests).filter(Contests.contest_id == cat_info['contest']).first()
        else:
            cont = db.session.query(Contests).filter(Contests.contest_name == cat_info['contest']).first()
        if cat_info['cat_id'] not in [catdir.cat_id for catdir in CatDirs.query.all()]:
            cat_dir = CatDirs(categ.cat_id, direct.direction_id, cont.contest_id)
            db.session.add(cat_dir)
        cat_info['cat_id'] = db.session.query(Categories).filter(
            Categories.cat_name == cat_info['cat_name']).first().cat_id
    if cat_info['cat_id'] in [cat_sup.cat_id for cat_sup in CatSupervisors.query.all()]:
        cat = db.session.query(CatSupervisors).filter(CatSupervisors.cat_id == cat_info['cat_id']).first()
        sup = db.session.query(Supervisors).filter(Supervisors.supervisor_id == cat_info['supervisor']).first()
        if sup is not None:
            cat.supervisor_id = sup.supervisor_id
    else:
        if type(cat_info['supervisor']) is int:
            sup = db.session.query(Supervisors).filter(Supervisors.supervisor_id == cat_info['supervisor']).first()
        else:
            if cat_info['supervisor'] is not None:
                sup_name = cat_info['supervisor'].split(' ')
                sup = db.session.query(Supervisors).filter(Supervisors.last_name == sup_name[0]
                                                           ).filter(Supervisors.first_name == sup_name[1]
                                                                    ).filter(Supervisors.patronymic == sup_name[2]
                                                                             ).first()
                db_cat = db.session.query(Categories).filter(Categories.cat_id == cat_info['cat_id']).first()
            else:
                sup = None
        if sup is not None:
            cat = CatSupervisors(db_cat.cat_id, sup.supervisor_id)
            db.session.add(cat)
    db.session.commit()
    return True


def one_category(categ):
    cat = {}
    cat_id = categ.cat_id
    cat['id'] = categ.cat_id
    cat['year'] = categ.year
    cat['name'] = categ.cat_name
    cat['short_name'] = categ.short_name
    cat['tg_channel'] = categ.tg_channel
    cat_dir = db.session.query(CatDirs).filter(CatDirs.cat_id == cat_id).first()
    direction = db.session.query(Directions).filter(Directions.direction_id == cat_dir.dir_id).first()
    cat['direction'] = direction.dir_name
    cat['dir_id'] = direction.direction_id
    contest = db.session.query(Contests).filter(Contests.contest_id == cat_dir.contest_id).first()
    cat['contest'] = contest.contest_name
    cat['cont_id'] = contest.contest_id
    cat['drive_link'] = categ.drive_link
    cat['cat_site_id'] = categ.cat_site_id
    if cat_id in [c.cat_id for c in CatSupervisors.query.all()]:
        sup = db.session.query(Supervisors).join(CatSupervisors).filter(CatSupervisors.cat_id == cat_id).first()
        cat['supervisor_id'] = sup.supervisor_id
        cat['supervisor'] = sup.last_name + ' ' + sup.first_name + ' ' + sup.patronymic
        cat['supervisor_email'] = sup.email
        cat['supervisor_tel'] = sup.tel
    if cat_id in [c.cat_id for c in CatSecretaries.query.all()]:
        user = db.session.query(Users).join(CatSecretaries).filter(CatSecretaries.cat_id == cat_id).first()
        cat['secretary_id'] = user.user_id
        cat['secretary'] = user.last_name + ' ' + user.first_name
        cat['secretary_full'] = user.last_name + ' ' + user.first_name + ' ' + user.patronymic
        cat['secretary_email'] = user.email
        cat['secretary_tel'] = user.tel
    if cat_id in [c.cat_id for c in OnlineSecretaries.query.all()]:
        user = db.session.query(Users).join(OnlineSecretaries).filter(OnlineSecretaries.cat_id == cat_id).first()
        cat['online_secretary_id'] = user.user_id
        cat['online_secretary'] = user.last_name + ' ' + user.first_name
        cat['online_secretary_full'] = user.last_name + ' ' + user.first_name + ' ' + user.patronymic
        cat['online_secretary_email'] = user.email
        cat['online_secretary_tel'] = user.tel
    if cat_id in [cat.cat_id for cat in ReportDates.query.all()]:
        dates_db = db.session.query(ReportDates).filter(ReportDates.cat_id == cat_id).first()
        dates = []
        if dates_db.day_1:
            dates.append(days[dates_db.day_1.strftime('%w')] + ' ' + dates_db.day_1.strftime('%d') + ' ' +
                         months_full[dates_db.day_1.strftime('%m')])
        if dates_db.day_2:
            dates.append(days[dates_db.day_2.strftime('%w')] + ' ' + dates_db.day_2.strftime('%d') + ' ' +
                         months_full[dates_db.day_2.strftime('%m')])
        if dates_db.day_3:
            dates.append(days[dates_db.day_3.strftime('%w')] + ' ' + dates_db.day_3.strftime('%d') + ' ' +
                         months_full[dates_db.day_3.strftime('%m')])
        cat['dates'] = ', '.join(dates)
    return cat


def categories_info(cat_id='all'):
    if cat_id == 'all':
        categories = db.session.query(Categories
                                      ).filter(Categories.year == curr_year
                                               ).join(CatDirs
                                                      ).join(Directions).join(Contests
                                                                              ).order_by(CatDirs.dir_id,
                                                                                         CatDirs.contest_id,
                                                                                         Categories.cat_name).all()

        dirs = {d.direction_id: {'dir_id': d.direction_id, 'dir_name': d.dir_name} for d in Directions.query.all()}
        conts = {cont.contest_id: {'cont_id': cont.contest_id, 'cont_name': cont.contest_name} for cont in
                 Contests.query.all()}
        cat_dir = {cd.cat_id: {'c_id': cd.cat_id, 'd_id': cd.dir_id, 'dir_name': dirs[cd.dir_id]['dir_name'],
                               'cont_id': cd.contest_id, 'cont_name': conts[cd.contest_id]['cont_name']} for cd in
                   CatDirs.query.all()}

        cat_sup = [c.cat_id for c in CatSupervisors.query.all()]
        cat_sec = [c.cat_id for c in CatSecretaries.query.all()]
        cat_online_sec = [c.cat_id for c in OnlineSecretaries.query.all()]
        cat_rep = [cat.cat_id for cat in ReportDates.query.all()]
        cats = []
        for categ in categories:

            cat = {}
            cat_id = categ.cat_id
            cat['id'] = cat_id
            cat['year'] = categ.year
            cat['name'] = categ.cat_name
            cat['short_name'] = categ.short_name
            cat['tg_channel'] = categ.tg_channel
            cat['direction'] = cat_dir[cat_id]['dir_name']
            cat['dir_id'] = cat_dir[cat_id]['d_id']
            cat['contest'] = cat_dir[cat_id]['cont_name']
            cat['cont_id'] = cat_dir[cat_id]['cont_id']
            cat['drive_link'] = categ.drive_link
            cat['cat_site_id'] = categ.cat_site_id
            if cat_id in cat_sup:
                sup = db.session.query(Supervisors).join(CatSupervisors).filter(
                    CatSupervisors.cat_id == cat_id).first()  # check
                cat['supervisor_id'] = sup.supervisor_id
                cat['supervisor'] = sup.last_name + ' ' + sup.first_name + ' ' + sup.patronymic
                cat['supervisor_email'] = sup.email
                cat['supervisor_tel'] = sup.tel
            if cat_id in cat_sec:
                user = db.session.query(Users).join(CatSecretaries).filter(
                    CatSecretaries.cat_id == cat_id).first()  # check
                cat['secretary_id'] = user.user_id
                cat['secretary'] = user.last_name + ' ' + user.first_name
                cat['secretary_full'] = user.last_name + ' ' + user.first_name + ' ' + user.patronymic
                cat['secretary_email'] = user.email
                cat['secretary_tel'] = user.tel
            if cat_id in cat_online_sec:
                user = db.session.query(Users).join(OnlineSecretaries).filter(
                    OnlineSecretaries.cat_id == cat_id).first()  # check
                cat['online_secretary_id'] = user.user_id
                cat['online_secretary'] = user.last_name + ' ' + user.first_name
                cat['online_secretary_full'] = user.last_name + ' ' + user.first_name + ' ' + user.patronymic
                cat['online_secretary_email'] = user.email
                cat['online_secretary_tel'] = user.tel
            if cat_id in cat_rep:
                dates_db = db.session.query(ReportDates).filter(ReportDates.cat_id == cat_id).first()  # check
                dates = []
                if dates_db.day_1:
                    dates.append(days[dates_db.day_1.strftime('%w')] + ', ' + dates_db.day_1.strftime('%d') + ' ' +
                                 months_full[dates_db.day_1.strftime('%m')])
                if dates_db.day_2:
                    dates.append(days[dates_db.day_2.strftime('%w')] + ', ' + dates_db.day_2.strftime('%d') + ' ' +
                                 months_full[dates_db.day_2.strftime('%m')])
                if dates_db.day_3:
                    dates.append(days[dates_db.day_3.strftime('%w')] + ', ' + dates_db.day_3.strftime('%d') + ' ' +
                                 months_full[dates_db.day_3.strftime('%m')])
                cat['dates'] = '; '.join(dates)

            cats.append(cat)
    else:
        category = db.session.query(Categories).filter(Categories.cat_id == cat_id).first()
        cats = one_category(category)
    cats_count = len(cats)
    return cats_count, cats


def get_supervisors():
    supervisors_list = db.session.query(Supervisors).order_by(Supervisors.last_name).all()
    sups = dict()
    for sup in supervisors_list:
        sups[sup.supervisor_id] = dict()
        sups[sup.supervisor_id]['id'] = sup.supervisor_id
        sups[sup.supervisor_id]['name'] = sup.last_name + ' ' + sup.first_name + ' ' + sup.patronymic
        sups[sup.supervisor_id]['email'] = sup.email
        sups[sup.supervisor_id]['tel'] = sup.tel
    return sups


def supervisor_info(sup_id):
    sup = db.session.query(Supervisors).filter(Supervisors.supervisor_id == sup_id).first()
    sup_info = dict()
    sup_info['id'] = sup.supervisor_id
    sup_info['name'] = sup.last_name + ' ' + sup.first_name + ' ' + sup.patronymic
    sup_info['last_name'] = sup.last_name
    sup_info['first_name'] = sup.first_name
    sup_info['patronymic'] = sup.patronymic
    sup_info['email'] = sup.email
    sup_info['tel'] = sup.tel
    categories = db.session.query(CatSupervisors).filter(CatSupervisors.supervisor_id == sup_info['id']).all()
    sup_categories = []
    cats_db = db.session.query(Categories)
    for cat in categories:
        sup_categories.append(cats_db.filter(Categories.cat_id == cat.cat_id).first().cat_name)
    sup_info['categories'] = ', '.join(sup_categories)
    return sup_info


def one_application(application):
    one = dict()
    categories = db.session.query(Categories)
    users = db.session.query(Users)
    user = users.filter(Users.user_id == application.user_id).first()
    one['user_id'] = user.user_id
    one['user'] = user.last_name + ' ' + user.first_name
    one['year'] = application.year
    one['role'] = application.role
    if application.category_1 != 'None':
        cat_1 = categories.filter(Categories.cat_id == application.category_1).first()
        one['category_1_id'] = cat_1.cat_id
        one['category_1'] = cat_1.cat_name
        one['category_1_short'] = cat_1.short_name
    if application.category_2 != 'None':
        cat_2 = categories.filter(Categories.cat_id == application.category_2).first()
        one['category_2_id'] = cat_2.cat_id
        one['category_2'] = cat_2.cat_name
        one['category_2_short'] = cat_2.short_name
    if application.category_3 != 'None':
        cat_3 = categories.filter(Categories.cat_id == application.category_3).first()
        one['category_3_id'] = cat_3.cat_id
        one['category_3'] = cat_3.cat_name
        one['category_3_short'] = cat_3.short_name
    one['any_category'] = application.any_category
    one['taken_part'] = application.taken_part
    one['considered'] = application.considered
    return one


def application_info(info_type, user, year=curr_year):
    if info_type == 'user':
        applications = db.session.query(Application).filter(Application.user_id == user).order_by(
            Application.year.desc())
    elif info_type == 'year':
        applications = db.session.query(Application).join(Users).filter(Application.year == year).order_by(
            Users.last_name)
    elif info_type == 'user-year':
        if user in [u.user_id for u in Application.query.filter(Application.year == year)]:
            applications = db.session.query(Application).filter(Application.user_id == user).filter(
                Application.year == year)
        else:
            applications = None
    else:
        applications = None
    appl = dict()
    if applications is not None:
        if info_type == 'user-year':
            appl = one_application(applications.first())
        else:
            for application in applications.all():
                if info_type == 'user':
                    key = application.year
                elif info_type == 'year':
                    key = application.user_id
                else:
                    key = application.user_id
                appl[key] = one_application(application)
    else:
        appl[curr_year] = dict()
        appl[curr_year]['role'], appl[curr_year]['category_1'], appl[curr_year]['category_2'], \
        appl[curr_year]['category_3'], appl[curr_year]['any_category'], appl[curr_year]['taken_part'], \
        appl[curr_year]['considered'] = None, None, None, None, None, None, None
    return appl


def one_news(news_id):
    n = db.session.query(News).filter(News.news_id == news_id).first()
    news = dict()
    news['news_id'] = news_id
    news['title'] = n.title
    news['content'] = n.content
    news['access'] = n.access
    news['publish'] = n.publish
    news['date'] = n.date_time.strftime('%d-%m-%Y')
    news['time'] = n.date_time.strftime('%H:%M')
    return news


def all_news():
    news_db = News.query.order_by(News.date_time.desc()).all()
    all_n = dict()
    for news in news_db:
        all_n[news.news_id] = one_news(news.news_id)
    return all_n


def work_info(work_id, additional_info=False, site_id=False, reports_info=False, analysis_info=False,
              w_payment_info=False, appl_info=False, cat_info=False, organisation_info=False, status_info=False,
              mail_info=False):
    work_id = int(work_id)
    work_db = db.session.query(Works).filter(Works.work_id == work_id).first()
    work = dict()
    work['work_id'] = work_id
    work['work_name'] = work_db.work_name
    work['author_1_name'] = work_db.author_1_name
    work['author_1_age'] = work_db.author_1_age
    work['author_1_class'] = work_db.author_1_class
    work['author_2_name'] = work_db.author_2_name
    work['author_2_age'] = work_db.author_2_age
    work['author_2_class'] = work_db.author_2_class
    work['author_3_name'] = work_db.author_3_name
    work['author_3_age'] = work_db.author_3_age
    work['author_3_class'] = work_db.author_3_class
    work['supervisor'] = work_db.teacher_name
    work['authors'] = work['author_1_name']
    if work['author_2_name'] is not None:
        work['authors'] += ', ' + work['author_2_name']
    if work['author_3_name'] is not None:
        work['authors'] += ', ' + work['author_3_name']
    if work_id in [w.work_id for w in WorkCategories.query.all()]:
        work['cat_id'] = WorkCategories.query.filter(WorkCategories.work_id == work_id).first().cat_id
        work['cat_short'] = Categories.query.filter(Categories.cat_id == work['cat_id']).first().short_name
    else:
        work['cat_id'] = None
        work['cat_short'] = None

    if cat_info is True:
        if work_id in [w.work_id for w in WorkCategories.query.all()]:
            work['cat_name'] = Categories.query.filter(Categories.cat_id == work['cat_id']).first().cat_name
        else:
            work['cat_name'] = None

    if status_info is True:
        work['status'] = ParticipationStatuses.query \
            .join(WorkStatuses, ParticipationStatuses.status_id == WorkStatuses.status_id) \
            .filter(WorkStatuses.work_id == work_id).first().status_name

    if organisation_info is True:
        if work_id in [w.work_id for w in WorkOrganisations.query.all()]:
            org_db = db.session.query(Organisations) \
                .join(WorkOrganisations, Organisations.organisation_id == WorkOrganisations.organisation_id) \
                .filter(WorkOrganisations.work_id == work_id).first()
            work['organisation_id'] = org_db.organisation_id
            work['organisation_name'] = org_db.name
            work['organisation_city'] = org_db.city
            work['organisation_country'] = org_db.country
        else:
            work['organisation_id'] = None
            work['organisation_name'] = None
            work['organisation_city'] = None
            work['organisation_country'] = None

    if site_id is True:
        work['site_id'] = work_db.work_site_id

    if additional_info is True:
        work['email'] = work_db.email
        work['tel'] = work_db.tel
        work['reg_tour'] = work_db.reg_tour

    if reports_info is True:
        work['reported'] = work_db.reported
        work['timeshift'] = work_db.msk_time_shift
        if work['timeshift']:
            if work['timeshift'] >= 0:
                work['timeshift'] = '+' + str(work['timeshift'])
            else:
                work['timeshift'] = str(work['timeshift'])
        if work_id in [w.work_id for w in ReportOrder.query.all()]:
            report = db.session.query(ReportOrder).filter(ReportOrder.work_id == work_id).first()
            work['report_day'] = report.report_day
            work['report_order'] = report.order

    if analysis_info is True:
        work['reg_tour'] = work_db.reg_tour
        if work_id in [w.work_id for w in RevAnalysis.query.all()]:
            if len(RevAnalysis.query.filter(RevAnalysis.work_id == work_id).all()
                   ) == len(
                RevCriteria.query.filter(RevCriteria.year == datetime.datetime.strptime(str(curr_year), '%Y').date()
                                         ).all()):
                work['analysis'] = True
            else:
                work['analysis'] = 'part'
        elif work_id in [w.work_id for w in PreAnalysis.query.all()]:
            pre = db.session.query(PreAnalysis).filter(PreAnalysis.work_id == work_id).first()
            if pre.has_review is False:
                work['analysis'] = True
            else:
                work['analysis'] = False
        else:
            work['analysis'] = False

    if w_payment_info is True:
        appl_for_online = [w.work_id for w in AppliedForOnline.query.all()]
        work['reg_tour'] = work_db.reg_tour
        if work['work_id'] in [w.work_id for w in Discounts.query.all()]:
            disc = db.session.query(Discounts).filter(Discounts.work_id == work['work_id']).first()
            work['fee'] = disc.payment
            work['format'] = disc.participation_format
        elif work['work_id'] in [w.work_id for w in WorksNoFee.query.all()]:
            work['fee'] = 0
        elif work['reg_tour'] is not None:
            work['fee'] = tour_fee
        elif work['work_id'] in [w.work_id for w in WorksNoFee.query.all()]:
            work['fee'] = 0
        else:
            work['fee'] = fee
        if work['work_id'] in [w.work_id for w in AppliedForOnline.query.all()]:
            work['format'] = 'online'
        if str(curr_year)[2:] not in [str(w.work_id)[:2] for w in ParticipatedWorks.query.all()]:
            if work_id in [w.work_id for w in ParticipatedWorks.query.all()] or \
                    work_id in [w.work_id for w in Applications2Tour.query.all() if w.appl_no is not None]:
                work['part_offline'] = True
                work['format'] = 'face-to-face'
            else:
                if work_id in appl_for_online:
                    work['part_offline'] = False
                else:
                    work['part_offline'] = None
        else:
            if work_id in [w.work_id for w in ParticipatedWorks.query.all()]:
                work['part_offline'] = True
                work['format'] = 'face-to-face'
            else:
                if work_id in appl_for_online:
                    work['part_offline'] = False
                else:
                    work['part_offline'] = None
        if work['work_id'] in [p.participant for p in PaymentRegistration.query.all()]:
            work['payed'] = True
            work['payment_id'] = PaymentRegistration.query.filter(PaymentRegistration.participant ==
                                                                  work['work_id']).first().payment_id
        elif work['fee'] == 0:
            work['payed'] = True
            work['payment_id'] = 'Работа участвует без оргвзноса'
        else:
            work['payed'] = False

    if appl_info is True:
        if work_id in [w.work_id for w in Applications2Tour.query.all()]:
            appl_db = db.session.query(Applications2Tour).filter(Applications2Tour.work_id == work_id).first()
            work['appl_no'] = appl_db.appl_no
            work['included'] = True
            if work['appl_no'] is None:
                work['appl_no'] = False
                work['included'] = False
            work['arrived'] = appl_db.arrived
        else:
            work['appl_no'] = False
            work['arrived'] = False
            work['included'] = False
        if str(work_id)[:5] == str(curr_year)[2:] + '000':
            work['copy_for_appl'] = str(work_id)[5:]
        elif str(work_id)[:4] == str(curr_year)[2:] + '00':
            work['copy_for_appl'] = str(work_id)[4:]
        elif str(work_id)[:3] == str(curr_year)[2:] + '0':
            work['copy_for_appl'] = str(work_id)[3:]
        else:
            work['copy_for_appl'] = str(work_id)[2:]

        if work['appl_no'] is False:
            if 'organisation_id' not in work.keys():
                if work_id in [w.work_id for w in WorkOrganisations.query.all()]:
                    work['organisation_id'] = WorkOrganisations.query.filter(WorkOrganisations.work_id == work_id) \
                        .first().organisation_id
                    if work['organisation_id'] in [o.organisation_id for o in OrganisationApplication.query.all()]:
                        work['org_arrived'] = OrganisationApplication.query \
                            .filter(OrganisationApplication.organisation_id == work['organisation_id']).first().arrived
            else:
                if work['organisation_id'] in [o.organisation_id for o in OrganisationApplication.query.all()]:
                    work['appl_no'] = OrganisationApplication.query \
                        .filter(OrganisationApplication.organisation_id == work['organisation_id']).first().appl_no
                    work['included'] = False
                    work['org_arrived'] = OrganisationApplication.query \
                        .filter(OrganisationApplication.organisation_id == work['organisation_id']).first().arrived
        if work['appl_no'] is None:
            work['appl_no'] = False
    if mail_info is True:
        work['mails'] = [{'mail_id': m.mail_id, 'email': m.email} for m
                         in Mails.query.join(WorkMail, Mails.mail_id == WorkMail.mail_id) \
                             .filter(WorkMail.work_id == work_id).all()]
        for a in work['mails']:
            a['sent'] = WorkMail.query.filter(WorkMail.work_id == work_id) \
                .filter(WorkMail.mail_id == a['mail_id']).first().sent
            if work_id in [w.work_id for w in Diplomas.query.all()]:
                c = Diplomas.query.filter(Diplomas.work_id == work_id).first().diplomas
                if not c:
                    a['sent'] = 'Дипломы не загружены'
    return work


def get_works(cat_id, status, mode='all', additional_info=False, site_id=False, reports_info=False, analysis_info=False,
              w_payment_info=False, appl_info=False, mail_info=False):
    works = dict()
    works_cat = [w.work_id for w in WorkCategories.query.filter(WorkCategories.cat_id == cat_id).all()]
    works_stat = [w.work_id for w in WorkStatuses.query.filter(WorkStatuses.status_id >= status).all() if w.work_id
                  in works_cat]
    if mode == 'online':
        works_searched = [w.work_id for w in AppliedForOnline.query.all() if w.work_id in works_stat]
    else:
        works_searched = works_stat
    for w in works_searched:
        works[w] = work_info(w, additional_info=additional_info, site_id=site_id, reports_info=reports_info,
                             analysis_info=analysis_info, w_payment_info=w_payment_info, appl_info=appl_info,
                             mail_info=mail_info)
    return works


def get_works_no_fee(cat_id):
    works_no_fee = {}
    cat_works = db.session.query(WorkCategories).filter(WorkCategories.cat_id == cat_id
                                                        ).order_by(WorkCategories.work_id).all()
    for work in cat_works:
        if work.work_id in [w.work_id for w in WorksNoFee.query.all()]:
            work_db = db.session.query(Works).filter(Works.work_id == work.work_id).first()
            w_no = work_db.work_id
            works_no_fee[w_no] = work_info(w_no)
    return works_no_fee


def get_criteria(year):
    criteria = dict()
    y = datetime.datetime.strptime(str(year), '%Y').date()
    crit_db = db.session.query(RevCriteria).filter(RevCriteria.year == y).all()
    crit_val_db = db.session.query(CriteriaValues)
    values_db = db.session.query(RevCritValues)
    for crit in crit_db:
        crit_info = dict()
        crit_id = crit.criterion_id
        crit_info['id'] = crit_id
        crit_info['name'] = crit.criterion_name
        crit_info['description'] = crit.criterion_description
        crit_info['year'] = datetime.datetime.strftime(crit.year, '%Y')
        crit_info['weight'] = crit.weight
        crit_values = crit_val_db.filter(CriteriaValues.criterion_id == crit_id).all()
        values = dict()
        for val in crit_values:
            v = dict()
            val_id = val.value_id
            value = values_db.filter(RevCritValues.value_id == val_id).first()
            v['value_id'] = val_id
            v['val_name'] = value.value_name
            v['comment'] = value.comment
            v['val_weight'] = value.weight
            values[val_id] = v
            if v['comment'] is not None and v['comment'] != '':
                crit_info['val_comment'] = True
            else:
                crit_info['val_comment'] = False
        crit_info['values'] = values
        crit_info['val_num'] = len(values)
        criteria[crit_id] = crit_info
    return criteria


def reg_works(cat_id='all', status=1):
    wks = []
    if status == 1:
        stat = ['Допущена до 1-го тура', 'Направлена на рецензирование', 'Отрецензирована',
                'Окончила 1-й тур. Не допущена до 2-го тура', 'Допущена до 2-го тура']
    elif status == 2:
        stat = ['Допущена до 2-го тура']
    if cat_id == 'all':
        categories = db.session.query(Categories
                                      ).filter(Categories.year == curr_year
                                               ).join(CatDirs
                                                      ).join(Directions
                                                             ).join(Contests
                                                                    ).order_by(CatDirs.dir_id, CatDirs.contest_id,
                                                                               Categories.cat_name).all()
        for cat in categories:
            cat_id = cat.cat_id
            works = {}
            cat_works = db.session.query(WorkCategories).filter(WorkCategories.cat_id == cat_id
                                                                ).order_by(WorkCategories.work_id).all()
            for work in cat_works:
                work_db = db.session.query(Works).filter(Works.work_id == work.work_id).first()
                w_no = work_db.work_id
                status_id = WorkStatuses.query.filter(WorkStatuses.work_id == w_no).first().status_id
                if ParticipationStatuses.query.filter(ParticipationStatuses.status_id == status_id).first().status_name \
                        in stat:
                    works[w_no] = work_info(w_no, analysis_info=True)
                    if works[w_no]['reg_tour'] is not None:
                        works[w_no]['pre_ana'] = get_pre_analysis(w_no)
                        works[w_no]['rk'], works[w_no]['ana_res'] = get_analysis(w_no)
                        wks.append(works[w_no])
    else:
        works = {}
        cat_works = db.session.query(WorkCategories).filter(WorkCategories.cat_id == int(cat_id)
                                                            ).order_by(WorkCategories.work_id).all()
        for work in cat_works:
            work_db = db.session.query(Works).filter(Works.work_id == work.work_id).first()
            w_no = work_db.work_id
            status_id = WorkStatuses.query.filter(WorkStatuses.work_id == w_no).first().status_id
            if ParticipationStatuses.query.filter(ParticipationStatuses.status_id == status_id).first().status_name \
                    in stat:
                works[w_no] = work_info(w_no, analysis_info=True)
                if works[w_no]['reg_tour'] is not None:
                    works[w_no]['pre_ana'] = get_pre_analysis(w_no)
                    works[w_no]['rk'], works[w_no]['ana_res'] = get_analysis(w_no)
                    wks.append(works[w_no])
    return wks


def get_pre_analysis(work_id):
    pre = dict()
    pre_ana = db.session.query(PreAnalysis).filter(PreAnalysis.work_id == int(work_id)).first()
    if pre_ana is not None:
        pre['good_work'] = pre_ana.good_work
        pre['research'] = pre_ana.research
        pre['has_review'] = pre_ana.has_review
        pre['rev_type'] = pre_ana.rev_type
        pre['pushed'] = pre_ana.pushed
        pre['work_comment'] = pre_ana.work_comment
        pre['rev_comment'] = pre_ana.rev_comment
    else:
        pre = None
    if pre == {}:
        pre = None
    return pre


def get_analysis(work_id, internal=None):
    rk = 0
    analysis = dict()
    if not internal:
        analysis_db = db.session.query(RevAnalysis).filter(RevAnalysis.work_id == work_id).all()
    else:
        analysis_db = db.session.query(InternalAnalysis).filter(InternalAnalysis.review_id == work_id).all()
    values_db = db.session.query(RevCritValues)
    criteria = db.session.query(RevCriteria)
    if analysis_db is not None:
        for criterion in analysis_db:
            crit = dict()
            crit['val_id'] = criterion.value_id
            crit['val_name'] = values_db.filter(RevCritValues.value_id == crit['val_id']).first().value_name
            analysis[criterion.criterion_id] = crit
            val_rk = values_db.filter(RevCritValues.value_id == crit['val_id']).first().weight
            cr_rk = criteria.filter(RevCriteria.criterion_id == criterion.criterion_id).first().weight
            c_v_rk = val_rk * cr_rk
            rk += c_v_rk
    else:
        analysis = None
    if analysis == {}:
        analysis = None
    return rk, analysis


def analysis_results():
    analysis_res = dict()
    criteria = db.session.query(RevCriteria).all()
    rev_ana = db.session.query(RevAnalysis)
    cats = db.session.query(Categories).all()
    for cat in cats:
        cat_works = get_works(cat.cat_id, 2)
        analysis_res.update(cat_works)
    for work in analysis_res.keys():
        if work in [w.work_id for w in RevAnalysis.query.all()]:
            for criterion in criteria:
                if criterion.criterion_id in \
                        [w.criterion_id for w in RevAnalysis.query.filter(RevAnalysis.work_id == work).all()]:
                    val = rev_ana.filter(RevAnalysis.work_id == work)
                    value = val.filter(RevAnalysis.criterion_id == criterion.criterion_id).first().value_id
                    analysis_res[work].update({criterion.criterion_id: value})
    crit_vals = get_criteria(curr_year)
    for work in analysis_res.keys():
        rk = 0
        if 'analysis' in analysis_res[work].keys() and analysis_res[work]['analysis'] is True:
            for key in analysis_res[work].keys():
                if key in crit_vals.keys():
                    rk += crit_vals[key]['weight'] * crit_vals[key]['values'][analysis_res[work][key]]['val_weight']
            analysis_res[work]['ana_rk'] = rk
    return analysis_res


def analysis_nums():
    c, cats = categories_info()
    ana_nums = []
    all_stats = {'regionals': 0,
                 'analysed': 0}
    works_db = WorkCategories.query.join(Works, WorkCategories.work_id == Works.work_id) \
        .join(WorkStatuses, WorkCategories.work_id == WorkStatuses.work_id) \
        .filter(Works.reg_tour != 0).filter(WorkStatuses.status_id >= 2)
    for cat in cats:
        cat_works = works_db.filter(WorkCategories.cat_id == cat['id']).all()
        analysed = set(w.work_id for w in RevAnalysis.query.all() if w.work_id in [c.work_id for c in cat_works])
        analysed.update(w.work_id for w in PreAnalysis.query.filter(PreAnalysis.has_review == 0).all()
                        if w.work_id in cat_works)
        cat_ana = {'cat_id': cat['id'], 'cat_name': cat['name'], 'analysed': len(analysed),
                   'regional_applied': len(cat_works)}
        cat_ana['left'] = cat_ana['regional_applied'] - cat_ana['analysed']
        ana_nums.append(cat_ana)
        all_stats['regionals'] += cat_ana['regional_applied']
        all_stats['analysed'] += cat_ana['analysed']
    all_stats['regions'] = len(set(w.reg_tour for w in Works.query.all()))
    all_stats['left'] = all_stats['regionals'] - all_stats['analysed']
    return ana_nums, all_stats


def check_analysis(cat_id):
    cat_works = [w.work_id for w in Works.query.select_from(WorkCategories
                                                            ).join(Works, Works.work_id == WorkCategories.work_id
                                                                   ).filter(WorkCategories.cat_id == cat_id).all()
                 if w.reg_tour is not None]
    for work in cat_works:
        if work not in [w.work_id for w in PreAnalysis.query.all()]:
            return True
        else:
            if PreAnalysis.query.filter(PreAnalysis.work_id == work).first().has_review is True:
                if len(RevAnalysis.query.filter(RevAnalysis.work_id == work).all()
                       ) != len(
                    RevCriteria.query.filter(RevCriteria.year == datetime.datetime.strptime(str(curr_year), '%Y').date()
                                             ).all()):
                    return True
    return False


def no_fee_nums():
    cats_no, cats = categories_info()
    total = 0
    all_no_fee = []
    for cat in cats:
        works = get_works_no_fee(cat['id'])
        cat['works'] = ', '.join([str(w) for w in works.keys()])
        total += len(works)
        a = [{'cat_name': cat['name'], 'work_id': w} for w in works.keys()]
        all_no_fee.extend(a)

    df = pd.DataFrame(data=all_no_fee)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/no_fee_works' + str(curr_year) + '.xlsx') as writer:
        df.to_excel(writer, sheet_name='Топ')
    return total, cats


def application_2_tour(appl):
    application = {'id': appl, 'works': [work_info(w.work_id, w_payment_info=True, appl_info=True) for w
                                         in Applications2Tour.query.filter(Applications2Tour.appl_no == appl).all()],
                   'participants': []}
    org = Organisations.query \
        .join(OrganisationApplication, Organisations.organisation_id == OrganisationApplication.organisation_id) \
        .filter(OrganisationApplication.appl_no == application['id']).first()

    if org is not None:
        application['organisation'] = org.name
        application['city'] = org.city
    else:
        application['organisation'] = ''
        application['city'] = ''
    for part in ParticipantsApplied.query.filter(ParticipantsApplied.appl_id == appl).all():
        part_db = db.session.query(ParticipantsApplied).filter(ParticipantsApplied.participant_id == part.participant_id
                                                               ).first()
        participant = {'id': part_db.participant_id, 'last_name': part_db.last_name, 'first_name': part_db.first_name,
                       'patronymic_name': part_db.patronymic_name, 'class': part_db.participant_class,
                       'role': part_db.role}
        p_name = (participant['last_name'] + ' ' + participant['first_name'] + ' ' + participant[
            'patronymic_name']).strip()
        if participant['id'] in [p.participant_id for p in Discounts.query.all()]:
            disc = db.session.query(Discounts).filter(Discounts.participant_id == participant['id']).first()
            participant['fee'] = disc.payment
            participant['format'] = disc.participation_format
        else:
            participant['fee'] = fee
            participant['format'] = 'face-to-face'
        if participant['id'] in [p.participant for p in PaymentRegistration.query.all()]:
            participant['payed'] = True
            participant['payment_id'] = PaymentRegistration.query.filter(PaymentRegistration.participant ==
                                                                         participant['id']).first().payment_id
        else:
            participant['payed'] = False
        if participant['id'] in [d.participant_id for d in Discounts.query.filter(Discounts.payment == 0).all()]:
            participant['payed'] = True
        else:
            pass
        application['participants'].append(participant)
    return application


def payment_info(payment_id):
    payment = db.session.query(BankStatement).filter(BankStatement.payment_id == int(payment_id)).first()
    date = datetime.datetime.strftime(payment.date, '%d.%m.%Y')
    payment_reg = db.session.query(PaymentRegistration)
    p_discounts = [p.participant_id for p in Discounts.query.all()]
    w_discounts = [p.work_id for p in Discounts.query.all()]
    works = {w.work_id: w.reg_tour for w in Works.query.all()}
    pays_for = {p.participant: p.for_work for p in payment_reg
    .filter(PaymentRegistration.payment_id == payment.payment_id).all()}
    all_part = [p.participant_id for p in ParticipantsApplied.query.all()]

    for p, t in pays_for.items():
        if t is True:
            if p not in works.keys():
                d = True
            else:
                d = False
        else:
            if p not in all_part:
                d = True
            else:
                d = False
        if d is True:
            db.session.query(PaymentRegistration).filter(PaymentRegistration.payment_id == payment_id) \
                .filter(PaymentRegistration.participant == p).delete()
            db.session.commit()
            pays_for = {p.participant: p.for_work for p in payment_reg
            .filter(PaymentRegistration.payment_id == payment.payment_id).all()}
        else:
            pass

    remainder = payment.debit
    if payment.payment_id in [p.payment_id for p in payment_reg.all()]:
        for participant in pays_for.keys():
            if participant in p_discounts:
                disc = db.session.query(Discounts).filter(Discounts.participant_id == participant).first()
                payed = disc.payment
            elif participant in w_discounts:
                disc = db.session.query(Discounts).filter(Discounts.work_id == participant).first()
                payed = disc.payment
            elif participant in works.keys() and works[participant] is not None:
                payed = tour_fee
            else:
                payed = fee
            remainder -= payed

    if remainder % 1 == 0:
        remainder = str(int(remainder)) + ' р.'
    else:
        remainder = str(remainder).replace('.', ',') + ' р.'
    if payment.debit % 1 == 0:
        debit = str(int(payment.debit)) + ' р.'
    else:
        debit = str(payment.debit).replace('.', ',') + ' р.'

    payees_list = []
    for part, t in pays_for.items():
        if t is True:
            a = 'работа ' + str(part)
        else:
            p = ParticipantsApplied.query.filter(ParticipantsApplied.participant_id == part).first()
            name = p.last_name + ' ' + p.first_name + ' ' + p.patronymic_name
            a = 'участник ' + str(part) + ' ' + name
        payees_list.append(a)
    if len(payees_list) > 0:
        payees = '; '.join(payees_list)
    else:
        payees = 'платеж еще никому не назначен'

    pay = {'payment_id': payment.payment_id, 'date': date, 'order_id': payment.order_id,
           'debit': debit, 'organisation': payment.organisation, 'tin': payment.tin, 'bic': payment.bic,
           'bank_name': payment.bank_name, 'account': payment.account, 'comment': payment.payment_comment,
           'remainder': remainder, 'payees': payees}
    return pay


def statement_info(payment_list):
    statement = []
    payment_reg = db.session.query(PaymentRegistration)
    for p in payment_list:
        payment = db.session.query(BankStatement).filter(BankStatement.payment_id == p).first()
        remainder = payment.debit
        if payment.payment_id in [p.payment_id for p in payment_reg.all()]:
            for participant in [p.participant for p
                                in payment_reg.filter(PaymentRegistration.payment_id == payment.payment_id).all()]:
                if participant in [p.participant_id for p in Discounts.query.all()]:
                    disc = db.session.query(Discounts).filter(Discounts.participant_id == participant).first()
                    payed = disc.payment
                elif participant in [p.work_id for p in Discounts.query.all()]:
                    disc = db.session.query(Discounts).filter(Discounts.work_id == participant).first()
                    payed = disc.payment
                elif participant in [w.work_id for w in Works.query.all()] \
                        and Works.query.filter(Works.work_id == participant).first().reg_tour is not None:
                    payed = tour_fee
                else:
                    payed = fee
                remainder -= payed
        # elif payment.payment_id in [p.payment_id for p in YaisWorkPayment.query.all()]:
        #     payed = 4940 * len(YaisWorkPayment.query.filter(YaisWorkPayment.payment_id == payment.payment_id).all())
        #     remainder = payment.debit - payed
        if payment.payment_id in [p.payment_id for p in PaymentTypes.query.all()]:
            payment_type = PaymentTypes.query.filter(PaymentTypes.payment_id == payment.payment_id) \
                .first().payment_type
        else:
            payment_type = None
        remainder = str(int(remainder))
        date = datetime.datetime.strftime(payment.date, '%d.%m.%Y')
        if payment.debit % 1 == 0:
            debit = str(int(payment.debit))
        else:
            debit = str(payment.debit).replace('.', ',')
        if payment.alternative is True:
            pay = {'payment_id': payment.payment_id, 'date': date, 'order_id': None,
                   'debit': debit, 'organisation': None, 'tin': None, 'bic': None,
                   'bank_name': None, 'account': None, 'comment': payment.payment_comment,
                   'remainder': remainder, 'payment_type': payment_type, 'alternative': payment.alternative,
                   'alternative_comment': payment.alternative_comment}
        else:
            pay = {'payment_id': payment.payment_id, 'date': date, 'order_id': payment.order_id,
                   'debit': debit, 'organisation': payment.organisation, 'tin': payment.tin, 'bic': payment.bic,
                   'bank_name': payment.bank_name, 'account': payment.account, 'comment': payment.payment_comment,
                   'remainder': remainder, 'payment_type': payment_type, 'alternative': None,
                   'alternative_comment': None}
        statement.append(pay)
    return statement


def document_set():
    document = Document()

    style = document.styles['Header']
    style.font.name = 'Calibri Light'
    style.font.size = Pt(16)
    style.font.bold = True

    style = document.styles['Heading 1']
    style.font.name = 'Calibri Light'
    style.font.size = Pt(16)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.left_indent = Pt(0)

    style = document.styles['Normal']
    style.font.name = 'Calibri Light'
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.left_indent = Pt(30)

    # style = document.styles['Normal']
    # style.font.name = 'Calibri Light'
    # style.font.size = Pt(14)
    # style.font.color.rgb = RGBColor(0, 0, 0)
    # style.paragraph_format.left_indent = Pt(36)

    return document


def write_work_date(cat_id, work_id, day):
    cat_id = int(cat_id)
    work_cat = cat_id
    ordered = [c.cat_id for c in ReportOrder.query.filter(ReportOrder.report_day == day).all()]

    if cat_id in [c.cat_id for c in CategoryUnions.query.all()]:
        union = CategoryUnions.query.filter(CategoryUnions.cat_id == cat_id).first().union_id
        cats = [c.cat_id for c in CategoryUnions.query.filter(CategoryUnions.union_id == union).all()]
        union = True
    else:
        cats = [cat_id]
        union = False

    if union is True:
        orders = []
        for cat in cats:
            orders.extend([w.order for w in ReportOrder.query.filter(ReportOrder.cat_id ==
                                                                     int(cat)).filter(ReportOrder.report_day ==
                                                                                      day).all()])
    else:
        orders = [w.order for w in ReportOrder.query.filter(ReportOrder.cat_id ==
                                                            int(cat_id)).filter(ReportOrder.report_day == day).all()]

    order = 1
    for cat_id in cats:
        if int(cat_id) in ordered:
            last_order = max(orders) + 1
            if last_order > order:
                order = last_order

    if work_id in [w.work_id for w in ReportOrder.query.all()]:
        db.session.query(ReportOrder).filter(ReportOrder.work_id == work_id
                                             ).update({ReportOrder.report_day: day,
                                                       ReportOrder.order: order,
                                                       ReportOrder.cat_id: work_cat})
        db.session.commit()
    else:
        o = ReportOrder(work_id, day, order, work_cat)
        db.session.add(o)
        db.session.commit()
    return 'done'


def get_responsibility(responsibility_id):
    responsibility_id = int(responsibility_id)
    resp_db = db.session.query(Responsibilities).filter(Responsibilities.responsibility_id == responsibility_id).first()
    assignees = [get_org_info(u.user_id) for u
                 in ResponsibilityAssignment.query
                 .filter(ResponsibilityAssignment.responsibility_id == responsibility_id).all()]
    assignees_ids = [u['user_id'] for u in assignees]
    responsibility = {'id': resp_db.responsibility_id, 'name': resp_db.name, 'description': resp_db.description,
                      'assignees': assignees, 'assignees_ids': assignees_ids}
    return responsibility


def make_pages(length, data, page):
    if length == 'all':
        return 1, data
    if len(data) == 1:
        return 1, data
    length = int(length)
    page = int(page)
    k = len(data) // length
    a = len(data) % length
    if page > k:
        chunk = data[length * k:length * k + a]
    else:
        page -= 1
        chunk = [data[length * i:length * (i + 1)] for i in range(k)][page]
    if a > 0:
        n = k + 1
    else:
        n = k
    return n, chunk


# САЙТ
# Главная страница
@app.route('/')
def main_page():
    renew_session()
    news = all_news()
    if 'access' in session.keys():
        access = session['access']
    else:
        access = 0
    access_list = [i for i in access_types.keys() if access_types[i] <= access]
    if 'type' in session.keys() and session['type'] in ['admin', 'org', 'manager']:
        without_cat = len([w.work_id for w in WorkCategories.query.filter(WorkCategories.cat_id == 0).all()])
    else:
        without_cat = []
    return render_template('main.html', news=news, access_list=access_list, without_cat=without_cat)


@app.route('/photo_archive')
def photo_archive():
    return render_template('photo_archive.html')


@app.route('/no_access', defaults={'url': '', 'message': None})
@app.route('/no_access/<url>/<message>')
def no_access(url, message):
    if type(url) == list:
        u = '/'.join(url)
    else:
        u = url
    return render_template('no_access.html', message=message, url=u)


@app.route('/secretary_reminder')
def secretary_reminder():
    access = check_access(5)
    if access is not True:
        return access
    return render_template('info_pages/secretaries_info/secretary_reminder.html')


@app.route('/secretary_job')
def secretary_job():
    return render_template('info_pages/secretaries_info/secretary_job.html')


# Страница авторизации
@app.route('/login', defaults={'wrong': None})
@app.route('/login/<wrong>')
def login(wrong):
    if 'user_id' in session.keys():
        return redirect(url_for('.main_page'))
    if request.referrer is not None:
        if request.referrer == request.url_root:
            url = ''
        else:
            url = request.referrer.replace(request.url_root, '').strip('/').split('/')
    else:
        url = ''
    if 'logging' in request.referrer:
        url = ''
    return render_template('registration, logging and applications/login.html', wrong=wrong, url=url)


# Страница регистрации на сайте
@app.route('/register', defaults={'message': None})
@app.route('/register/<message>')
def register(message):
    return render_template('registration, logging and applications/registration_form.html', message=message)


# Обработка данных формы регистрации на сайте
@app.route('/registration_res', methods=['POST'])
def registration_res():
    # Извлечение данных формы
    user = personal_info_form()
    # Проверка существования email и номера телефона в уже зарегистрированных пользователях.
    # При наличии пользователя с такими данными выводится ошибка через переменную exists.
    if user['email'] in [user.email for user in Users.query.all()]:
        return redirect(url_for('.register', message='email'))
    elif user['tel'] in [user.tel for user in Users.query.all()]:
        return redirect(url_for('.register', message='tel'))
    # Извлечение из формы и шифрование пароля
    user['password'] = encrypt(request.form['password'])
    user['user_type'] = 'user'
    user['approved'] = False
    # Запись полученных данных пользователя в БД, таблица users
    write_user(user)
    # Запись сессии пользователя
    session['user_id'] = db.session.query(Users).filter(Users.email == user['email']).first().user_id
        # Отправка письма для подтверждения регистрации
    send_email(user['email'])
    # Вывод страницы с информацией профиля
    return redirect(url_for('.profile_info', message='first_time'))


@app.route('/password_reset_page')
def password_reset_page():
    query = request.values.get('query', str)
    message = request.values.get('message', str)
    return render_template('registration, logging and applications/reset_password.html', query=query, message=message)


@app.route('/reset_password', methods=['POST'])
def reset_password():
    user_got = request.form['user']
    user = find_user(user_got)
    if user is None:
        return redirect(url_for('.password_reset_page', message='wrong_user', query=user_got))
    else:
        reset_key = ''.join([random.choice(string.ascii_lowercase + string.digits) for _ in range(50)])
        r = PassworsResets(user.user_id, datetime.datetime.now(), reset_key)
        db.session.add(r)
        db.session.commit()
        link = request.url_root + 'new_password/' + str(user.user_id) + '/' + reset_key

        # Генерируем HTML-содержимое письма
        html_body = render_template('mails/user_management/mail_reset_password.html', link=link)

        # Параметры письма
        subject = 'Сброс пароля'
        sender = 'info@vernadsky.info'  # адрес, от которого отправляем

        # Получаем сервис Gmail API и формируем сообщение
        service = get_service()
        message = create_message_html(
            sender=sender,
            to=user.email,
            subject=subject,
            html_body=html_body
        )

        # Отправляем письмо
        try:
            send_message(service, "me", message)
            return redirect(url_for('.login', wrong='sent'))
        except BaseException:
            return redirect(url_for('.login', wrong='mail_failed'))


# Обработка данных формы авторизации
@app.route('/logging', defaults={'url': ''})
@app.route('/logging/<url>')
def logging(url):
    try:
        if 'team_application' in url:
            url = 'team_application'
        elif '[' in url:
            url = json.loads(url.replace("'", "\""))
    except json.decoder.JSONDecodeError:
        url = ''
    if type(url) == list:
        if 'change_pwd' in url:
            u = ''
        else:
            u = '/'.join(url)
    else:
        u = url
    if 'logging' in url:
        u = ''
    # Извлечение данных формы
    user_got = request.values.get('user', str)
    pwd = request.values.get('password', str)
    password = pwd
    user = find_user(user_got)
    if user is None:
        return render_template('registration, logging and applications/login.html', wrong='user')
    # Проверка соответствия пароля записи в БД. Если совпал, записываем сессию пользователя
    else:
        if decrypt(user.password) == password:
            app.permanent_session_lifetime = datetime.timedelta(hours=1)
            session.permanent = True
            session['user_id'] = user.user_id
        else:
            # Если пароль не совпал, выводим страницу авторизации с описанием ошибки
            return redirect(url_for('.login', wrong='password'))
        user = db.session.query(Users).filter(Users.user_id == session['user_id']).first()
        user.last_login = datetime.datetime.now()
        db.session.commit()
        renew_session()
        if u == '' or u == [''] or u == '''['']''' or u == request.url_root or 'http' in u:
            return redirect(url_for('.main_page'))
        else:
            return redirect(request.url_root + u)


# Выход из учетной записи
@app.route('/logout')
def logout():
    # Удаление сессии пользователя
    session.pop('user_id', None)
    session.pop('type', None)
    session.pop('profile', None)
    session.pop('secretary', None)
    session.pop('supervisor', None)
    session.pop('cat_id', None)
    session.pop('approved', None)
    session.pop('application', None)
    session.pop('access', None)
    # Перенаправление на главную страницу
    return redirect(url_for('main_page'))


# Страница подтверждения регистрации (из email)
@app.route('/approve/<user_id>', defaults={'page': 'main'})
@app.route('/approve/<user_id>/<page>')
def approve(user_id, page):
    # Изменение статуса пользователя на "подтвержден"
    user = db.session.query(Users).filter(Users.user_id == int(user_id)).first()
    user.approved = True
    db.session.commit()
    if page == 'adm':
        return redirect(url_for('.user_page', user=user_id))
    else:
        renew_session()
        # Перенаправление на главную страницу
        return redirect(url_for('.main_page'))


# Страница с информацией профиля
@app.route('/profile_info', defaults={'message': None})
@app.route('/profile_info/<message>')
def profile_info(message):
    access = check_access(1)
    if access is not True:
        return access
    user = get_user_info(session['user_id'])
    profile = get_profile_info(session['user_id'])
    if profile['born'] is not None:
        profile['born'] = profile['born'].strftime('%d.%m.%Y')
    return render_template('registration, logging and applications/profile_info.html', profile=profile, user=user,
                           access=access, message=message)


# Форма изменения информации пользователя (email, телефон, ФИО, дата рождения)
@app.route('/edit_user/<user_id>', defaults={'message': None})
@app.route('/edit_user/<user_id>/<message>')
def edit_user(user_id, message):
    access = check_access(2)
    if access is not True:
        return access
    # Получение информации текущего пользователя из БД
    user = get_user_info(int(user_id))
    url = request.referrer.replace(request.url_root, '').strip('/').split('/')
    # Вывод формы изменения информации пользователя с предзаполненными из БД полями
    return render_template('registration, logging and applications/edit_user.html', user=user, message=message, url=url)


# Обработка информации из формы изменения информации пользователя
@app.route('/edited_user/<url>', methods=['POST'])
def edited_user(url):
    # Получение новых данных пользователя из формы и запись их в БД
    user_info = personal_info_form()
    message = write_user(user_info)
    if message == 'email' or message == 'tel':
        return redirect(url_for('.edit_user', user_id=user_info['user_id'], message=message))
    if url == 'rofile_info':
        return redirect(url_for('.profile_info'))
    else:
        return redirect(url_for('.users_list'))


# Форма редактирования информации профиля
@app.route('/edit_profile/<user_id>')
def edit_profile(user_id):
    access = check_access(2)
    if access is not True:
        return access
    # Извлечение информации профиля из БД (если она заполнен)
    profile = get_profile_info(user_id)
    if profile['born'] is not None:
        profile['born'] = profile['born'].strftime('%Y-%m-%d')
    renew_session()
    # Вывод страницы профиля с информацией пользователя и профиля из БД
    return render_template('registration, logging and applications/edit_profile.html', profile=profile)


# Обработка данных формы редактирования профиля
@app.route('/write_profile', methods=['POST'])
def write_profile():
    if 'user_id' in request.form.keys():
        user_id = int(request.form['user_id'])
    else:
        user_id = int(session['user_id'])
    if 'occupation' in request.form:
        occupation = request.form['occupation']
    else:
        occupation = None
    if 'place_of_w' in request.form:
        place_of_w = request.form['place_of_w']
        if place_of_w == 'None':
            place_of_w = None
    else:
        place_of_w = None
    if 'place_of_work' in request.form:
        place_of_work = request.form['place_of_work']
    else:
        place_of_work = place_of_w
    if 'involved' in request.form:
        inv = request.form['involved']
    else:
        inv = None
    if 'school' in request.form:
        involved = request.form['school']
        if involved == 'None':
            involved = None
    else:
        involved = inv
    if 'grade' in request.form:
        grade = request.form['grade']
    else:
        grade = None
    if 'year' in request.form:
        year = request.form['year']
    else:
        year = None
    vk = re.sub(r'^vk.com/|^https://vk.com/', '', request.form['vk'])
    if 'telegram' in request.form:
        tg = re.sub(r'https://t.me/|@', '', request.form['telegram'])
    else:
        tg = None
    if 'vernadsky_username' in request.form:
        username = request.form['vernadsky_username']
    else:
        username = None
    if 'born' in request.form.keys():
        born = datetime.datetime.strptime(request.form['born'], '%Y-%m-%d').date()
    else:
        born = None

    if user_id not in [prof.user_id for prof in Profile.query.all()]:
        prof = Profile(user_id, occupation, place_of_work, involved, grade, year, vk, tg, username, born)
        db.session.add(prof)
        db.session.commit()
        if user_id == session['user_id']:
            return redirect(url_for('.team_application'))
        else:
            return redirect(url_for('.user_page', user=user_id))
    else:
        db.session.query(Profile).filter(Profile.user_id == user_id).update(
            {Profile.occupation: occupation, Profile.place_of_work: place_of_work, Profile.involved: involved,
             Profile.grade: grade, Profile.year: year, Profile.vk: vk, Profile.telegram: tg,
             Profile.vernadsky_username: username, Profile.born: born})
        db.session.commit()
        return redirect(url_for('.profile_info'))


@app.route('/new_password/<user_id>/<key>')
def new_password(user_id, key):
    user_id = int(user_id)
    now = datetime.datetime.now()
    delta = datetime.timedelta(minutes=15, seconds=0)
    if user_id in [u.user_id for u in PassworsResets.query.all()]:
        if key in [u.reset_key for u in PassworsResets.query.filter(PassworsResets.user_id == user_id).all()]:
            t = PassworsResets.query.filter(PassworsResets.user_id == user_id).filter(
                PassworsResets.reset_key == key).first().request_time
            if t + delta >= now:
                for to_del in PassworsResets.query.filter(PassworsResets.user_id == user_id).all():
                    db.session.delete(to_del)
                    db.session.commit()
                return redirect(url_for('.change_pwd', user_id=user_id, mode='reset', success=None))
    return redirect(url_for('.login', wrong='invalid_key'))


@app.route('/change_pwd/<mode>/<user_id>', defaults={'success': None})
@app.route('/change_pwd/<mode>/<user_id>/<success>')
def change_pwd(mode, user_id, success):
    if mode == 'change':
        access = check_access(2)
        if access is not True:
            return access
    return render_template('registration, logging and applications/change_pwd.html', mode=mode, success=success,
                           user_id=user_id)


@app.route('/new_pwd', methods=['GET'])
def new_pwd():
    new = request.values.get('new_password', str)
    confirm = request.values.get('confirm_password', str)
    user_id = request.values.get('user_id', int)
    user = db.session.query(Users).filter(Users.user_id == user_id).first()
    if 'old_password' in request.values.keys():
        mode = 'change'
        old = request.values.get('old_password', str)
        old_check = decrypt(user.password)
        if old == old_check:
            validate = True
        else:
            validate = False
    elif 'valid_key' in request.values.keys():
        validate = True
        mode = 'reset'
    else:
        mode = 'reset'
        validate = False
    if validate is True:
        if new == confirm:
            user.password = encrypt(new)
            db.session.commit()
            success = True
        else:
            success = 'unmatched'
    else:
        success = 'wrong_old'
    if mode == 'change':
        return redirect(url_for('.change_pwd', success=success, mode=mode, user_id=user_id))
    else:
        return redirect(url_for('.login', wrong='password_changed'))


@app.route('/change_user_password/<user_id>', defaults={'message': None})
@app.route('/change_user_password/<user_id>/<message>')
def change_user_password(user_id, message):
    access = check_access(8)
    if access is not True:
        return access
    return render_template('user_management/change_user_password.html', user=user_id, message=message)


@app.route('/new_user_password')
def new_user_password():
    access = check_access(8)
    if access is not True:
        return access
    new = request.values.get('new_password', str)
    confirm = request.values.get('confirm_password', str)
    user_id = int(request.values.get('user_id', str))
    user = db.session.query(Users).filter(Users.user_id == user_id).first()
    if new == confirm:
        user.password = encrypt(new)
        db.session.commit()
        message = 'password_changed'
    else:
        return redirect(url_for('.change_user_password', user_id=user_id, message='unmatched'))
    renew_session()
    return redirect(url_for('.user_page', user=user_id, message=message))


@app.route('/admin')
def admin():
    access = check_access(8)
    if access is not True:
        return access
    renew_session()
    return render_template('admin.html')


@app.route('/categories')
def categories_list():
    cats_count, cats = categories_info()
    cats_ids = [categ['id'] for categ in cats]
    with_secretary = 0
    for categ in db.session.query(CatSecretaries).all():
        if categ.cat_id in cats_ids:
            with_secretary += 1
    no_secr = cats_count - with_secretary
    return render_template('categories/categories.html', cats_count=cats_count, categories=cats, no_secr=no_secr)


@app.route('/download_categories')
def download_categories():
    cats_count, cats = categories_info()
    categories = []
    for c in cats:
        if 'supervisor_id' not in c.keys():
            c['supervisor'] = ''
            c['supervisor_email'] = ''
            c['supervisor_tel'] = ''
        if 'secretary_id' not in c.keys():
            c['secretary_full'] = ''
            c['secretary_email'] = ''
            c['secretary_tel'] = ''
        if 'online_secretary_id' not in c.keys():
            c['online_secretary_full'] = ''
            c['online_secretary_email'] = ''
            c['online_secretary_tel'] = ''
        cat = {'Направление': c['direction'], 'Название секции': c['name'], 'Короткое название': c['short_name'],
               'Telegram-канал': '@' + c['tg_channel'], 'Руководитель': c['supervisor'],
               'e-mail руководителя': c['supervisor_email'], 'Телефон руководиотеля': c['supervisor_tel'],
               'Секретарь': c['secretary_full'], 'e-mail секретаря': c['secretary_email'],
               'Телефон секретаря': c['secretary_tel'],
               'Секретарь онлайн': c['online_secretary_full'], 'e-mail секретаря онлайн': c['online_secretary_email'],
               'Телефон секретаря онлайн': c['online_secretary_tel']}
        if 'dates' in c.keys():
            cat['Даты заседаний'] = c['dates']
        categories.append(cat)
    df = pd.DataFrame(data=categories)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/categories.xlsx') as writer:
        df.to_excel(writer, sheet_name='Секции ' + str(curr_year) + ' года')
    return send_file('static/files/generated_files/categories.xlsx', as_attachment=True)


@app.route('/edit_category', defaults={'cat_id': None})
@app.route('/edit_category/<cat_id>')
def edit_category(cat_id):
    if cat_id is None:
        cat_id = ''
    access = check_access(10)
    if access is not True:
        return access
    sups = get_supervisors()
    dirs = dict()
    conts = dict()
    directions = db.session.query(Directions).all()
    contests = db.session.query(Contests)
    for direct in directions:
        dir_id = direct.direction_id
        dirs[dir_id] = dict()
        dirs[dir_id]['id'] = direct.direction_id
        dirs[dir_id]['name'] = direct.dir_name
    for cont in contests:
        dir_id = cont.contest_id
        conts[dir_id] = dict()
        conts[dir_id]['id'] = cont.contest_id
        conts[dir_id]['name'] = cont.contest_name
    if cat_id is not None and cat_id != '':
        category = one_category(db.session.query(Categories).filter(Categories.cat_id == cat_id).first())
    else:
        category = None
    renew_session()
    return render_template('categories/add_category.html', supervisors=sups, directions=dirs, contests=conts,
                           category=category)


@app.route('/edited_cat', methods=['POST'])
def edited_category():
    cat_info = dict()
    cat_id = request.form['cat_id']
    if cat_id != '' and cat_id is not None:
        cat_info['cat_id'] = int(cat_id)
    else:
        cat_info['cat_id'] = None
    cat_info['cat_name'] = request.form['category_name']
    cat_info['short_name'] = request.form['short_name']
    supervisor = request.form['supervisor']
    if supervisor != 'Руководитель секции':
        cat_info['supervisor'] = int(supervisor)
    else:
        cat_info['supervisor'] = None
    cat_info['tg_channel'] = re.sub(r'https://t.me/|@', '', request.form['tg_channel'])
    cat_info['direction'] = int(request.form['direction'])
    cat_info['contest'] = int(request.form['contest'])
    cat_site_id = request.form['cat_site_id']
    if cat_site_id != '' and cat_site_id is not None:
        cat_info['cat_site_id'] = int(cat_site_id)
    else:
        cat_info['cat_site_id'] = None
    if 'drive_link' in request.form.keys():
        cat_info['drive_link'] = request.form['drive_link']
    else:
        cat_info['drive_link'] = None
    write_category(cat_info)
    renew_session()
    return redirect(url_for('.categories_list'))


@app.route('/add_categories')
def add_categories():
    access = check_access(10)
    if access is not True:
        return access
    return render_template('categories/add_categories.html')


@app.route('/many_categs', methods=['POST'])
def many_categs():
    text = request.form['text']
    cat_text = text.split('\n')
    for cat in cat_text:
        if cat != '':
            c = cat.split('\t')
            cat_info = dict()
            cat_info['cat_id'] = None
            cat_info['cat_name'] = c[0].strip('\r')
            cat_info['short_name'] = c[1].strip('\r')
            cat_info['supervisor'] = c[2].strip('\r')
            cat_info['tg_channel'] = re.sub(r'https://t.me/|@', '', c[3].strip('\r'))
            direction = c[4].strip('\r')
            if direction == 'вернак' or direction == 'Вернак':
                cat_info['direction'] = 'Конкурс им. В. И. Вернадского'
            elif direction == 'тропа' or direction == 'Тропа':
                cat_info['direction'] = 'Тропой открытий В. И. Вернадского'
            else:
                cat_info['direction'] = direction
            cat_info['contest'] = c[5].strip('\r')
            cat_info['cat_site_id'] = ''
            cat_info['drive_link'] = ''
            write_category(cat_info)
    return redirect(url_for('.categories_list'))


@app.route('/supervisors')
def supervisors():
    sups = get_supervisors()
    c, cats = categories_info()
    relevant = [cat['supervisor_id'] for cat in cats if 'supervisor_id' in cat.keys()]
    relevant.append(21)  # Добавление Свешниковой
    relevant.append(44)  # Добавление Марусяк
    renew_session()
    if 'access' in session.keys():
        access = session['access']
    else:
        access = 0
    return render_template('supervisors/supervisors.html', supervisors=sups, access=access, relevant=relevant)


@app.route('/download_supervisors')
def download_supervisors():
    sups = get_supervisors()
    c, cats = categories_info()
    relevant = [cat['supervisor_id'] for cat in cats if 'supervisor_id' in cat.keys()]
    relevant.append(21)  # Добавление Свешниковой
    relevant.append(44)  # Добавление Марусяк
    supers = [sup for sup in sups.values() if sup['id'] in relevant]
    df = pd.DataFrame(data=supers)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/supervisors.xlsx') as writer:
        df.to_excel(writer, sheet_name='Руководители секций')
    return send_file('static/files/generated_files/supervisors.xlsx', as_attachment=True)


@app.route('/edit_supervisor', defaults={'sup_id': ''})
@app.route('/edit_supervisor/<sup_id>')
def edit_supervisor(sup_id):
    access = check_access(10)
    if access is not True:
        return access
    if sup_id != '':
        supervisor = supervisor_info(sup_id)
    else:
        supervisor = None
    renew_session()
    return render_template('supervisors/add_supervisor.html', supervisor=supervisor)


@app.route('/adding_supervisor', methods=['POST'])
def edited_supervisor():
    supervisor = personal_info_form()
    supervisor['sup_info'] = request.form['supervisor_info']
    sup_id = request.form['supervisor_id'].strip()
    if sup_id != '' and sup_id is not None and '\r\n' not in sup_id:
        supervisor_id = int(sup_id)
        if supervisor_id in [sup.supervisor_id for sup in Supervisors.query.all()]:
            db.session.query(Supervisors).filter(Supervisors.supervisor_id == supervisor_id).update(
                {Supervisors.last_name: supervisor['last_name'], Supervisors.first_name: supervisor['first_name'],
                 Supervisors.patronymic: supervisor['patronymic'], Supervisors.email: supervisor['email'],
                 Supervisors.tel: supervisor['tel'], Supervisors.supervisor_info: supervisor['sup_info']})
    else:
        supervisor = Supervisors(supervisor['last_name'], supervisor['first_name'], supervisor['patronymic'],
                                 supervisor['email'], supervisor['tel'], supervisor['sup_info'])
        db.session.add(supervisor)
    db.session.commit()
    renew_session()
    return redirect(url_for('.supervisors'))


@app.route('/confirm_sup_deletion/<sup_id>')
def confirm_sup_deletion(sup_id):
    access = check_access(8)
    if access is not True:
        return access
    sup_info = supervisor_info(sup_id)
    return render_template('supervisors/confirm_supervisor_deletion.html', supervisor=sup_info)


@app.route('/delete_supervisor/<sup_id>')
def delete_supervisor(sup_id):
    supervisor = db.session.query(Supervisors).filter(Supervisors.supervisor_id == sup_id).first()
    db.session.delete(supervisor)
    db.session.commit()
    renew_session()
    return redirect(url_for('.supervisors'))


@app.route('/add_supervisors')
def add_supervisors():
    access = check_access(10)
    if access is not True:
        return access
    renew_session()
    return render_template('supervisors/add_supervisors.html')


@app.route('/many_sups', methods=['POST'])
def many_sups():
    text = request.form['text']
    sup_text = text.split('\n')
    for sup in sup_text:
        if sup != '':
            s = sup.split('\t')
            tel = re.sub(r'^8|^7|^(?=9)', '+7', ''.join([n for n in s[4] if n not in tel_unneeded]))
            supervisor = Supervisors(s[0].strip(' '), s[1].strip(' '), s[2].strip(' '), s[3].strip(' '), tel, None)
            db.session.add(supervisor)
    db.session.commit()
    sups = get_supervisors()
    renew_session()
    return redirect(url_for('.supervisors'))


# @app.route('/add_supervisors', methods=['GET'])
# def add_supervisors():
#     if check_access() < 10:
#         return redirect(url_for('.no_access'))
#     file = request.files.get('file', None)
#     # file.save(os.path.join('static/', file.txt))
#     # data = genfromtxt(file, delimiter='\t', encoding='utf-8', dtype=None, names=True).tolist()
#     # for row in data:
#     #     tel = re.sub(r'^8|^7|^(?=9)', '+7', ''.join([n for n in row[4] if n not in tel_unneeded]))
#     #     supervisor = Supervisors(row[0].strip(' '), row[1].strip(' '), row[2].strip(' '), row[3].strip(' '),
#     #                              tel, row[5].strip(' '))
#     #     db.session.add(supervisor)
#     # db.session.commit()
#     sups = get_supervisors()
#     renew_session()
#     return render_template('supervisors.html', supervisors=sups)


@app.route('/supervisor_profile/<supervisor_id>')
def supervisor_profile(supervisor_id):
    access = check_access(3)
    if access is not True:
        return access
    sup_info = supervisor_info(supervisor_id)
    renew_session()
    return render_template('supervisors/supervisor_profile.html', supervisor=sup_info, access=access)


@app.route('/team_application')
def team_application():
    access = check_access(2)
    if access is not True:
        return access
    if 'profile' not in session.keys():
        return redirect(url_for('.edit_profile', user_id=session['user_id']))
    cats_count, categs = categories_info()
    if session['user_id'] in [a.user_id for a in Application.query.filter(Application.year == curr_year).all()]:
        application = application_info('user-year', user=session['user_id'])
        # if application == {curr_year: {'role': None, 'category_1': None, 'category_2': None, 'category_3': None,
        #                                'any_category': None, 'taken_part': None, 'considered': None}}:
        #     application = application[curr_year]
    else:
        application = None
    renew_session()
    return render_template('registration, logging and applications/team_application.html', application=application,
                           categories=categs)


@app.route('/application_process', methods=['POST'])
def application_process():
    role = request.form['role']
    category_1 = request.form['category_1']
    category_2 = request.form['category_2']
    category_3 = request.form['category_3']
    if 'any_category' in request.form:
        any_cat = request.form['any_category']
        any_category = bool(any_cat)
    else:
        any_category = False
    if 'taken_part' in request.form:
        taken_part = request.form['taken_part']
    else:
        taken_part = 'not_filled'
    if session['user_id'] in [user.user_id for user in Application.query.filter(Application.year == curr_year).all()]:
        db.session.query(Application).filter(Application.user_id == session['user_id']).update(
            {Application.role: role, Application.category_1: category_1, Application.category_2: category_2,
             Application.category_3: category_3, Application.any_category: any_category,
             Application.taken_part: taken_part})
    else:
        appl_id = max([appl.appl_id for appl in Application.query.all()]) + 1
        cat_sec = Application(appl_id, session['user_id'], curr_year, role, category_1, category_2, category_3,
                              any_category,
                              taken_part, 'False')
        db.session.add(cat_sec)
    db.session.commit()
    renew_session()
    return redirect(url_for('.application_page'))


@app.route('/my_applications')
def application_page():
    access = check_access(2)
    if access is not True:
        return access
    appl_info = application_info('user', user=session['user_id'])
    renew_session()
    return render_template('registration, logging and applications/my_applications.html', application=appl_info)


@app.route('/view_applications')
def view_applications():
    access = check_access(8)
    if access is not True:
        return access
    appl = application_info('year', user=session['user_id'])
    users = all_users()
    renew_session()
    secretaries = [a['user_id'] for a in appl.values() if a['role'] == 'secretary']
    volunteers = [a['user_id'] for a in appl.values() if a['role'] == 'volunteer']
    source = [{'user_id': u.user_id, 'involved': u.involved, 'occupation': u.occupation, 'grade': u.grade} for u in Profile.query
    .join(Application, Profile.user_id == Application.user_id).filter(Application.year == curr_year).all()]
    msu_school = [a['user_id'] for a in source if a['involved'] == 'MSU_School' and a['occupation'] == 'scholar']
    lyceum = [a['user_id'] for a in source if a['involved'] == '1553' and a['occupation'] == 'scholar']
    graduates = [a['user_id'] for a in source if a['occupation'] == 'student']
    unseen = [a['user_id'] for a in appl.values() if a['considered'] == 'False']
    for s in source:
        appl[s['user_id']]['occupation'] = s['occupation']
        appl[s['user_id']]['involved'] = s['involved']
        appl[s['user_id']]['grade'] = s['grade']
    return render_template('application management/view_applications.html', applications=appl, year=curr_year,
                           users=users, secretaries=secretaries, volunteers=volunteers, msu_school=msu_school,
                           lyceum=lyceum, graduates=graduates, unseen=unseen)


@app.route('/download_team_contacts')
def download_team_contacts():
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    team = [get_user_info(u.user_id) for u in Application.query.filter(Application.year == curr_year).all()]
    with open('static/files/generated_files/team_vcards_' + str(curr_year) + '.vcf', 'w') as file:
            file.write('')
    for p in team:
        vcard = vobject.vCard()
        o = vcard.add('fn')
        o.value = p['last_name'] + ' ' + p['first_name'] + ' ' + p['patronymic']

        o = vcard.add('n')
        o.value = vobject.vcard.Name(family=p['last_name'], given=p['first_name'], additional=p['patronymic'])

        o = vcard.add('tel')
        o.type_param = "cell"
        o.value = p['tel']

        o = vcard.add('email')
        o.type_param = "work"
        o.value = p['email']

        o = vcard.add('url')
        o.value = 'org.vernadsky.info/user_page/' + str(p['user_id'])

        with open('static/files/generated_files/team_vcards_' + str(curr_year) + '.vcf', 'a') as file:
            file.write(vcard.serialize())
    return send_file('static/files/generated_files/team_vcards_' + str(curr_year) + '.vcf', as_attachment=True)


@app.route('/download_contact/<user>')
def download_contact(user):
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    user_info = get_user_info(int(user))

    vcard = vobject.vCard()
    o = vcard.add('fn')
    o.value = user_info['last_name'] + ' ' + user_info['first_name'] + ' ' + user_info['patronymic']

    o = vcard.add('n')
    o.value = vobject.vcard.Name(family=user_info['last_name'], given=user_info['first_name'], additional=user_info['patronymic'])

    o = vcard.add('tel')
    o.type_param = "cell"
    o.value = user_info['tel']

    o = vcard.add('email')
    o.type_param = "work"
    o.value = user_info['email']

    o = vcard.add('url')
    o.value = 'org.vernadsky.info/user_page/' + str(user_info['user_id'])

    with open('static/files/generated_files/one_vcard.vcf', 'w') as file:
        file.write(vcard.serialize())
    return send_file('static/files/generated_files/one_vcard.vcf', as_attachment=True)


@app.route('/one_application/<year>/<user>')
def see_one_application(year, user):
    user = int(user)
    access = check_access(8)
    if access is not True:
        return access
    application = application_info('user-year', user=user, year=year)
    user_info = get_user_info(user)
    profile = get_profile_info(user)
    if profile['born'] is not None:
        profile['born'] = profile['born'].strftime('%d.%m.%Y')
    cats_count, cats = categories_info()
    renew_session()
    return render_template('application management/one_application.html', application=application, year=curr_year,
                           user=user_info,
                           profile=profile, categories=cats)


@app.route('/confirm_application_deletion/<year>/<user>')
def confirm_application_deletion(year, user):
    application = application_info('user-year', user, year)
    user_info = get_user_info(user)
    return render_template('application management/confirm_application_deletion.html', application=application,
                           year=year, user=user_info)


@app.route('/manage_application/<year>/<user>/<action>', defaults={'page': 'all'})
@app.route('/manage_application/<year>/<user>/<action>/<page>')
def manage_application(year, user, action, page):
    access = check_access(8)
    if access is not True:
        return access
    appl_db = db.session.query(Application).filter(Application.user_id == user).filter(Application.year == year).first()
    user_db = db.session.query(Users).filter(Users.user_id == user).first()
    if action == 'accept':
        appl_db.considered = 'True'
        if user_db.user_type == 'user':
            user_db.user_type = 'team'
    elif action == 'decline':
        appl_db.considered = 'False'
    elif action == 'await':
        appl_db.considered = 'in_process'
    elif action == 'delete':
        db.session.delete(appl_db)
    db.session.commit()
    renew_session()
    if page == 'all':
        return redirect(url_for('.view_applications'))
    else:
        return redirect(url_for('.see_one_application', year=year, user=user))


@app.route('/assign_category/<user>/<category>')
def assign_category(user, category):
    access = check_access(8)
    if access is not True:
        return access
    user_info = get_user_info(user)
    cats_count, cats = categories_info(category)
    renew_session()
    return render_template('application management/confirm_assignment.html', user=user_info, category=cats)


@app.route('/assign_online_category/<user>/<category>')
def assign_online_category(user, category):
    access = check_access(8)
    if access is not True:
        return access
    user_info = get_user_info(user)
    cats_count, cats = categories_info(category)
    renew_session()
    return render_template('application management/confirm_online_assignment.html', user=user_info, category=cats)


@app.route('/confirm_assignment/<user>/<category>')
def confirm_assignment(user, category):
    access = check_access(8)
    if access is not True:
        return access
    user = int(user)
    category = int(category)
    if category in [cat.cat_id for cat in CatSecretaries.query.all()]:
        cat = db.session.query(CatSecretaries).filter(CatSecretaries.cat_id == category).first()
        cat.secretary_id = user
    else:
        cat_sec = CatSecretaries(category, user)
        db.session.add(cat_sec)
    db.session.commit()
    renew_session()
    return redirect(url_for('.view_applications'))


@app.route('/confirm_online_assignment/<user>/<category>')
def confirm_online_assignment(user, category):
    access = check_access(8)
    if access is not True:
        return access
    user = int(user)
    category = int(category)
    if category in [cat.cat_id for cat in OnlineSecretaries.query.all()]:
        cat = db.session.query(OnlineSecretaries).filter(OnlineSecretaries.cat_id == category).first()
        cat.secretary_id = user
    else:
        cat_sec = OnlineSecretaries(category, user)
        db.session.add(cat_sec)
    db.session.commit()
    renew_session()
    return redirect(url_for('.view_applications'))


@app.route('/users_list', defaults={'query': 'all', 'length': 50, 'page': 1})
@app.route('/users_list/<query>/<length>/<page>')
def users_list(query, length, page):
    if isinstance(query, str) and '%' in query:
        query = unquote(query)
    access = check_access(8)
    if access is not True:
        return access
    if query == 'not_found' or query == []:
        found = None
        users = None
        pages = 1
    else:
        if query == 'all' or query is None or query == []:
            pages, us = make_pages(length, [u.user_id for u in Users.query.order_by(Users.user_id.desc()).all()], page)
            users = [get_user_info(u) for u in us]
        else:
            if type(query) == str:
                query = json.loads(query)
            users = [get_user_info(u) for u in query]
            pages = 1
        found = 'ugu'
    return render_template('user_management/users_list.html', users=users, length=length, page=page, pages=pages,
                           found=found, link='users_list/all')


@app.route('/search_user', methods=['GET'])
def search_user():
    query = request.values.get('query', str)
    tel = re.sub(
        r'(^\+7|^8|^7|^9)(-|\(|\)|\s)*(?P<a>\d+)(-|\(|\)|\s)*(?P<b>\d+)(-|\(|\)|\s)*(?P<c>\d+)(-|\(|\)|\s)*(?P<d>\d+)',
        r'+7\g<a>\g<b>\g<c>\g<d>', query)
    users = []
    if tel in [u.tel for u in Users.query.all()]:
        users.extend([u.user_id for u in Users.query.filter(Users.tel == tel).order_by(Users.user_id.desc()).all()])
    try:
        query = int(query)
        users.extend([u.user_id for u in Users.query.filter(Users.user_id == query)
                     .order_by(Users.user_id.desc()).all()])
    except ValueError:
        if query == 'secretary':
            users.extend(
                [u.secretary_id for u in CatSecretaries.query.order_by(CatSecretaries.secretary_id.desc()).all()])
        elif query == 'supervisor':
            users.extend([u.user_id for u in SupervisorUser.query.order_by(SupervisorUser.user_id.desc()).all()])
        elif query in access_types.keys():
            us = []
            for val in [val for val in access_types.values() if val >= access_types[query]]:
                us.extend([u.user_id for u in Users.query
                          .filter(Users.user_type == list(access_types.keys())[list(access_types.values()).index(val)])
                          .order_by(Users.user_id.desc()).all()])
            us.sort(reverse=True)
            users.extend(us)
        if query in [u.email for u in Users.query.all()]:
            users.extend(
                [u.user_id for u in Users.query.filter(Users.email == query).order_by(Users.user_id.desc()).all()])
        users.extend([u.user_id for u in Users.query.order_by(Users.user_id.desc()).all()
                      if query.lower() == u.last_name.lower()[:len(query)]])
        users.extend([u.user_id for u in Users.query.order_by(Users.user_id.desc()).all()
                      if query.lower() == u.first_name.lower()[:len(query)]])
        users.extend([u.user_id for u in Users.query.order_by(Users.user_id.desc()).all()
                      if query.lower() == u.patronymic.lower()[:len(query)]])
    if users:
        us = set(users)
        users = [u for u in us]
    else:
        users = 'not_found'
    return redirect(url_for('.users_list', query=users, length='all', page=1))


@app.route('/user_page/<user>', defaults={'message': None})
@app.route('/user_page/<user>/<message>')
def user_page(user, message):
    access = check_access(3)
    if access is not True:
        return access
    user_info = get_user_info(user)
    profile = get_profile_info(user)
    if profile['born'] is not None:
        profile['born'] = profile['born'].strftime('%d.%m.%Y')
    cats_count, cats = categories_info()
    supers = get_supervisors()
    classes = [{'class_id': c.class_id, 'school': c.school, 'class_name': c.class_name}
               for c in SchoolClasses.query.filter(SchoolClasses.year == curr_year).all()]
    return render_template('user_management/user_page.html', user=user_info, profile=profile, categories=cats,
                           message=message, supervisors=supers, curr_year=curr_year, classes=classes)


@app.route('/assign_user_type/<user>', methods=['GET'])
def assign_user_type(user):
    renew_session()
    assign_type = request.values.get('assign_type', str)
    user_db = db.session.query(Users).filter(Users.user_id == user).first()
    user_db.user_type = assign_type
    db.session.commit()
    return redirect(url_for('.user_page', user=user))


@app.route('/remove_secretary/<user_id>/<cat_id>')
def remove_secretary(user_id, cat_id):
    access = check_access(8)
    if access is not True:
        return access
    cat_sec = CatSecretaries.query.filter(CatSecretaries.secretary_id == user_id
                                          ).filter(CatSecretaries.cat_id == cat_id).first()
    db.session.delete(cat_sec)
    db.session.commit()
    return redirect(url_for('.user_page', user=user_id))


@app.route('/remove_oline_secretary/<user_id>/<cat_id>')
def remove_oline_secretary(user_id, cat_id):
    access = check_access(8)
    if access is not True:
        return access
    cat_sec = OnlineSecretaries.query.filter(OnlineSecretaries.secretary_id == user_id
                                             ).filter(OnlineSecretaries.cat_id == cat_id).first()
    db.session.delete(cat_sec)
    db.session.commit()
    return redirect(url_for('.user_page', user=user_id))


@app.route('/category_page/<cat_id>', defaults={'errors': None})
@app.route('/category_page/<cat_id>/<errors>')
def category_page(cat_id, errors):
    category = one_category(db.session.query(Categories).filter(Categories.cat_id == cat_id).first())
    renew_session()
    need_analysis = check_analysis(cat_id=cat_id)
    works_no_fee = get_works_no_fee(cat_id)
    show_top_100 = True
    # works_1_tour = [work_info(w.work_id) for w in Works.query
    # .join(WorkCategories, Works.work_id == WorkCategories.work_id)
    # .filter(WorkCategories.cat_id == cat_id).all()]
    return render_template('categories/category_page.html', category=category, need_analysis=need_analysis,
                           errors=errors, works_no_fee=works_no_fee, show_top_100=show_top_100)
    # , works_1_tour=works_1_tour)


@app.route('/news_list')
def news_list():
    access = check_access(8)
    if access is not True:
        return access
    news = all_news()
    return render_template('news/news_list.html', news=news)


@app.route('/edit_news', defaults={'news_id': None})
@app.route('/edit_news/<news_id>')
def edit_news(news_id):
    access = check_access(8)
    if access is not True:
        return access
    if news_id == 'None' or not news_id:
        news = {'news_id': None}
    else:
        news = one_news(news_id)
    return render_template('news/edit_news.html', news=news)


@app.route('/editing_news', methods=['POST'])
def editing_news():
    renew_session()
    news = dict()
    news_id = request.form['news_id']
    news['title'] = request.form['title']
    news['content'] = request.form['content']
    news['access'] = request.form['access']
    if news_id != 'None':
        news_id = int(news_id)
        if news_id in [n.news_id for n in News.query.all()]:
            db.session.query(News).filter(News.news_id == news_id).update(
                {News.title: news['title'], News.content: news['content'],
                 News.access: news['access']})
    else:
        news['publish'] = False
        new_news = News(news['title'], news['content'], news['access'], news['publish'])
        db.session.add(new_news)
    db.session.commit()
    return redirect(url_for('.news_list'))


@app.route('/publish_news/<news_id>')
def publish_news(news_id):
    access = check_access(8)
    if access is not True:
        return access
    news = db.session.query(News).filter(News.news_id == news_id).first()
    if news.publish is True:
        news.publish = False
    elif news.publish is False:
        news.publish = True
    db.session.commit()
    return redirect(url_for('.news_list'))


@app.route('/supervisor_user/<user_id>', methods=['GET'])
def supervisor_user(user_id):
    renew_session()
    sup_id = request.values.get('user_supervisor')
    user_id = int(user_id)
    user = db.session.query(Users).filter(Users.user_id == user_id).first()
    if sup_id != 'None':
        sup_id = int(sup_id)
        supervisor = db.session.query(Supervisors).filter(Supervisors.supervisor_id == sup_id).first()
        if user_id in [u.user_id for u in SupervisorUser.query.all()]:
            user_db = db.session.query(SupervisorUser).filter(SupervisorUser.user_id == user_id).first()
            user_db.supervisor_id = supervisor.supervisor_id
            db.session.commit()
        else:
            sup_user = SupervisorUser(user.user_id, supervisor.supervisor_id)
            db.session.add(sup_user)
            db.session.commit()
    else:
        if user_id in [u.user_id for u in SupervisorUser.query.all()]:
            superv = SupervisorUser.query.filter(SupervisorUser.user_id == user_id).first().supervisor_id
            user_sup = SupervisorUser.query.filter(SupervisorUser.user_id == user_id
                                                   ).filter(SupervisorUser.supervisor_id == superv).first()
            db.session.delete(user_sup)
            db.session.commit()
    return redirect(url_for('.user_page', user=user_id))


@app.route('/tutor_user/<user_id>', methods=['GET'])
def tutor_user(user_id):
    renew_session()
    class_id = request.values.get('class_id')
    user_id = int(user_id)
    if class_id != 'None':
        class_id = int(class_id)
        if user_id in [u.user_id for u in TutorUser.query.all()]:
            user_db = db.session.query(TutorUser).filter(TutorUser.user_id == user_id).first()
            user_db.class_id = class_id
            db.session.commit()
        else:
            tut_user = TutorUser(user_id, class_id)
            db.session.add(tut_user)
            db.session.commit()
    else:
        if user_id in [u.user_id for u in TutorUser.query.all()]:
            user_sup = TutorUser.query.filter(TutorUser.user_id == user_id) \
                .filter(TutorUser.supervisor_id == user_id).first()
            db.session.delete(user_sup)
            db.session.commit()
    return redirect(url_for('.user_page', user=user_id))


@app.route('/organising_committee')
def organising_committee():
    membs = [get_org_info(u.user_id) for u
             in OrganisingCommittee.query.filter(OrganisingCommittee.year == curr_year).all()]
    m = sorted(membs, key=lambda u: u['first_name'])
    members = sorted(m, key=lambda u: u['last_name'])
    if 'access' in session.keys():
        access = session['access']
    else:
        access = 0
    return render_template('organising_committee/organising_committee.html', members=members,
                           access=access, curr_year=curr_year)


@app.route('/set_orgcom')
def set_orgcom():
    users = []
    user_ids = [u.user_id for u in Users.query.order_by(Users.last_name).order_by(Users.first_name).all()
                if u.user_type in [u for u in access_types.keys() if access_types[u] >= 8]]
    for u in user_ids:
        users.append(get_user_info(u))
    org = [o.user_id for o in OrganisingCommittee.query.filter(OrganisingCommittee.year == curr_year).all()]
    return render_template('organising_committee/set_orgcom.html', curr_year=curr_year, users=users, org=org)


@app.route('/save_orgcom', methods=['POST'])
def save_orgcom():
    org = request.form.getlist('orgcom')
    orgcom = [int(o) for o in org]
    for org in orgcom:
        if org not in [o.user_id for o in OrganisingCommittee.query.filter(OrganisingCommittee.year ==
                                                                           curr_year).all()]:
            member = OrganisingCommittee(user_id=org, year=curr_year)
            db.session.add(member)
            db.session.commit()
    check = [u.user_id for u in OrganisingCommittee.query.all() if u.user_id not in orgcom]
    if check:
        for user in check:
            OrganisingCommittee.query.filter(OrganisingCommittee.user_id == user).delete()
            db.session.commit()
    return redirect(url_for('.organising_committee'))


@app.route('/responsibilities')
def responsibilities():
    resps = [get_responsibility(r.responsibility_id) for r
             in Responsibilities.query.filter(Responsibilities.year == curr_year).all()]
    respons = sorted(resps, key=lambda u: u['name'])
    return render_template('organising_committee/responsibilities.html', responsibilities=respons, curr_year=curr_year)


@app.route('/add_responsibilities', defaults={'responsibility_id': ''})
@app.route('/add_responsibilities/<responsibility_id>')
def add_responsibilities(responsibility_id):
    orgcom = [get_org_info(u.user_id) for u
              in OrganisingCommittee.query.filter(OrganisingCommittee.year == curr_year).all()]
    m = sorted(orgcom, key=lambda u: u['first_name'])
    orgcom = sorted(m, key=lambda u: u['last_name'])
    if responsibility_id:
        responsibility = get_responsibility(responsibility_id)
    else:
        responsibility = None
    return render_template('organising_committee/add_responsibilities.html', curr_year=curr_year, orgcom=orgcom,
                           responsibility=responsibility)


@app.route('/save_responsibilities', methods=['POST'])
def save_responsibilities():
    year = curr_year
    name = request.form['name']
    if 'description' in request.form.keys():
        description = request.form['description']
    else:
        description = None
    if 'responsibility_id' in request.form.keys() and request.form['responsibility_id'] != '':
        responsibility_id = int(request.form['responsibility_id'])
        if responsibility_id in [r.responsibility_id for r in Responsibilities.query.all()]:
            db.session.query(Responsibilities).filter(Responsibilities.responsibility_id == responsibility_id) \
                .update({Responsibilities.name: name, Responsibilities.description: description})
        db.session.commit()
    else:
        resp = Responsibilities(name, description, year)
        db.session.add(resp)
        db.session.commit()
        responsibility_id = Responsibilities.query.filter(Responsibilities.year == curr_year) \
            .filter(Responsibilities.name == name).first().responsibility_id
    if 'assignees' in request.form.keys() and request.form['assignees']:
        ass = request.form.getlist('assignees')
        assignees = [int(assn) for assn in ass]
        for org in assignees:
            if org not in [r.user_id for r
                           in ResponsibilityAssignment.query.filter(ResponsibilityAssignment.responsibility_id
                                                                    == responsibility_id).all()]:
                ra = ResponsibilityAssignment(org, responsibility_id)
                db.session.add(ra)
                db.session.commit()
        org_users = [o.user_id for o
                     in ResponsibilityAssignment.query.filter(ResponsibilityAssignment.responsibility_id
                                                              == responsibility_id).all()]
        for user in org_users:
            if user not in assignees:
                to_del = db.session.query(ResponsibilityAssignment) \
                    .filter(ResponsibilityAssignment.responsibility_id == responsibility_id) \
                    .filter(ResponsibilityAssignment.user_id == user).first()
                db.session.delete(to_del)
            db.session.commit()
    return redirect(url_for('.responsibilities'))


@app.route('/delete_responsibility/<resp_id>')
def delete_responsibility(resp_id):
    resp_id = int(resp_id)
    if resp_id in [r.responsibility_id for r in Responsibilities.query.all()]:
        to_del = db.session.query(Responsibilities).filter(Responsibilities.responsibility_id == resp_id).first()
        db.session.delete(to_del)
        db.session.commit()
    if resp_id in [r.responsibility_id for r in ResponsibilityAssignment.query.all()]:
        for user in [u.user_id for u
                     in ResponsibilityAssignment.query
                             .filter(ResponsibilityAssignment.responsibility_id == resp_id).all()]:
            to_del = db.session.query(ResponsibilityAssignment) \
                .filter(ResponsibilityAssignment.responsibility_id == resp_id) \
                .filter(ResponsibilityAssignment.user_id == user).first()
            db.session.delete(to_del)
            db.session.commit()
    return redirect(url_for('.responsibilities'))


@app.route('/rev_analysis')
def rev_analysis_management():
    access = check_access(10)
    if access is not True:
        return access
    return render_template('rev_analysis/analysis_menu.html')


@app.route('/rev_analysis_results')
def rev_analysis_results():
    rev_criteria = get_criteria(curr_year)
    works = reg_works('all', 1)
    c, cats = categories_info()
    wks = sorted(works, key=lambda d: d['rk'], reverse=True)
    works = sorted(wks, key=lambda d: d['reg_tour'])
    cr_n = len(rev_criteria)
    return render_template('rev_analysis/rev_analysis_results.html', criteria=rev_criteria, works=works, cats=cats,
                           cr_n=cr_n)


@app.route('/analysis_state')
def analysis_state():
    renew_session()
    access = check_access(5)
    if access is not True:
        return access
    ana_nums, all_stats = analysis_nums()
    return render_template('rev_analysis/analysis_state.html', ana_nums=ana_nums, all_stats=all_stats)


@app.route('/analysis_criteria')
def analysis_criteria():
    renew_session()
    access = check_access(8)
    if access is not True:
        return access
    criteria = get_criteria(curr_year)
    return render_template('rev_analysis/analysis_criteria.html', criteria=criteria)


@app.route('/add_criteria')
def add_criteria():
    renew_session()
    access = check_access(10)
    if access is not True:
        return access
    return render_template('rev_analysis/add_criteria.html')


@app.route('/download_criteria')
def download_criteria():
    access = check_access(10)
    if access is not True:
        return access
    prev_year = curr_year - 1
    criteria = get_criteria(prev_year)
    crit = [v for v in criteria.values()]
    lines = []
    for c in crit:
        lines.append(c['name'] + '\t' + c['description'] + '\t' + str(c['year']) + '\t' + str(c['weight']))
    if not os.path.isdir('static/files/rev_crit'):
        os.mkdir('static/files/rev_crit')
    f_name = 'static/files/rev_crit/criteria_' + str(prev_year) + '.txt'
    with open(f_name, 'w', encoding='utf-8') as f:
        f.writelines([line + '\n' for line in lines])
    path = f_name
    return send_file(path, as_attachment=True)


@app.route('/adding_criteria', methods=['POST'])
def adding_criteria():
    renew_session()
    data = request.form['data']
    for line in data.split('\n'):
        if line != '':
            criteria = line.split('\t')
            name = criteria[0].strip()
            description = criteria[1].strip()
            year = datetime.datetime.strptime(criteria[2].strip(), '%Y')
            weight = criteria[3].strip()
            crit = RevCriteria(name, description, year, weight)
            db.session.add(crit)
            db.session.commit()
    return redirect(url_for('.analysis_criteria'))


@app.route('/edit_criterion/<crit_id>')
def edit_criterion(crit_id):
    access = check_access(10)
    if access is not True:
        return access
    criterion = get_criteria(curr_year)[int(crit_id)]
    return render_template('rev_analysis/edit_criterion.html', criterion=criterion)


@app.route('/write_criterion', methods=['POST'])
def write_criterion():
    renew_session()
    crit_id = int(request.form['id'])
    crit_name = request.form['name']
    if 'description' in request.form.keys():
        description = request.form['description']
    else:
        description = None
    crit_weight = int(request.form['weight'])
    if crit_id in [c.criterion_id for c in RevCriteria.query.all()]:
        db.session.query(RevCriteria).filter(RevCriteria.criterion_id == crit_id
                                             ).update({RevCriteria.criterion_name: crit_name,
                                                       RevCriteria.criterion_description: description,
                                                       RevCriteria.weight: crit_weight})
        db.session.commit()
    return redirect(url_for('.analysis_criteria'))


@app.route('/edit_value/<val_id>')
def edit_value(val_id):
    access = check_access(10)
    if access is not True:
        return access
    val = db.session.query(RevCritValues).filter(RevCritValues.value_id == int(val_id)).first()
    value = dict()
    value['id'] = val.value_id
    value['name'] = val.value_name
    value['comment'] = val.comment
    value['weight'] = val.weight
    return render_template('rev_analysis/edit_value.html', value=value)


@app.route('/write_value', methods=['POST'])
def write_value():
    renew_session()
    val_id = int(request.form['id'])
    val_name = request.form['name']
    if 'comment' in request.form.keys():
        comment = request.form['description']
    else:
        comment = None
    val_weight = int(request.form['weight'])
    if val_id in [v.value_id for v in RevCritValues.query.all()]:
        db.session.query(RevCritValues).filter(RevCritValues.value_id == val_id
                                               ).update({RevCritValues.value_name: val_name,
                                                         RevCritValues.comment: comment,
                                                         RevCritValues.weight: val_weight})
        db.session.commit()
    return redirect(url_for('.analysis_criteria'))


@app.route('/add_values')
def add_values():
    renew_session()
    access = check_access(10)
    if access is not True:
        return access
    return render_template('rev_analysis/add_values.html')


@app.route('/adding_values', methods=['POST'])
def adding_values():
    renew_session()
    data = request.form['data']
    for line in data.split('\n'):
        if line != '':
            values = line.split('\t')
            criterion = values[0].strip()
            value = values[1].strip()
            comment = values[2].strip()
            weight = values[3].strip()
            vals = RevCritValues(value, comment, weight)
            db.session.add(vals)
            db.session.commit()
            criterion_id = RevCriteria.query.filter(RevCriteria.year == datetime.datetime.strptime(str(curr_year), '%Y'
                                                                                                   ).date()
                                                    ).filter(RevCriteria.criterion_name == criterion
                                                             ).first().criterion_id
            vals = RevCritValues.query.order_by(RevCritValues.value_id.desc()
                                                ).filter(RevCritValues.value_name == value).all()
            val_ids = [v.value_id for v in vals if v.value_id not in [va.value_id for va in CriteriaValues.query.all()]]
            value_id = sorted(val_ids, reverse=True)[0]
            crit_val = CriteriaValues(criterion_id, value_id)
            db.session.add(crit_val)
            db.session.commit()
    return redirect(url_for('.analysis_criteria'))


@app.route('/download_values')
def download_values():
    access = check_access(10)
    if access is not True:
        return access
    prev_year = curr_year - 1
    criteria = get_criteria(prev_year)
    crit = [v for v in criteria.values()]
    lines = []
    for c in crit:
        crit_name = c['name']
        for v in c['values'].values():
            if v['comment'] is None:
                comment = ''
            else:
                comment = v['comment']
            lines.append(crit_name + '\t' + v['val_name'] + '\t' + comment + '\t' + str(v['val_weight']))
    if not os.path.isdir('static/files/rev_crit'):
        os.mkdir('static/files/rev_crit')
    f_name = 'static/files/rev_crit/values_' + str(prev_year) + '.txt'
    with open(f_name, 'w', encoding='utf-8') as f:
        f.writelines([line + '\n' for line in lines])
    path = f_name
    return send_file(path, as_attachment=True)


@app.route('/analysis_works/<cat_id>')
def analysis_works(cat_id):
    access = check_access(6)
    if access is not True:
        return access
    works = get_works(cat_id, 2, analysis_info=True)
    category = one_category(db.session.query(Categories).filter(Categories.cat_id == cat_id).first())
    need_analysis = check_analysis(cat_id)
    return render_template('rev_analysis/analysis_works.html', works=works, category=category,
                           need_analysis=need_analysis)


@app.route('/review_analysis/<work_id>')
def review_analysis(work_id):
    access = check_access(5)
    if access is not True:
        return access
    work = work_info(work_id, analysis_info=True)
    rk, analysis = get_analysis(work_id)
    criteria = get_criteria(curr_year)
    pre_ana = get_pre_analysis(work_id)
    if pre_ana is None:
        return redirect(url_for('.pre_analysis', work_id=work_id))
    work_id = int(work_id)
    if work_id in [w.work_id for w in RevComment.query.all()]:
        wo = db.session.query(RevComment).filter(RevComment.work_id == work_id).first()
        work_comment = wo.work_comment
        rev_comment = wo.rev_comment
    else:
        work_comment = None
        rev_comment = None
    return render_template('rev_analysis/review_analysis.html', work=work, analysis=analysis, criteria=criteria,
                           pre_ana=pre_ana, work_comment=work_comment, rev_comment=rev_comment)


@app.route('/pre_analysis/<work_id>')
def pre_analysis(work_id):
    access = check_access(6)
    if access is not True:
        return access
    work = work_info(work_id)
    work_id = int(work_id)
    pre = get_pre_analysis(work_id)
    if work_id in [w.work_id for w in RevComment.query.all()]:
        work_comment = RevComment.query.filter(RevComment.work_id == work_id).first().work_comment
        if work_comment is None:
            work_comment = ''
    else:
        work_comment = ''
    return render_template('/rev_analysis/pre_analysis.html', work=work, pre_ana=pre, work_comment=work_comment)


@app.route('/write_pre_analysis', methods=['POST'])
def write_pre_analysis():
    renew_session()
    work_id = request.form['work_id']
    work_id = int(work_id)
    if 'good_work' in request.form.keys():
        if request.form['good_work'] == 'True':
            good_work = True
        else:
            good_work = False
    else:
        good_work = None
    if 'research' in request.form.keys():
        research = request.form['research']
    else:
        research = None
    if 'has_review' in request.form.keys():
        if request.form['has_review'] == 'True':
            has_review = True
            rev_type = None
        elif request.form['has_review'] == 'points':
            has_review = False
            rev_type = 'points'
        else:
            has_review = False
            rev_type = None
    else:
        has_review = None
        rev_type = None
    if 'pushed' in request.form.keys():
        pushed = request.form['pushed']
        if pushed == 'True':
            pushed = True
        elif pushed == 'False':
            pushed = False
        else:
            pushed = None
    if 'work_comment' in request.form.keys() and request.form['work_comment'] != '':
        work_comment = request.form['work_comment']
    else:
        work_comment = None
    if work_id in [w.work_id for w in PreAnalysis.query.all()]:
        db.session.query(PreAnalysis).filter(PreAnalysis.work_id == int(work_id)).update(
            {PreAnalysis.good_work: good_work, PreAnalysis.research: research,
             PreAnalysis.has_review: has_review, PreAnalysis.rev_type: rev_type, PreAnalysis.pushed: pushed})
        db.session.commit()
    else:
        pre_ana = PreAnalysis(work_id, good_work, research, has_review, rev_type, pushed, None, None)
        db.session.add(pre_ana)
        db.session.commit()
    if work_id in [w.work_id for w in RevComment.query.all()]:
        db.session.query(RevComment).filter(RevComment.work_id == int(work_id)).update(
            {RevComment.work_comment: work_comment})
        db.session.commit()
    elif work_comment is not None:
        w_comm = RevComment(work_id, work_comment, None)
        db.session.add(w_comm)
        db.session.commit()
    if has_review is True:
        return redirect(url_for('.analysis_form', work_id=work_id))
    else:
        cat_id = WorkCategories.query.filter(WorkCategories.work_id == work_id).first().cat_id
        if work_id in [a.work_id for a in RevAnalysis.query.all()]:
            rev_ana = db.session.query(RevAnalysis).filter(RevAnalysis.work_id == work_id).all()
            for ana in rev_ana:
                db.session.delete(ana)
                db.session.commit()
        return redirect(url_for('.analysis_works', cat_id=cat_id))


@app.route('/analysis_form/<work_id>', defaults={'internal': None})
@app.route('/analysis_form/<work_id>/<internal>')
def analysis_form(work_id, internal):
    renew_session()
    access = check_access(6)
    if access is not True:
        return access
    criteria = get_criteria(curr_year)
    if not internal:
        work_id = int(work_id)
        work = work_info(work_id)
        rk, analysis = get_analysis(work_id)
        if work_id in [w.work_id for w in RevComment.query.all()]:
            rev_comment = RevComment.query.filter(RevComment.work_id == work_id).first().rev_comment
            if rev_comment is None:
                rev_comment = ''
        else:
            rev_comment = ''
        return render_template('/rev_analysis/analysis_form.html', criteria=criteria, work=work, analysis=analysis,
                               rev_comment=rev_comment, internal=None)
    else:
        work = {'work_id': int(work_id)}
        soup = BeautifulSoup(requests.post(url='https://vernadsky.info/review/' + str(work['work_id']),
                                           headers=mail_data.headers).text, 'html.parser')
        text = str(soup).split('<br/><br/>')
        rk, analysis = get_analysis(work_id, 'internal')
        if int(work_id) in [r.review_id for r in InternalReviewComments.query.all()]:
            rev_comment = InternalReviewComments.query.filter(InternalReviewComments.review_id == int(work_id)) \
                .first().comment
        else:
            rev_comment = ''
        return render_template('/rev_analysis/analysis_form.html', criteria=criteria, work=work, internal=internal,
                               analysis=analysis, rev_comment=rev_comment, text=text)


@app.route('/write_analysis/<internal>', methods=['POST'])
def write_analysis(internal):
    if internal == 'None':
        internal = None
    renew_session()
    work_id = int(request.form['work_id'])
    criteria_ids = [criterion.criterion_id for criterion in RevCriteria.query.all()]
    for criterion_id in criteria_ids:
        if str(criterion_id) in request.form.keys():
            value = int(request.form[str(criterion_id)])
            if not internal:
                if work_id in [w.work_id for w in RevAnalysis.query.all()]:
                    if criterion_id in [c.criterion_id for c in
                                        RevAnalysis.query.filter(RevAnalysis.work_id == work_id).all()]:
                        d = db.session.query(RevAnalysis).filter(RevAnalysis.work_id == work_id)
                        d.filter(RevAnalysis.criterion_id == criterion_id).update({RevAnalysis.value_id: value})
                        db.session.commit()
                    else:
                        crit_value = RevAnalysis(work_id, criterion_id, value)
                        db.session.add(crit_value)
                        db.session.commit()
                else:
                    crit_value = RevAnalysis(work_id, criterion_id, value)
                    db.session.add(crit_value)
                    db.session.commit()
            else:
                review_id = work_id
                if review_id in [r.review_id for r in InternalAnalysis.query.all()]:
                    if criterion_id in [c.criterion_id for c in
                                        InternalAnalysis.query.filter(InternalAnalysis.review_id == review_id).all()]:
                        d = db.session.query(InternalAnalysis).filter(InternalAnalysis.review_id == review_id)
                        d.filter(InternalAnalysis.criterion_id == criterion_id) \
                            .update({InternalAnalysis.value_id: value})
                        db.session.commit()
                    else:
                        crit_value = InternalAnalysis(review_id, criterion_id, value)
                        db.session.add(crit_value)
                        db.session.commit()
                else:
                    crit_value = InternalAnalysis(review_id, criterion_id, value)
                    db.session.add(crit_value)
                    db.session.commit()
    if not internal:
        if 'rev_comment' in request.form.keys() and request.form['rev_comment'] != '':
            rev_comment = request.form['rev_comment']
        else:
            rev_comment = None
        if work_id in [w.work_id for w in RevComment.query.all()]:
            db.session.query(RevComment).filter(RevComment.work_id == int(work_id)).update(
                {RevComment.rev_comment: rev_comment})
            db.session.commit()
        elif rev_comment is not None:
            rev_comm = RevComment(work_id, None, rev_comment)
            db.session.add(rev_comm)
            db.session.commit()
        cat_id = WorkCategories.query.filter(WorkCategories.work_id == work_id).first().cat_id
        return redirect(url_for('.analysis_works', cat_id=cat_id))
    else:
        review_id = int(work_id)
        if 'rev_comment' in request.form.keys() and request.form['rev_comment'] != '':
            rev_comment = request.form['rev_comment']
        else:
            rev_comment = None
        if review_id in [r.review_id for r in InternalReviewComments.query.all()]:
            db.session.query(InternalReviewComments).filter(InternalReviewComments.review_id == review_id).update(
                {InternalReviewComments.comment: rev_comment})
            db.session.commit()
        elif rev_comment:
            rev_comm = InternalReviewComments(review_id=review_id, comment=rev_comment)
            db.session.add(rev_comm)
            db.session.commit()
        reviewer_id = InternalReviews.query.filter(InternalReviews.review_id == review_id).first().reviewer_id
        return redirect(url_for('.see_reviews', reviewer_id=reviewer_id))


@app.route('/reviewer_comment/<reviewer_id>', methods=['POST'])
def reviewer_comment(reviewer_id):
    reviewer_id = int(reviewer_id)
    comment = request.form['text']
    if reviewer_id not in [r.reviewer_id for r in InternalReviewerComments.query.all()]:
        r = InternalReviewerComments(reviewer_id, comment)
        db.session.add(r)
        db.session.commit()
    else:
        db.session.query(InternalReviewerComments).filter(InternalReviewerComments.reviewer_id == reviewer_id) \
            .update({InternalReviewerComments.comment: comment})
        db.session.commit()
    return redirect(url_for('.see_reviews', reviewer_id=reviewer_id))


@app.route('/add_reviews', defaults={'done': None})
@app.route('/add_reviews/<done>')
def add_reviews(done):
    return render_template('internal_reviews/add_reviews.html', done=done)


@app.route('/save_reviews')
def save_reviews():
    response = json.loads(requests.post(url="https://vernadsky.info/all-works-json/" + str(curr_year) + "/",
                                        headers=mail_data.headers).text)
    for work in response:
        work_id = int(work['number'])
        if work['reviews']:
            for revs in work['reviews']:
                reviewer = revs['reviewer'].strip()
                review_id = int(revs['id'].strip())
                if reviewer not in [r.reviewer for r in InternalReviewers.query.all()]:
                    ir = InternalReviewers(reviewer)
                    db.session.add(ir)
                    db.session.commit()
                reviewer_id = InternalReviewers.query.filter(InternalReviewers.reviewer == reviewer).first().reviewer_id
                if review_id not in [r.review_id for r in InternalReviews.query.all()]:
                    rev = InternalReviews(review_id=review_id, reviewer_id=reviewer_id)
                    db.session.add(rev)
                    db.session.commit()
                else:
                    db.session.query(InternalReviews).filter(InternalReviews.review_id == review_id) \
                        .update({InternalReviews.reviewer_id: reviewer_id})
                    db.session.commit()
                if review_id not in [w.review_id for w in
                                     WorkReviews.query.filter(WorkReviews.work_id == work_id).all()]:
                    to_add = WorkReviews(work_id, review_id)
                    db.session.add(to_add)
                    db.session.commit()
                else:
                    db.session.query(WorkReviews).filter(WorkReviews.review_id == review_id) \
                        .update({WorkReviews.work_id: work_id})
                    db.session.commit()
    done = True
    return redirect(url_for('.add_reviews', done=done))


@app.route('/int_analysis')
def int_analysis():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('internal_reviews/int_analysis.html')


@app.route('/reviewers_to_review')
def reviewers_to_review():
    access = check_access(8)
    if access is not True:
        return access
    c, cats = categories_info()
    for cat in cats:
        cat_works = [w.work_id for w in WorkCategories.query.filter(WorkCategories.cat_id == cat['id'])]
        reviews = [{'work_id': w.work_id, 'reviewers': w.reviewers} for w in InternalReviews.query.all()
                   if w.work_id in cat_works]
        reviewers = []
        for rev in reviews:
            reviewers.extend([rev.strip() for rev in rev['reviewers'].split(',')])
        reviewers = set(reviewers)
        cat['reviewers'] = sorted(list(reviewers))
    return render_template('internal_reviews/reviewers_to_review.html', cats=cats)


@app.route('/internal_reviews')
def internal_reviews():
    access = check_access(8)
    if access is not True:
        return access
    reviewers = [{'id': r.reviewer_id, 'name': r.reviewer} for r in InternalReviewers.query.all()]
    for reviewer in reviewers:
        if reviewer['id'] in [r.reviewer_id for r in ReadingReviews.query.all()]:
            u = ReadingReviews.query.filter(ReadingReviews.reviewer_id == reviewer['id']).first()
            user = db.session.query(Users).filter(Users.user_id == u.reader_id).first()
            reader = {'id': user.user_id, 'name': user.first_name, 'l_name': user.last_name}
            reviewer['reader'] = reader
        revs = [r.review_id for r in
                InternalReviews.query.join(InternalAnalysis, InternalReviews.review_id == InternalAnalysis.review_id)
                .filter(InternalReviews.reviewer_id == reviewer['id']).all()]
        reviewer['analysed'] = len(revs)
    return render_template('internal_reviews/internal_reviews.html', reviewers=reviewers)


@app.route('/see_reviews/<reviewer_id>')
def see_reviews(reviewer_id):
    access = check_access(8)
    if access is not True:
        return access
    reviews = []
    for r in InternalReviews.query.filter(InternalReviews.reviewer_id == int(reviewer_id)).all():
        if r.review_id in [a.review_id for a in InternalAnalysis.query.all()]:
            read = True
        else:
            read = False

        soup = BeautifulSoup(requests.post(url='https://vernadsky.info/review/' + str(r.review_id),
                                           headers=mail_data.headers).text, 'html.parser')
        text = str(soup).split('<br/><br/>')
        reviews.append({'id': r.review_id, 'text': text, 'read': read})
    rev_no = len(reviews)
    read = len([r for r in reviews if r['read'] is True])
    if int(reviewer_id) in [r.reviewer_id for r in ReadingReviews.query.all()]:
        read_by = ReadingReviews.query.filter(ReadingReviews.reviewer_id == reviewer_id).first().reader_id
    else:
        read_by = None
    if int(reviewer_id) in [r.reviewer_id for r in InternalReviewerComments.query.all()]:
        comment = InternalReviewerComments.query.filter(InternalReviewerComments.reviewer_id == int(reviewer_id)) \
            .first().comment
    else:
        comment = None
    return render_template('internal_reviews/see_reviews.html', reviews=reviews, reviewer_id=reviewer_id,
                           rev_no=rev_no, read_by=read_by, comment=comment, read=read)


@app.route('/assign_reviewer/<do>/<reviewer_id>/<user_id>')
def assign_reviewer(do, reviewer_id, user_id):
    reviewer_id = int(reviewer_id)
    user_id = int(user_id)
    if do == 'do':
        if reviewer_id not in [r.reviewer_id for r in ReadingReviews.query.all()]:
            read = ReadingReviews(reviewer_id, user_id)
            db.session.add(read)
            db.session.commit()
        else:
            db.session.query(ReadingReviews).filter(ReadingReviews.reviewer_id == reviewer_id) \
                .update({ReadingReviews.reader_id: user_id})
            db.session.commit()
    elif do == 'undo':
        if reviewer_id in [r.reviewer_id for r in
                           ReadingReviews.query.filter(ReadingReviews.reader_id == user_id).all()]:
            to_del = db.session.query(ReadingReviews).filter(ReadingReviews.reviewer_id == reviewer_id) \
                .filter(ReadingReviews.reader_id == user_id).first()
            db.session.delete(to_del)
            db.session.commit()
    return redirect(url_for('.see_reviews', reviewer_id=reviewer_id))


@app.route('/internal_results')
def internal_results():
    c, cats = categories_info()
    for cat in cats:
        cat['works'] = [{'work_id': w.work_id} for w in
                        WorkCategories.query.filter(WorkCategories.cat_id == cat['id']).all()]
        for work in cat['works']:
            if WorkReviews.query.filter(WorkReviews.work_id == work['work_id']).first() is not None:
                work['review_id'] = WorkReviews.query.filter(WorkReviews.work_id == work['work_id']).first().review_id
                work['reviewer'] = {'reviewer_id': InternalReviews.query
                .filter(InternalReviews.review_id == work['review_id']).first().reviewer_id}
                if work['review_id'] in [r.review_id for r in InternalAnalysis.query.all()]:
                    work['analysed'] = True
                else:
                    work['analysed'] = False
            else:
                work['review_id'] = None
                work['reviewer'] = None
                work['analysed'] = False
            if work['analysed'] is True:
                cat['reviews'] = len(set([w['review_id'] for w in cat['works'] if w['review_id'] is not None]))
                cat['reviewers'] = len(
                    set([w['reviewer']['reviewer_id'] for w in cat['works'] if w['reviewer'] is not None]))
                cat['reviews_analysed'] = len(set([w['review_id'] for w in cat['works'] if w['analysed'] is True]))
                cat['reviewers_analysed'] = len(
                    set([w['reviewer']['reviewer_id'] for w in cat['works'] if w['analysed'] is True]))
    cats_all = {'works': sum([len(c['works']) for c in cats]),
                'reviews': sum([len([w['review_id'] for w in c['works'] if w['review_id'] is not None]) for c in cats]),
                'reviews_analysed': sum(
                    [len([w['review_id'] for w in c['works'] if w['analysed'] is True]) for c in cats]),
                'reviewers': sum(
                    [len(set(w['reviewer']['reviewer_id'] for w in c['works'] if w['reviewer'] is not None)) for c in
                     cats]),
                'reviewers_analysed': sum(
                    [len(set(w['reviewer']['reviewer_id'] for w in c['works'] if w['analysed'] is True)) for c in
                     cats])}
    return render_template('internal_reviews/internal_results.html', cats=cats, cats_all=cats_all)


@app.route('/add_works', defaults={'works_added': None, 'works_edited': None})
@app.route('/add_works/<works_added>/<works_edited>')
def add_works(works_added, works_edited):
    renew_session()
    access = check_access(8)
    if access is not True:
        return access
    return render_template('works/add_works.html', works_added=works_added, works_edited=works_edited)


@app.route('/applications_2_tour')
def applications_2_tour():
    renew_session()
    access = check_access(8)
    if access is not True:
        return access
    return render_template('works/applications_2_tour.html', year=curr_year)


@app.route('/many_works', methods=['POST'])
def many_works():
    renew_session()
    text = '{"works": ' + request.form['text'].strip('\n') + '}'
    works_added = 0
    works_edited = 0

    works = json.loads(text)
    w = works['works']

    for n in w:
        edited = False
        work_site_id = int(n['id'])
        work_id = int(n['number'])
        email = n['contacts']['email']
        tel = n['contacts']['phone']
        work_name = n['title']
        cat = n['section']['id']
        country = n['organization']['country']
        region = n['organization']['region']
        city = n['organization']['city']
        country_db = db.session.query(Cities)
        if country in [c.country for c in country_db.all()]:
            region_db = country_db.filter(Cities.country == country)
            if region in [r.region for r in region_db.all()]:
                city_db = region_db.filter(Cities.region == region)
                if city in [c.city for c in city_db.all()]:
                    timeshift = city_db.filter(Cities.city == city).first().msk_time_shift
                else:
                    timeshift = None
            else:
                timeshift = None
        else:
            timeshift = None
        if cat == 0:
            cat_id = None
        else:
            cat_id = Categories.query.filter(Categories.cat_site_id == cat).first().cat_id
        authors = n['authors']
        author_1_name = authors[0]['name']
        author_1_age = authors[0]['age']
        author_1_class = authors[0]['class']
        if len(authors) > 1:
            author_2_name = authors[1]['name']
            author_2_age = authors[1]['age']
            author_2_class = authors[1]['class']
        else:
            author_2_name = None
            author_2_age = None
            author_2_class = None
        if len(authors) > 2:
            author_3_name = authors[2]['name']
            author_3_age = authors[2]['age']
            author_3_class = authors[2]['class']
        else:
            author_3_name = None
            author_3_age = None
            author_3_class = None
        teacher_name = n['teacher']['name']
        if teacher_name == '':
            teacher_name = None
        status_id = int(n['status']['id'])
        status_name = n['status']['value']
        reg_tour = n['regional_tour']
        if work_id in [w.work_id for w in Works.query.all()]:
            db.session.query(Works).filter(Works.work_id == work_id).update({Works.work_name: work_name,
                                                                             Works.work_site_id: work_site_id,
                                                                             Works.email: email, Works.tel: tel,
                                                                             Works.author_1_name: author_1_name,
                                                                             Works.author_1_age: author_1_age,
                                                                             Works.author_1_class: author_1_class,
                                                                             Works.author_2_name: author_2_name,
                                                                             Works.author_2_age: author_2_age,
                                                                             Works.author_2_class: author_2_class,
                                                                             Works.author_3_name: author_3_name,
                                                                             Works.author_3_age: author_3_age,
                                                                             Works.author_3_class: author_3_class,
                                                                             Works.teacher_name: teacher_name,
                                                                             Works.reg_tour: reg_tour,
                                                                             Works.msk_time_shift: timeshift})
            edited = True
            db.session.commit()
        else:
            work_write = Works(work_id, work_name, work_site_id, email, tel, author_1_name, author_1_age,
                               author_1_class,
                               author_2_name, author_2_age, author_2_class, author_3_name, author_3_age, author_3_class,
                               teacher_name, reg_tour, timeshift, None)
            db.session.add(work_write)
            works_added += 1
            db.session.commit()
        if status_id not in [s.status_id for s in ParticipationStatuses.query.all()]:
            part_status = ParticipationStatuses(status_id, status_name)
            db.session.add(part_status)
            db.session.commit()
        if work_id in [s.work_id for s in WorkStatuses.query.all()]:
            db.session.query(WorkStatuses).filter(WorkStatuses.work_id == work_id
                                                  ).update({WorkStatuses.status_id: status_id})
            db.session.commit()
            edited = True
        else:
            work_status = WorkStatuses(work_id, status_id)
            db.session.add(work_status)
            db.session.commit()
        if work_id in [w.work_id for w in WorkCategories.query.all()]:
            if not cat_id:
                work_cat = db.session.query(WorkCategories).filter(WorkCategories.work_id == work_id).first()
                db.session.delete(work_cat)
                db.session.commit()
                edited = True
            else:
                db.session.query(WorkCategories).filter(WorkCategories.work_id == work_id
                                                        ).update({WorkCategories.cat_id: cat_id})
                db.session.commit()
                edited = True
        else:
            if cat_id:
                work_cat = WorkCategories(work_id, cat_id)
                db.session.add(work_cat)
                db.session.commit()
        if edited:
            works_edited += 1
    return redirect(url_for('.add_works', works_added=works_added, works_edited=works_edited))


@app.route('/button_works', defaults={'cat_id': 'all'})
@app.route('/button_works/<cat_id>')
def button_works(cat_id):
    url = request.referrer.replace(request.url_root, '').strip('/').split('/')

    response = json.loads(requests.post(url="https://vernadsky.info/all-works-json/" + str(curr_year) + "/",
                                        headers=mail_data.headers).text)
    works_added = 0
    works_edited = 0
    if cat_id == 'all':
        cats = [c.cat_site_id for c in Categories.query.filter(Categories.year == curr_year).all()]
        work_cats_list = []
        for cat in cats:
            work_cats_list.extend([w.work_id for w in WorkCategories.query.filter(WorkCategories.cat_id == cat).all()])
    elif cat_id == 'en':
        cats = [c.cat_site_id for c in Categories.query.join(CatDirs, Categories.cat_id == CatDirs.cat_id)
        .join(Directions, CatDirs.dir_id == Directions.direction_id).filter(Categories.year == curr_year)
        .filter(Directions.dir_name == 'Естественнонаучное').all()]
        work_cats_list = []
        for cat in cats:
            work_cats_list.extend([w.work_id for w in WorkCategories.query.filter(WorkCategories.cat_id == cat).all()])
    elif cat_id == 'gum':
        cats = [c.cat_site_id for c in Categories.query.join(CatDirs, Categories.cat_id == CatDirs.cat_id)
        .join(Directions, CatDirs.dir_id == Directions.direction_id).filter(Categories.year == curr_year)
        .filter(Directions.dir_name == 'Гуманитарное').all()]
        work_cats_list = []
        for cat in cats:
            work_cats_list.extend([w.work_id for w in WorkCategories.query.filter(WorkCategories.cat_id == cat).all()])
    else:
        cats = [Categories.query.filter(Categories.cat_id == int(cat_id)).first().cat_site_id]
        work_cats_list = []
        for cat in cats:
            work_cats_list.extend([w.work_id for w in WorkCategories.query.filter(WorkCategories.cat_id == cat).all()])

    work_id_list = [w.work_id for w in Works.query.all()]
    status_id_list = [s.status_id for s in ParticipationStatuses.query.all()]
    work_statuses_list = [s.work_id for s in WorkStatuses.query.all()]
    work_categories_list = [w.work_id for w in WorkCategories.query.all()]
    applications_2_tour_list = [w.work_id for w in Applications2Tour.query.all()]
    mails = {m.email: m.mail_id for m in Mails.query.all()}
    work_mails = [{'work_id': w_m.work_id, 'mail_id': w_m.mail_id} for w_m in WorkMail.query.all()]
    site_cats = {c.cat_site_id: c.cat_id for c in Categories.query.all()}
    works_reports = {w.work_id: w.reported for w in Works.query.all()}
    work_statuses = {w.work_id: w.status_id for w in WorkStatuses.query.all()}
    work_cats = {w.work_id: w.cat_id for w in WorkCategories.query.all()}
    # country_db = db.session.query(Cities)
    tz_regions = {t.region.lower(): t.tz for t in TimeZones.query.all()}
    tz_areas = {t.area.lower(): t.tz for t in TimeZones.query.filter(TimeZones.area != 0).all()}
    tz_countries = {t.country.lower(): t.tz for t in TimeZones.query.filter(TimeZones.region == 0).all()}
    works_pulled = []

    for n in response:
        if int(n['section']['id']) in cats:
            edited = False
            work_site_id = int(n['id'])
            work_id = int(n['number'])
            email = n['contacts']['email']
            tel = n['contacts']['phone']
            work_name = n['title']
            cat = n['section']['id']
            country = n['organization']['country']
            region = n['organization']['region']
            city = n['organization']['city']
            works_pulled.append(work_id)
            if city.lower() in tz_regions.keys() and city != '':
                timeshift = tz_regions[city.lower()]
            elif region.lower() in tz_regions.keys() and region != '':
                timeshift = tz_regions[region.lower()]
            elif country.lower() in tz_countries and country != '':
                timeshift = tz_countries[country.lower()]
            else:
                timeshift = None
                # if len(country) > 0:
                #     tz_country = country[0].upper() + country[1:].lower()
                #     tz_countries[tz_country.lower()] = timeshift
                # else:
                #     tz_country = ''
                # if len(region) > 0:
                #     tz_region = region[0].upper() + region[1:].lower()
                #     tz_regions[tz_region.lower()] = timeshift
                # else:
                #     tz_region = ''
                # if len(city) > 0:
                #     tz_area = city[0].upper() + city[1:].lower()
                #     tz_areas[tz_area.lower()] = timeshift
                # else:
                #     tz_area = ''
                # ta = TimeZones(country=tz_country, region=tz_region, area=tz_area, tz=timeshift)
                # db.session.add(ta)
                # db.session.commit()
            # if country in [c.country for c in country_db.all()]:
            #     region_db = country_db.filter(Cities.country == country)
            #     if region in [r.region for r in region_db.all()]:
            #         city_db = region_db.filter(Cities.region == region)
            #         if city in [c.city for c in city_db.all()]:
            #             timeshift = city_db.filter(Cities.city == city).first().msk_time_shift
            #         else:
            #             timeshift = None
            #     else:
            #         timeshift = None
            # else:
            #     timeshift = None

            if cat == 0:
                cat_id = None
            else:
                cat_id = site_cats[cat]
            authors = n['authors']
            author_1_name = authors[0]['name']
            author_1_age = authors[0]['age']
            author_1_class = authors[0]['class']
            if len(authors) > 1:
                author_2_name = authors[1]['name']
                author_2_age = authors[1]['age']
                author_2_class = authors[1]['class']
            else:
                author_2_name = None
                author_2_age = None
                author_2_class = None
            if len(authors) > 2:
                author_3_name = authors[2]['name']
                author_3_age = authors[2]['age']
                author_3_class = authors[2]['class']
            else:
                author_3_name = None
                author_3_age = None
                author_3_class = None
            teacher_name = n['teacher']['name']
            if teacher_name == '':
                teacher_name = None
            status_id = int(n['status']['id'])
            status_name = n['status']['value']
            reg_tour = n['regional_tour']
            if work_id in work_id_list:
                rep = works_reports[work_id]
            else:
                rep = False
            d = Works(work_id=work_id, work_name=work_name, work_site_id=work_site_id, email=email, tel=tel,
                      author_1_name=author_1_name, author_1_age=author_1_age, author_1_class=author_1_class,
                      author_2_name=author_2_name, author_2_age=author_2_age, author_2_class=author_2_class,
                      author_3_name=author_3_name, author_3_age=author_3_age, author_3_class=author_3_class,
                      teacher_name=teacher_name, reg_tour=reg_tour, msk_time_shift=timeshift, reported=rep)
            if work_id in work_id_list:
                db.session.query(Works).filter(Works.work_id == work_id).update({Works.work_name: work_name,
                                                                                 Works.work_site_id: work_site_id,
                                                                                 Works.email: email, Works.tel: tel,
                                                                                 Works.author_1_name: author_1_name,
                                                                                 Works.author_1_age: author_1_age,
                                                                                 Works.author_1_class: author_1_class,
                                                                                 Works.author_2_name: author_2_name,
                                                                                 Works.author_2_age: author_2_age,
                                                                                 Works.author_2_class: author_2_class,
                                                                                 Works.author_3_name: author_3_name,
                                                                                 Works.author_3_age: author_3_age,
                                                                                 Works.author_3_class: author_3_class,
                                                                                 Works.teacher_name: teacher_name,
                                                                                 Works.reg_tour: reg_tour,
                                                                                 Works.msk_time_shift: timeshift})
                edited = True
                db.session.commit()
            else:
                db.session.add(d)
                works_added += 1
                db.session.commit()
                work_id_list.append(work_id)
            if status_id not in status_id_list:
                part_status = ParticipationStatuses(status_id, status_name)
                db.session.add(part_status)
                db.session.commit()
                status_id_list.append(status_id)
            if work_id in work_statuses_list:
                if work_statuses[work_id]:
                    db.session.query(WorkStatuses).filter(WorkStatuses.work_id == work_id
                                                          ).update({WorkStatuses.status_id: status_id})
                    db.session.commit()
                    edited = True
            else:
                s = WorkStatuses(work_id, status_id)
                db.session.add(s)
                db.session.commit()
                work_statuses_list.append(work_id)
            w_cat = WorkCategories(work_id, cat_id)
            if w_cat not in work_cats.keys():
                if work_id in work_categories_list:
                    if not cat_id:
                        work_cat = work_cats[work_id]
                        db.session.delete(work_cat)
                        db.session.commit()
                        edited = True
                    else:
                        db.session.query(WorkCategories).filter(WorkCategories.work_id == work_id
                                                                ).update({WorkCategories.cat_id: cat_id})
                        db.session.commit()
                        edited = True
                else:
                    if cat_id:
                        db.session.add(w_cat)
                        db.session.commit()
            if email not in mails:
                m = Mails(email)
                db.session.add(m)
                db.session.commit()
                db.session.flush(m)
                mail_id = m.mail_id
            else:
                mail_id = mails[email]
            if {'work_id': work_id, 'mail_id': mail_id} not in work_mails:
                w_m = WorkMail(work_id, mail_id, False)
                db.session.add(w_m)
                db.session.commit()
            if edited:
                works_edited += 1

            if work_id not in applications_2_tour_list:
                w = Applications2Tour(work_id, None, False)
                db.session.add(w)
                db.session.commit()

    for w in work_cats_list:
        if w not in works_pulled:
            db.session.query(WorkCategories).filter(WorkCategories.work_id == w).update({WorkCategories.cat_id: 0})
            db.session.commit()

    if type(url) == list and 'category_page' in url:
        # errs = 'Обновлено успешно'
        if url[1] != '':
            return redirect(url_for('.category_page', cat_id=int(url[1])))
        else:
            return redirect(url_for('.categories_list'))
    else:
        return redirect(url_for('.add_works', works_added=works_added, works_edited=works_edited))


@app.route('/view_works', defaults={'cat': 'all'})
@app.route('/view_works/<cat>')
def view_works(cat):
    if cat == 'all':
        cats = []
        categories = db.session.query(Categories
                                      ).filter(Categories.year == curr_year
                                               ).join(CatDirs
                                                      ).join(Directions).join(Contests
                                                                              ).order_by(CatDirs.dir_id,
                                                                                         CatDirs.contest_id,
                                                                                         Categories.cat_name).all()
        for c in categories:
            curr_cat = {'cat_id': c.cat_id, 'cat_name': c.cat_name}
            works_db = Works.query.join(WorkCategories, WorkCategories.work_id == Works.work_id)\
            .join(WorkStatuses, WorkStatuses.work_id == Works.work_id).filter(WorkStatuses.status_id == 2)\
            .filter(WorkCategories.cat_id == c.cat_id).all()
            curr_cat['works'] = [{'work_id': w.work_id, 'name': w.work_name} for w in works_db]
            cats.append(curr_cat)
    else:
        cats = []
        c = Categories.query.filter(Categories.cat_id == int(cat)).first()
        curr_cat = {'cat_id': c.cat_id, 'cat_name': c.cat_name}
        works_db = Works.query.join(WorkCategories, WorkCategories.work_id == Works.work_id).filter(WorkCategories.cat_id == int(cat)).all()
        curr_cat['works'] = [{'work_id': w.work_id, 'name': w.work_name} for w in works_db]
        cats.append(curr_cat)
    return render_template('works/view_works.html', cats = cats)



@app.route('/timezones', defaults={'e': None})
@app.route('/timezones/<e>')
def timezones(e):
    all_t = [{'tz_id': t.tz_id, 'country': t.country, 'region': t.region, 'area': t.area, 'tz': t.tz}
             for t in TimeZones.query.all()]
    tz = sorted(all_t, key=lambda x: x['area'])
    all_t = sorted(tz, key=lambda x: x['region'])
    tz = sorted(all_t, key=lambda x: x['country'])
    if e is not None:
        t = TimeZones.query.filter(TimeZones.tz_id == int(e)).first()
        edit = {'tz_id': t.tz_id, 'country': t.country, 'region': t.region, 'area': t.area, 'tz': t.tz}
    else:
        edit = None
    return render_template('online_reports/timezones.html', tz=tz, edit=edit)


@app.route('/save_a_timezone', methods=['POST'])
def save_a_timezone():
    country = request.form['country']
    region = request.form['region']
    area = request.form['area']
    tz = int(request.form['tz'])
    if 'tz_id' in request.form.keys():
        tz_id = int(request.form['tz_id'])
        db.session.query(TimeZones).filter(TimeZones.tz_id == tz_id).update({TimeZones.country: country,
                                                                             TimeZones.region: region,
                                                                             TimeZones.area: area,
                                                                             TimeZones.tz: tz})
        db.session.commit()
    else:
        ta = TimeZones(country=country, region=region, area=area, tz=tz)
        db.session.add(ta)
        db.session.commit()
    return redirect(url_for('.timezones'))


@app.route('/save_timezones', methods=['POST'])
def save_timezones():
    data = request.files['file'].read().decode('mac_cyrillic').replace('\r', '')
    lines = data.split('\n')
    all_t = [{'country': t.country, 'region': t.region, 'area': t.area, 'tz': t.tz}
             for t in TimeZones.query.all()]
    for line in lines[2:]:
        if line != '':
            sta = {name: value for name, value in zip(lines[0].split('\t'), line.split('\t'))}
            if sta['tz'] == 'МСК':
                t_z = 0
            else:
                t_z = int(sta['tz'].strip('МСК+'))
            t = {'country': sta['country'], 'region': sta['region'], 'area': sta['area'], 'tz': t_z}
            if t not in all_t:
                ta = TimeZones(country=t['country'], region=t['region'], area=t['area'], tz=t['tz'])
                db.session.add(ta)
                db.session.commit()
    return redirect(url_for('.timezones'))


@app.route('/del_timezone/<tz_id>')
def del_timezone(tz_id):
    db.session.query(TimeZones).filter(TimeZones.tz_id == tz_id).delete()
    db.session.commit()
    return redirect(url_for('.timezones'))


@app.route('/many_applications', methods=['POST'])
def many_applications():
    renew_session()
    text = '{"works": ' + request.form['text'].strip('\n') + '}'
    works = json.loads(text)
    w = works['works']
    works_applied = []
    participants = []
    organisations = []
    for n in w:
        organisation = {'organisation_id': int(n['organization']['id']), 'name': n['organization']['name'],
                        'city': n['organization']['city'], 'country': n['organization']['country'],
                        'appl_no': int(n['id']), 'arrived': bool(n['arrival'])}
        organisations.append(organisation)
        works = [{'work': int(a['number']), 'appl': int(n['id']), 'arrived': bool(n['arrival'])} for a in n['works']]
        works_applied.extend(works)
        part_s = [{'id': int(p['id']), 'appl': int(n['id']), 'last_name': p['last_name'],
                   'first_name': p['first_name'], 'patronymic_name': p['patronymic_name'],
                   'participant_class': p['class'], 'role': p['role']} for p in n['delegation']['members']]
        participants.extend(part_s)
        for participant in ParticipantsApplied.query.filter(ParticipantsApplied.appl_id == int(n['id'])).all():
            if participant.participant_id not in [p['id'] for p in participants]:
                part_to_del = db.session.query(ParticipantsApplied).filter(ParticipantsApplied.participant_id ==
                                                                           participant.participant_id).first()
                db.session.delete(part_to_del)
                db.session.commit()
    for work in works_applied:
        if Applications2Tour(work['work'], work['appl'], False) not in Applications2Tour.query.all():
            if work['work'] in [wo.work_id for wo in Applications2Tour.query.all()]:
                db.session.query(Applications2Tour).filter(Applications2Tour.work_id == work['work']
                                                           ).update({Applications2Tour.appl_no: work['appl'],
                                                                     Applications2Tour.arrived: work['arrived']})
                db.session.commit()
            else:
                if work['work'] in [wo.work_id for wo in Works.query.all()]:
                    wo = Works.query.filter(Works.work_id == work['work']).first().work_id
                    appl = Applications2Tour(wo, work['appl'], False)
                    db.session.add(appl)
                    db.session.commit()
    for participant in participants:
        if ParticipantsApplied(participant['id'], participant['appl'], participant['last_name'],
                               participant['first_name'], participant['patronymic_name'],
                               participant['participant_class'], participant['role'], None) \
                not in ParticipantsApplied.query.all():
            if participant['id'] in [p.participant_id for p in ParticipantsApplied.query.all()]:
                db.session.query(ParticipantsApplied
                                 ).filter(ParticipantsApplied.participant_id == participant['id']
                                          ).update({ParticipantsApplied.appl_id: participant['appl'],
                                                    ParticipantsApplied.last_name: participant['last_name'],
                                                    ParticipantsApplied.first_name: participant['first_name'],
                                                    ParticipantsApplied.patronymic_name: participant['patronymic_name'],
                                                    ParticipantsApplied.participant_class: participant[
                                                        'participant_class'],
                                                    ParticipantsApplied.role: participant['role']})
                db.session.commit()
            else:
                part = ParticipantsApplied(participant['id'], participant['appl'], participant['last_name'],
                                           participant['first_name'], participant['patronymic_name'],
                                           participant['participant_class'], participant['role'], None)
                db.session.add(part)
                db.session.commit()

    for organisation in organisations:
        if Organisations(organisation['organisation_id'], organisation['name'], organisation['city'],
                         organisation['country']) not in Organisations.query.all():
            if organisation['organisation_id'] in [o.organisation_id for o in Organisations.query.all()]:
                db.session.query(Organisations).filter(Organisations.organisation_id == organisation['organisation_id']) \
                    .update({Organisations.name: organisation['name'],
                             Organisations.city: organisation['city'],
                             Organisations.country: organisation['country']})
                db.session.commit()
            else:
                o = Organisations(organisation['organisation_id'], organisation['name'], organisation['city'],
                                  organisation['country'])
                db.session.add(o)
                db.session.commit()
        if OrganisationApplication(organisation['organisation_id'], organisation['appl_no'], organisation['arrived']) \
                not in OrganisationApplication.query.all():
            if organisation['organisation_id'] in [o.organisation_id for o in OrganisationApplication.query.all()]:
                db.session.query(OrganisationApplication) \
                    .filter(OrganisationApplication.organisation_id == organisation['organisation_id']) \
                    .update({OrganisationApplication.appl_no: organisation['appl_no'],
                             OrganisationApplication.arrived: organisation['arrived']})
                db.session.commit()
            else:
                o = OrganisationApplication(organisation['organisation_id'], organisation['appl_no'],
                                            organisation['arrived'])
                db.session.add(o)
                db.session.commit()

    applications_to_del = [a.appl_no for a in Applications2Tour.query.all()
                           if a.appl_no not in [o['appl_no'] for o in organisations]]
    for a in applications_to_del:
        if a in [wo.appl_no for wo in Applications2Tour.query.all()]:
            for to_del in db.session.query(Applications2Tour).filter(Applications2Tour.appl_no == a).all():
                db.session.delete(to_del)
                db.session.commit()
        if a in [wo.appl_id for wo in ParticipantsApplied.query.all()]:
            for to_del in db.session.query(ParticipantsApplied).filter(ParticipantsApplied.appl_id == a).all():
                db.session.delete(to_del)
                db.session.commit()
        if a in [wo.appl_no for wo in OrganisationApplication.query.all()]:
            for to_del in db.session.query(OrganisationApplication).filter(OrganisationApplication.appl_no == a).all():
                db.session.delete(to_del)
                db.session.commit()
    return redirect(url_for('.applications_2_tour'))


@app.route('/button_applications')
def button_applications():
    response = json.loads(requests.post(url="https://vernadsky.info/second-tour-requests-json/" + str(curr_year) + "/",
                                        headers=mail_data.headers).text)
    works_applied = [w.work_id for w in AppliedForOnline.query.filter(str(AppliedForOnline.work_id)[:2]
                                                                      == str(curr_year)[2:]).all()]
    participants = []
    organisations = []

    appl_2_tour = [(w.work_id, w.appl_no, w.arrived) for w in Applications2Tour.query.all()]
    part_ids = [p.participant_id for p in ParticipantsApplied.query.all()]
    work_applied_ids = [wo.work_id for wo in Applications2Tour.query.all()]
    work_ids = [wo.work_id for wo in Works.query.all()]
    part_appl = [(p.participant_id, p.appl_id, p.last_name, p.first_name, p.patronymic_name,
                  p.participant_class, p.role, None) for p in ParticipantsApplied.query.all()]
    org_appl = [(o.organisation_id, o.name, o.city, o.country) for o in Organisations.query.all()]
    org_ids = [o.organisation_id for o in Organisations.query.all()]
    org_appl_db = [(o.organisation_id, o.appl_no, o.arrived) for o in OrganisationApplication.query.all()]
    org_appl_ids = [o.organisation_id for o in OrganisationApplication.query.all()]

    for n in response:
        organisation = {'organisation_id': int(n['organization']['id']), 'name': n['organization']['name'],
                        'city': n['organization']['city'], 'country': n['organization']['country'],
                        'appl_no': int(n['id']), 'arrived': bool(n['arrival'])}
        organisations.append(organisation)
        works = [{'work': int(a['number']), 'appl': int(n['id']), 'arrived': bool(n['arrival'])} for a in n['works']]
        works_applied.extend(works)
        part_s = [{'id': int(p['id']), 'appl': int(n['id']), 'last_name': p['last_name'],
                   'first_name': p['first_name'], 'patronymic_name': p['patronymic_name'],
                   'participant_class': p['class'], 'role': p['role']} for p in n['delegation']['members']]
        participants.extend(part_s)
    for participant in part_ids:
        if participant not in [p['id'] for p in participants]:
            part_to_del = db.session.query(ParticipantsApplied).filter(ParticipantsApplied.participant_id ==
                                                                       participant).first()
            db.session.delete(part_to_del)
            db.session.commit()
    for work in works_applied:
        if (work['work'], work['appl'], work['arrived']) not in appl_2_tour:
            if work['work'] in work_applied_ids:
                db.session.query(Applications2Tour).filter(Applications2Tour.work_id == work['work']
                                                           ).update({Applications2Tour.appl_no: work['appl'],
                                                                     Applications2Tour.arrived: work['arrived']})
                db.session.commit()
            else:
                if work['work'] in work_ids:
                    appl = Applications2Tour(work['work'], work['appl'], False)
                    db.session.add(appl)
                    db.session.commit()
    for participant in participants:
        if (participant['id'], participant['appl'], participant['last_name'], participant['first_name'],
            participant['participant_class'], participant['role'], None) not in part_appl:
            if participant['id'] in part_ids:
                db.session.query(ParticipantsApplied
                                 ).filter(ParticipantsApplied.participant_id == participant['id']
                                          ).update({ParticipantsApplied.appl_id: participant['appl'],
                                                    ParticipantsApplied.last_name: participant['last_name'],
                                                    ParticipantsApplied.first_name: participant['first_name'],
                                                    ParticipantsApplied.patronymic_name: participant['patronymic_name'],
                                                    ParticipantsApplied.participant_class: participant[
                                                        'participant_class'],
                                                    ParticipantsApplied.role: participant['role']})
                db.session.commit()
            else:
                part = ParticipantsApplied(participant['id'], participant['appl'], participant['last_name'],
                                           participant['first_name'], participant['patronymic_name'],
                                           participant['participant_class'], participant['role'], None)
                db.session.add(part)
                db.session.commit()

    for organisation in organisations:
        if (organisation['organisation_id'], organisation['name'], organisation['city'], organisation['country']) \
                not in org_appl:
            if organisation['organisation_id'] in org_ids:
                db.session.query(Organisations).filter(Organisations.organisation_id == organisation['organisation_id']) \
                    .update({Organisations.name: organisation['name'],
                             Organisations.city: organisation['city'],
                             Organisations.country: organisation['country']})
                db.session.commit()
            else:
                o = Organisations(organisation['organisation_id'], organisation['name'], organisation['city'],
                                  organisation['country'])
                db.session.add(o)
                db.session.commit()
        if (organisation['organisation_id'], organisation['appl_no'], organisation['arrived']) not in org_appl_db:
            if organisation['organisation_id'] in org_appl_ids:
                db.session.query(OrganisationApplication) \
                    .filter(OrganisationApplication.organisation_id == organisation['organisation_id']) \
                    .update({OrganisationApplication.appl_no: organisation['appl_no'],
                             OrganisationApplication.arrived: organisation['arrived']})
                db.session.commit()
            else:
                o = OrganisationApplication(organisation['organisation_id'], organisation['appl_no'],
                                            organisation['arrived'])
                db.session.add(o)
                db.session.commit()

    for org in org_ids:
        if org not in org_appl_ids:
            o = OrganisationApplication(org, None, None)
            db.session.add(o)
            db.session.commit()

    applications_to_del = [a.appl_no for a in Applications2Tour.query.all()
                           if a.appl_no not in [o['appl_no'] for o in organisations] and a.appl_no is not None]
    for a in applications_to_del:
        if a in [wo.appl_no for wo in Applications2Tour.query.all()]:
            a_w = [w.work_id for w in Applications2Tour.query.filter(Applications2Tour.appl_no == a).all()]
            for k in a_w:
                Applications2Tour.query.filter(Applications2Tour.appl_no == a).filter(Applications2Tour.work_id == k) \
                    .update({Applications2Tour.appl_no: None,
                             Applications2Tour.arrived: False})
                db.session.commit()
        if a in [wo.appl_id for wo in ParticipantsApplied.query.all()]:
            for to_del in db.session.query(ParticipantsApplied).filter(ParticipantsApplied.appl_id == a).all():
                db.session.delete(to_del)
                db.session.commit()
        if a in [wo.appl_no for wo in OrganisationApplication.query.all()]:
            for to_del in db.session.query(OrganisationApplication).filter(OrganisationApplication.appl_no == a).all():
                db.session.delete(to_del)
                db.session.commit()
    return redirect(url_for('.applications_2_tour'))


@app.route('/top_100')
def top_100():
    access = check_access(5)
    if access is not True:
        return access
    total, no_fee = no_fee_nums()
    return render_template('works/top_100.html', no_fee=no_fee, total=total)


@app.route('/top_100_excel')
def top_100_excel():
    return send_file('static/files/generated_files/no_fee_works' + str(curr_year) + '.xlsx', as_attachment=True)


@app.route('/top_100_for_site')
def top_100_for_site():
    total, no_fee = no_fee_nums()
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with open('static/files/generated_files/top_100_for_site.html', 'w'
              ) as writer:
        writer.write(render_template('works/top_100_for_site.html', no_fee=no_fee, total=total))
    return send_file('static/files/generated_files/top_100_for_site.html', as_attachment=True)


@app.route('/apply_for_online', defaults={'errs_a': None, 'errs_b': None})
@app.route('/apply_for_online/<errs_a>/<errs_b>')
def apply_for_online(errs_a, errs_b):
    access = check_access(8)
    if access is not True:
        return access
    if errs_b == 'a':
        errs_b = None
    elif errs_b is not None:
        errs_b = errs_b.split('\n')
    if errs_a == 'a':
        errs_a = None
    elif errs_a is not None:
        errs_a = errs_a.split('\n')
    return render_template('online_reports/apply_for_online.html', errors_a=errs_a, errors_b=errs_b)


@app.route('/applied_for_online', methods=['POST'])
def applied_for_online():
    works = request.form['works']
    works_list = []
    success = False
    all_works = [w.work_id for w in Works.query.all()]
    participated = [w.work_id for w in ParticipatedWorks.query.all()]
    applied = [w.work_id for w in AppliedForOnline.query.all()]
    w_stat = {w.work_id: w.status_id for w in WorkStatuses.query.all()}
    if not participated:
        return redirect(url_for('.apply_for_online', errs_a='Не загрузились участвовавшие работы'))
    if ',' in works:
        works_list.extend(works.split(','))
    else:
        works_list.append(works)
    if '' in works_list:
        works_list.remove('')
    errors = {}
    for work in set(works_list):
        try:
            work = int(work.strip())
            if work in all_works:
                if w_stat[work] < 6:
                    errors[work] = 'работа не прошла во Второй тур'
                else:
                    if work in participated:
                        errors[work] = 'работа уже участвовала во 2 туре'
                    else:
                        if work not in applied:
                            w = AppliedForOnline(work)
                            db.session.add(w)
                            db.session.commit()
                            success = True
            else:
                errors[work] = 'работа не найдена'
        except ValueError:
            if success is False and work not in errors.keys():
                errors[work] = 'некорректный номер работы'
            pass
    errs = ''
    if errors != {}:
        for work, error in errors.items():
            errs += str(work) + ' - ' + error + '\n'
    else:
        errs = 'a'
    return redirect(url_for('.apply_for_online', errs_a=errs, errs_b='a'))


@app.route('/participated', methods=['POST'])
def participated():
    works = request.form['works']
    works_list = []
    success = False
    if ',' in works:
        works_list.extend(works.split(','))
    else:
        works_list.append(works)
    errors = {}
    for work in set(works_list):
        try:
            work = int(work.strip())
            if work in [w.work_id for w in Works.query.all()]:
                if work not in [w.work_id for w in ParticipatedWorks.query.all()]:
                    work_db = db.session.query(Works).filter(Works.work_id == work).first()
                    w = ParticipatedWorks(work_db.work_id)
                    db.session.add(w)
                    db.session.commit()
                    success = True
            else:
                errors[work] = 'работа не найдена'
        except ValueError:
            if success is False and work not in errors.keys():
                errors[work] = 'некорректный номер работы'
            pass
    errs = ''
    if errors != {}:
        for work, error in errors.items():
            errs += str(work) + ' - ' + error + '\n'
    else:
        errs = 'a'
    return redirect(url_for('.apply_for_online', errs_b=errs, errs_a='a'))


@app.route('/online_applicants')
def online_applicants():
    cats = []

    text = '<h2>Работы, заявленные для участия в Дополнительном онлайн-конкурсе</h2>\n'
    text += '''<p>В список выступающих работа будет включена только после оплаты оргвзноса.
        Если вы оплатили оргвзнос 3 или больше рабочих дня назад, и это не отражено в таблице,
        пришлите чек оплаты оргвзноса на <a href="info@vernadsky.info" target="_blank">info@vernadsky.info</a>.
        <br>Если вы подали заявку на участие, но не были включены в список ниже,
        напишите об этом на <a href="info@vernadsky.info" target="_blank">info@vernadsky.info</a>.</p>\n'''
    table = '''<table frame="void" border="1px" bordercolor="#4962A4" cellpadding="3px" cellspacing="0px">
        <tr>
            <td width="6%" align="сenter"><b>
                Номер работы
            </b></td>
            <td width="59%" align="сenter"><b>
                Название
            </b></td>
            <td width="25%" align="сenter"><b>
                Авторы
            </b></td>
            <td width="10%" align="сenter"><b>
                Оргвзнос
            </b></td>
        </tr>'''

    for cat in [c.cat_id for c in Categories.query.filter(Categories.year == curr_year).all()]:
        c, cat_info = categories_info(cat)
        cat_works = [work_info(w.work_id, w_payment_info=True) for w in
                     Works.query.join(WorkCategories, Works.work_id == WorkCategories.work_id)
                     .filter(WorkCategories.cat_id == cat).all() if w.work_id in [wo.work_id for wo
                                                                                  in AppliedForOnline.query.all()]]
        cat_info['works'] = cat_works
        cats.append(cat_info)
        if cat_info['works']:
            table += '''\n<tr><td align="сenter" colspan="4"><b>'''
            cat_row = cat_info['name'] + '''</b></td>\n</tr>'''
            table += cat_row
            for work in cat_info['works']:
                table += '''\n<tr><td align="сenter">'''
                to_add = str(work['work_id']) + '''</td><td>'''
                table += to_add
                to_add = work['work_name'] + '''</td><td>'''
                table += to_add
                to_add = work['authors'] + '''</td><td align="сenter">'''
                table += to_add
                if work['payed'] is True:
                    to_add = '''Оплачен</td></tr>'''
                else:
                    to_add = '''Не оплачен</td></tr>'''
                table += to_add

    table += '''\n</table>'''
    text += table
    with open('static/files/generated_files/online_applicants.html', 'w', encoding='utf-8') as f:
        f.write(text)
    return render_template('online_reports/online_applicants.html', cats=cats)


def create_report_dates_html(cat_dates):
    for cat in cat_dates:
        da = []
        if 'd_1' in cat.keys() and cat['d_1'] is not None:
            da.append(cat['d_1'])
            if cat['d_2'] is not None:
                da.append(cat['d_2'])
                if cat['d_3'] is not None:
                    da.append(cat['d_3'])
                else:
                    cat['d_3'] = ''
            else:
                cat['d_2'] = ''
        else:
            cat['d_1'] = ''
        if da and type(da) == list:
            d = '; '.join(da)
        else:
            d = 'Не назначены'
        cat['all_dates'] = d

    all_dates = set()
    all_dates.update([c['day_1_date'] for c in cat_dates])
    all_dates.update([c['day_2_date'] for c in cat_dates])
    all_dates.update([c['day_3_date'] for c in cat_dates])
    all_dates.remove(None)
    dates = sorted(list(all_dates))

    with open('static/files/generated_files/report_dates_' + str(curr_year) + '.html', 'w', encoding='utf-8') as f:
        f.write(render_template('online_reports/report_dates_html.html', cat_dates=cat_dates))

    table_1 = [{'Название секции': cat['cat_name'], 'Telegram-канал': cat['tg_channel'], 'День 1': cat['d_1'],
                'День 2': cat['d_2'], 'День 3': cat['d_3']} for cat in cat_dates]
    table_2 = []
    for cat in cat_dates:
        c_d = [cat['day_1_date'], cat['day_2_date'], cat['day_3_date']]
        c = {'Название секции': cat['cat_name'], 'Telegram-канал': cat['tg_channel']}
        for d in dates:
            day_name = days[d.strftime('%w')] + ', ' + d.strftime('%d.%m')
            if d in c_d:
                c[day_name] = True
            else:
                c[day_name] = ''
        table_2.append(c)

    t_1 = pd.DataFrame(data=table_1)
    t_2 = pd.DataFrame(data=table_2)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/report_dates_' + str(curr_year) + '.xlsx') as writer:
        t_1.to_excel(writer, sheet_name='Даты')
        t_2.to_excel(writer, sheet_name='Сетка')
    return 'ok'


@app.route('/download_applicants')
def download_applicants():
    return send_file('static/files/generated_files/online_applicants.html', as_attachment=True)


@app.route('/download_report_dates_html')
def download_report_dates_html():
    return send_file('static/files/generated_files/report_dates_' + str(curr_year) + '.html', as_attachment=True)


@app.route('/download_report_dates_excel')
def download_report_dates_excel():
    return send_file('static/files/generated_files/report_dates_' + str(curr_year) + '.xlsx', as_attachment=True)


@app.route('/works_for_free/<cat_id>', methods=['POST'])
def works_for_free(cat_id):
    works = request.form['works']
    works_list = []
    success = False
    if ',' in works:
        works_list.extend(works.split(','))
    else:
        works_list.append(works)
    errors = {}
    for work in works_list:
        try:
            work = int(work.strip())
            if work in [w.work_id for w in Works.query.all()]:
                work_db = db.session.query(Works).filter(Works.work_id == work).first()
                if work_db.work_id in [w.work_id for w
                                       in WorkCategories.query.filter(WorkCategories.cat_id == cat_id).all()]:
                    if WorkStatuses.query.filter(WorkStatuses.work_id == work).first().status_id < 2:
                        errors[work] = 'работа не прошла на Конкурс'
                    else:
                        if work_db.reg_tour:
                            errors[work] = 'работа регионального тура, нельзя отменить оргвзнос'
                        else:
                            if work not in [w.work_id for w in WorksNoFee.query.all()]:
                                wnf = WorksNoFee(work_db.work_id)
                                db.session.add(wnf)
                                db.session.commit()
                                success = True
                else:
                    errors[work] = 'работа не из вашей секции'
            else:
                errors[work] = 'работа не найдена'
        except BaseException:
            if success is False and work not in errors.keys():
                errors[work] = 'некорректный номер работы'
            pass
    errs = ''
    if errors != {}:
        for work, error in errors.items():
            if isinstance(error, str) and '%' in error:
                error = unquote(error)
            errs += str(work) + ' - ' + error + '\n'
    else:
        errs = None
    return redirect(url_for('.category_page', cat_id=cat_id, errors=errs))


@app.route('/remove_no_fee/<cat_id>/<work_id>')
def remove_no_fee(cat_id, work_id):
    work = db.session.query(WorksNoFee).filter(WorksNoFee.work_id == work_id).first()
    db.session.delete(work)
    db.session.commit()
    return redirect(url_for('.category_page', cat_id=cat_id))


@app.route('/no_fee_list')
def no_fee_list():
    cats = []
    for cat in Categories.query.join(CatDirs).order_by(CatDirs.dir_id, CatDirs.contest_id, Categories.cat_name).all():
        cat_db = db.session.query(Categories).filter(Categories.cat_id == cat.cat_id).first()
        works = []
        for work_id in [w.work_id for w in WorkCategories.query.filter(WorkCategories.cat_id == cat_db.cat_id).all()]:
            if work_id in [w.work_id for w in WorksNoFee.query.all()]:
                work_db = db.session.query(Works).filter(Works.work_id == work_id).first()
                works.append({'work_id': work_id, 'work_name': work_db.work_name,
                              'authors': ', '.join([str(work_db.author_1_name), str(work_db.author_2_name),
                                                    str(work_db.author_3_name)]).replace(', None', '')})
        cats.append({'cat_id': cat_db.cat_id, 'cat_name': cat_db.cat_name, 'works': works})
    return render_template('works/no_fee_list.html', cats=cats)


@app.route('/drive_links')
def drive_links():
    count, categories = categories_info()
    return render_template('categories/drive_links.html', categories=categories)


@app.route('/set_report_dates', defaults={'message': None})
@app.route('/set_report_dates/<message>')
def set_report_dates(message):
    c, cats = categories_info()
    cat_dates = []
    for cat in cats:
        c_dates = {'cat_id': cat['id'], 'cat_name': cat['name'], 'tg_channel': cat['tg_channel']}
        if cat['id'] in [c.cat_id for c in ReportDates.query.all()]:
            dates_db = db.session.query(ReportDates).filter(ReportDates.cat_id == cat['id']).first()
            if dates_db.day_1:
                c_dates['day_1'] = dates_db.day_1.strftime('%Y-%m-%d')
                c_dates['d_1'] = days[dates_db.day_1.strftime('%w')] + ', ' + dates_db.day_1.strftime('%d.%m')
                c_dates['day_1_date'] = dates_db.day_1
            else:
                c_dates['day_1'] = None
                c_dates['d_1'] = None
                c_dates['day_1_date'] = None
            if dates_db.day_2:
                c_dates['day_2'] = dates_db.day_2.strftime('%Y-%m-%d')
                c_dates['d_2'] = days[dates_db.day_2.strftime('%w')] + ', ' + dates_db.day_2.strftime('%d.%m')
                c_dates['day_2_date'] = dates_db.day_2
            else:
                c_dates['day_2'] = None
                c_dates['d_2'] = None
                c_dates['day_2_date'] = None
            if dates_db.day_3:
                c_dates['day_3'] = dates_db.day_3.strftime('%Y-%m-%d')
                c_dates['d_3'] = days[dates_db.day_3.strftime('%w')] + ', ' + dates_db.day_3.strftime('%d.%m')
                c_dates['day_3_date'] = dates_db.day_3
            else:
                c_dates['day_3'] = None
                c_dates['d_3'] = None
                c_dates['day_3_date'] = None
        else:
            c_dates['day_1'] = None
            c_dates['day_2'] = None
            c_dates['day_3'] = None
            c_dates['d_1'] = None
            c_dates['d_2'] = None
            c_dates['d_3'] = None
            c_dates['day_1_date'] = None
            c_dates['day_2_date'] = None
            c_dates['day_3_date'] = None
        cat_dates.append(c_dates)
    create_report_dates_html(cat_dates)
    return render_template('online_reports/set_report_dates.html', cat_dates=cat_dates, message=unquote(message))


@app.route('/save_report_dates', methods=['POST'])
def save_report_dates():
    dates = []
    day_1 = None
    day_2 = None
    day_3 = None
    for cat_id in [c.cat_id for c in Categories.query.filter(Categories.year == curr_year).all()]:
        if str(cat_id) + '_day_1' in request.form.keys():
            if request.form[str(cat_id) + '_day_1'] != '':
                day_1 = datetime.datetime.strptime(request.form[str(cat_id) + '_day_1'], '%Y-%m-%d').date()
            else:
                day_1 = None
        elif cat_id in [c.cat_id for c in ReportDates.query.all()] and not day_1:
            day_1 = ReportDates.query.filter(ReportDates.cat_id == cat_id).first().day_1
        else:
            day_1 = None
        if str(cat_id) + '_day_2' in request.form.keys():
            if request.form[str(cat_id) + '_day_2'] != '':
                day_2 = datetime.datetime.strptime(request.form[str(cat_id) + '_day_2'], '%Y-%m-%d').date()
            else:
                day_2 = None
        elif cat_id in [c.cat_id for c in ReportDates.query.all()] and not day_2:
            day_2 = ReportDates.query.filter(ReportDates.cat_id == cat_id).first().day_2
        else:
            day_2 = None
        if str(cat_id) + '_day_3' in request.form.keys():
            if request.form[str(cat_id) + '_day_3'] != '':
                day_3 = datetime.datetime.strptime(request.form[str(cat_id) + '_day_3'], '%Y-%m-%d').date()
            else:
                day_3 = None
        elif cat_id in [c.cat_id for c in ReportDates.query.all()] and not day_3:
            day_3 = ReportDates.query.filter(ReportDates.cat_id == cat_id).first().day_3
        else:
            day_3 = None
        dates.append({'cat_id': cat_id, 'day_1': day_1, 'day_2': day_2, 'day_3': day_3})
    for date in dates:
        if date['cat_id'] in [c.cat_id for c in ReportDates.query.all()]:
            db.session.query(ReportDates).filter(ReportDates.cat_id == date['cat_id']
                                                 ).update({ReportDates.day_1: date['day_1'],
                                                           ReportDates.day_2: date['day_2'],
                                                           ReportDates.day_3: date['day_3']})
            db.session.commit()
        else:
            rep_d = ReportDates(date['cat_id'], date['day_1'], date['day_2'], date['day_3'])
            db.session.add(rep_d)
            db.session.commit()
            db.session.commit()
        success = 'Даты добавлены'
    return redirect(url_for('.set_report_dates', message=success))


@app.route('/reports_order/<cat_id>')
def reports_order(cat_id):
    cat_id = int(cat_id)
    if cat_id in [c.cat_id for c in CategoryUnions.query.all()]:
        union = CategoryUnions.query.filter(CategoryUnions.cat_id == cat_id).first().union_id
        cats = [c.cat_id for c in CategoryUnions.query.filter(CategoryUnions.union_id == union).all()]
        union = True
    else:
        cats = [cat_id]
        union = False
    categories = []
    works_unordered = []
    approved_for_2 = 0
    participating = 0
    c_dates = []
    works = {}
    appl_for_online = [w.work_id for w in AppliedForOnline.query.all()]
    reported = [w.work_id for w in ReportOrder.query.all()]

    for cat_id in cats:
        dates_db = db.session.query(ReportDates).filter(ReportDates.cat_id == cat_id).first()
        cat_name = Categories.query.filter(Categories.cat_id == cat_id).first().cat_name
        categories.append({'cat_id': cat_id, 'cat_name': cat_name})
        works.update(
            get_works(cat_id, 2, 'online', appl_info=True, w_payment_info=True, reports_info=True, site_id=True))
        approved_for_2 += len(works)
        if not dates_db:
            return redirect(url_for('.set_report_dates'))
        if dates_db.day_1:
            d_1 = {'d': 'day_1', 'day': days[dates_db.day_1.strftime('%w')],
                   'day_full': days_full[dates_db.day_1.strftime('%w')] + ', ' + dates_db.day_1.strftime('%d') + ' ' + \
                               months_full[dates_db.day_1.strftime('%m')]}
            if d_1 not in c_dates:
                c_dates.append(d_1)
        if dates_db.day_2:
            d_2 = {'d': 'day_2', 'day': days[dates_db.day_2.strftime('%w')],
                   'day_full': days_full[dates_db.day_2.strftime('%w')] + ', ' + dates_db.day_2.strftime(
                       '%d') + ' ' + \
                               months_full[dates_db.day_2.strftime('%m')]}
            if d_2 not in c_dates:
                c_dates.append(d_2)
        if dates_db.day_3:
            d_3 = {'d': 'day_3', 'day': days[dates_db.day_3.strftime('%w')],
                   'day_full': days_full[dates_db.day_3.strftime('%w')] + ', ' + dates_db.day_3.strftime('%d') + ' ' + \
                               months_full[dates_db.day_3.strftime('%m')]}
            if d_3 not in c_dates:
                c_dates.append(d_3)

    for work in works.keys():
        if work in appl_for_online:
            participating += 1
        if works[work]['work_id'] not in reported:
            works_unordered.append(works[work])

    for day in c_dates:
        day_works = []
        for work in works.values():
            if 'report_day' in work.keys() and work['report_day'] == day['d']:
                day_works.append(work)
        day['works'] = sorted(day_works, key=lambda w: w['report_order'])
        if [w['report_order'] for w in day['works']]:
            day['max_order'] = max([w['report_order'] for w in day['works']])

    return render_template('online_reports/reports_order.html', works_unordered=works_unordered,
                           participating=participating, c_dates=c_dates, approved_for_2=approved_for_2,
                           categories=categories, union=union)


@app.route('/works_list_schedule/<cat_id>')
def works_list_schedule(cat_id):
    cat_id = int(cat_id)
    if cat_id in [c.cat_id for c in CategoryUnions.query.all()]:
        union = CategoryUnions.query.filter(CategoryUnions.cat_id == cat_id).first().union_id
        cats = [c.cat_id for c in CategoryUnions.query.filter(CategoryUnions.union_id == union).all()]
        union = True
    else:
        cats = [cat_id]
        union = False
    categories = []
    works_unordered = []
    approved_for_2 = 0
    participating = 0
    c_dates = []
    works = {}

    for cat_id in cats:
        dates_db = db.session.query(ReportDates).filter(ReportDates.cat_id == cat_id).first()
        cat_name = Categories.query.filter(Categories.cat_id == cat_id).first().cat_name
        categories.append({'cat_id': cat_id, 'cat_name': cat_name})
        works.update(
            get_works(cat_id, 2, 'online', appl_info=True, w_payment_info=True, reports_info=True, site_id=True))
        approved_for_2 += len(works)
        if dates_db.day_1:
            d_1 = {'d': 'day_1', 'day': days[dates_db.day_1.strftime('%w')],
                   'day_full': days_full[dates_db.day_1.strftime('%w')] + ', ' + dates_db.day_1.strftime('%d') + ' ' + \
                               months_full[dates_db.day_1.strftime('%m')]}
            if d_1 not in c_dates:
                c_dates.append(d_1)
        if dates_db.day_2:
            d_2 = {'d': 'day_2', 'day': days[dates_db.day_2.strftime('%w')],
                   'day_full': days_full[dates_db.day_2.strftime('%w')] + ', ' + dates_db.day_2.strftime(
                       '%d') + ' ' + \
                               months_full[dates_db.day_2.strftime('%m')]}
            if d_2 not in c_dates:
                c_dates.append(d_2)
        if dates_db.day_3:
            d_3 = {'d': 'day_3', 'day': days[dates_db.day_3.strftime('%w')],
                   'day_full': days_full[dates_db.day_3.strftime('%w')] + ', ' + dates_db.day_3.strftime('%d') + ' ' + \
                               months_full[dates_db.day_3.strftime('%m')]}
            if d_3 not in c_dates:
                c_dates.append(d_3)

    for work in works.keys():
        if works[work]['appl_no']:
            participating += 1
        if works[work]['work_id'] not in [w.work_id for w in ReportOrder.query.all()]:
            works_unordered.append(works[work])

    for day in c_dates:
        day_works = []
        for work in works.values():
            if 'report_day' in work.keys() and work['report_day'] == day['d']:
                day_works.append(work)
        day['works'] = sorted(day_works, key=lambda w: w['report_order'])
        if [w['report_order'] for w in day['works']]:
            day['max_order'] = max([w['report_order'] for w in day['works']])

    return render_template('online_reports/works_list_schedule.html', works_unordered=works_unordered,
                           participating=participating, c_dates=c_dates, approved_for_2=approved_for_2,
                           categories=categories, union=union)


@app.route('/work_date/<cat_id>/<work_id>/<day>/<page>')
def work_date(cat_id, work_id, day, page):
    write_work_date(cat_id=cat_id, work_id=work_id, day=day)
    return redirect(url_for('.' + page, cat_id=cat_id))


@app.route('/report_order_many/<cat_id>', methods=['POST'])
def report_order_many(cat_id):
    cat_id = int(cat_id)
    if cat_id in [c.cat_id for c in CategoryUnions.query.all()]:
        union = CategoryUnions.query.filter(CategoryUnions.cat_id == cat_id).first().union_id
        cats = [c.cat_id for c in CategoryUnions.query.filter(CategoryUnions.union_id == union).all()]
    else:
        cats = [cat_id]

    for cat_id in cats:
        works = get_works(cat_id, 2)
        schedule = {}
        for work in works.values():
            if str(work['work_id']) in request.form.keys():
                schedule[work['work_id']] = request.form[str(work['work_id'])]
        for w in schedule.keys():
            write_work_date(cat_id=cat_id, work_id=w, day=schedule[w])
    return redirect(url_for('.reports_order', cat_id=cat_id))


@app.route('/unorder/<cat_id>/<work_id>')
def unorder(cat_id, work_id):
    cat_id = int(cat_id)
    if cat_id in [c.cat_id for c in CategoryUnions.query.all()]:
        union = CategoryUnions.query.filter(CategoryUnions.cat_id == cat_id).first().union_id
        cats = [c.cat_id for c in CategoryUnions.query.filter(CategoryUnions.union_id == union).all()]
        union = True
    else:
        cats = [cat_id]
        union = False

    work_id = int(work_id)
    work_db = ReportOrder.query.filter(ReportOrder.work_id == work_id).first()
    order_deleted = work_db.order
    if work_id in [w.work_id for w in ReportOrder.query.all()]:
        work = ReportOrder.query.filter(ReportOrder.work_id == work_id).first()
        db.session.delete(work)
        db.session.query(Works).filter(Works.work_id == work_id).update({Works.reported: False})
        db.session.commit()

    works = []
    for cat_id in cats:
        works.extend([w.work_id for w in ReportOrder.query.filter(ReportOrder.cat_id == cat_id)
                     .filter(ReportOrder.report_day == work_db.report_day).all() if w.order > order_deleted])
    for work in works:
        new = ReportOrder.query.filter(ReportOrder.work_id == work).first().order - 1
        db.session.query(ReportOrder).filter(ReportOrder.work_id == work
                                             ).update({ReportOrder.order: new})
        db.session.commit()
    return redirect(url_for('.reports_order', cat_id=cats[0]))


@app.route('/confirm_clear_schedule/<cat_id>')
def confirm_clear_schedule(cat_id):
    dates_db = db.session.query(ReportDates).filter(ReportDates.cat_id == cat_id).first()
    c_dates = []
    if dates_db.day_1:
        d_1 = {'d': 'day_1', 'day': days[dates_db.day_1.strftime('%w')],
               'day_full': days_full[dates_db.day_1.strftime('%w')] + ', ' + dates_db.day_1.strftime('%d') + ' ' +
                           months_full[dates_db.day_1.strftime('%m')]}
        c_dates.append(d_1)
    if dates_db.day_2:
        d_2 = {'d': 'day_2', 'day': days[dates_db.day_2.strftime('%w')],
               'day_full': days_full[dates_db.day_2.strftime('%w')] + ', ' + dates_db.day_2.strftime('%d') + ' ' +
                           months_full[dates_db.day_2.strftime('%m')]}
        c_dates.append(d_2)
    if dates_db.day_3:
        d_3 = {'d': 'day_3', 'day': days[dates_db.day_3.strftime('%w')],
               'day_full': days_full[dates_db.day_3.strftime('%w')] + ', ' + dates_db.day_3.strftime('%d') + ' ' +
                           months_full[dates_db.day_3.strftime('%m')]}
        c_dates.append(d_3)
    return render_template('online_reports/confirm_clear_schedule.html', cat_id=cat_id, c_dates=c_dates)


@app.route('/clear_schedule/<cat_id>/<day>')
def clear_schedule(cat_id, day):
    cat_id = int(cat_id)
    if cat_id in [c.cat_id for c in CategoryUnions.query.all()]:
        union = CategoryUnions.query.filter(CategoryUnions.cat_id == cat_id).first().union_id
        cats = [c.cat_id for c in CategoryUnions.query.filter(CategoryUnions.union_id == union).all()]
        union = True
    else:
        cats = [cat_id]
        union = False

    for cat_id in cats:
        works_db = db.session.query(ReportOrder).filter(ReportOrder.cat_id == int(cat_id))
        works = [w.work_id for w in ReportOrder.query.filter(ReportOrder.cat_id == int(cat_id)).all()]
        for work in works:
            db.session.query(Works).filter(Works.work_id == work).update({Works.reported: False})
            db.session.commit()
        if day == 'all':
            to_delete = works_db
        else:
            to_delete = works_db.filter(ReportOrder.report_day == day)
        to_delete.delete()
        db.session.commit()
    return redirect(url_for('.reports_order', cat_id=cat_id))


@app.route('/reorder/<cat_id>/<work_id>/<direction>')
def reorder(cat_id, work_id, direction):
    cat_id = int(cat_id)
    if cat_id in [c.cat_id for c in CategoryUnions.query.all()]:
        union = CategoryUnions.query.filter(CategoryUnions.cat_id == cat_id).first().union_id
        cats = [c.cat_id for c in CategoryUnions.query.filter(CategoryUnions.union_id == union).all()]
        union = True
    else:
        cats = [cat_id]
        union = False

    order_db = db.session.query(ReportOrder).filter(ReportOrder.work_id == work_id).first()
    order_1 = order_db.order
    day = order_db.report_day
    if direction == 'up':
        order_2 = order_1 - 1
    else:
        order_2 = order_1 + 1

    day_ordered = [o.order for o in ReportOrder.query.filter(ReportOrder.cat_id ==
                                                             cat_id).filter(ReportOrder.report_day == day).all()]
    for cat_id in cats:
        if order_2 in day_ordered:
            db.session.query(ReportOrder).filter(ReportOrder.cat_id == cat_id
                                                 ).filter(ReportOrder.report_day == day
                                                          ).filter(ReportOrder.order == order_2
                                                                   ).update({ReportOrder.order: order_1})
            db.session.commit()
    db.session.query(ReportOrder).filter(ReportOrder.work_id == work_id).update({ReportOrder.order: order_2})
    db.session.commit()
    return redirect(url_for('.reports_order', cat_id=cat_id))


@app.route('/download_schedule/<cat_id>')
def download_schedule(cat_id):
    cat_id = int(cat_id)
    if cat_id in [c.cat_id for c in CategoryUnions.query.all()]:
        union = CategoryUnions.query.filter(CategoryUnions.cat_id == cat_id).first().union_id
        cats = [c.cat_id for c in CategoryUnions.query.filter(CategoryUnions.union_id == union).all()]
        union = True
    else:
        cats = [cat_id]
        union = False

    works = {}
    c_dates = []
    categories = []
    for cat_id in cats:
        cat = Categories.query.filter(Categories.cat_id == cat_id).first()
        categories.append({'cat_id': cat_id, 'cat_name': cat.cat_name, 'short_name': cat.short_name})
        dates_db = db.session.query(ReportDates).filter(ReportDates.cat_id == cat_id).first()
        works.update(
            get_works(cat_id, 2, 'online', appl_info=True, w_payment_info=True, reports_info=True, site_id=True))
        if dates_db.day_1:
            d_1 = {'d': 'day_1', 'day': days[dates_db.day_1.strftime('%w')],
                   'day_full': days_full[dates_db.day_1.strftime('%w')] + ', ' + dates_db.day_1.strftime('%d') + ' ' + \
                               months_full[dates_db.day_1.strftime('%m')]}
            if d_1 not in c_dates:
                c_dates.append(d_1)
        if dates_db.day_2:
            d_2 = {'d': 'day_2', 'day': days[dates_db.day_2.strftime('%w')],
                   'day_full': days_full[dates_db.day_2.strftime('%w')] + ', ' + dates_db.day_2.strftime(
                       '%d') + ' ' + \
                               months_full[dates_db.day_2.strftime('%m')]}
            if d_2 not in c_dates:
                c_dates.append(d_2)
        if dates_db.day_3:
            d_3 = {'d': 'day_3', 'day': days[dates_db.day_3.strftime('%w')],
                   'day_full': days_full[dates_db.day_3.strftime('%w')] + ', ' + dates_db.day_3.strftime('%d') + ' ' + \
                               months_full[dates_db.day_3.strftime('%m')]}
            if d_3 not in c_dates:
                c_dates.append(d_3)

    for day in c_dates:
        day_works = []
        for work in works.values():
            if 'report_day' in work.keys() and work['report_day'] == day['d']:
                day_works.append(work)
        day['works'] = sorted(day_works, key=lambda w: w['report_order'])
        if [w['report_order'] for w in day['works']]:
            day['max_order'] = max([w['report_order'] for w in day['works']])

    if not os.path.exists('static/files/generated_files/schedules/' + str(curr_year)):
        os.makedirs('static/files/generated_files/schedules/' + str(curr_year))
    path = 'static/files/generated_files/schedules/' + str(curr_year) + '/' + 'Расписание ' \
           + ', '.join([c['short_name'] for c in categories]) + '.docx'

    document = document_set()

    if union is True:
        a = 'Расписание заседания секций'
    else:
        a = 'Расписание заседания секции'
    h = a + '\n' + '\n'.join([c['cat_name'] for c in categories])

    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = h
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for day in c_dates:
        document.add_heading(day['day_full'], level=1)
        for work in day['works']:
            document.add_paragraph(str(work['report_order']) + '. ' + str(work['work_id']) + ' – ' + work['work_name'] +
                                   ' – ' + work['authors'], style='Normal')

    document.save(path)
    return send_file(path, as_attachment=True)


@app.route('/category_unions')
def category_unions():
    access = check_access(6)
    if access is not True:
        return access
    unions = []
    for u_id in set([u.union_id for u in CategoryUnions.query.filter(CategoryUnions.year == curr_year).all()]):
        cats = [categories_info(c.cat_id)[1] for c in
                CategoryUnions.query.filter(CategoryUnions.union_id == u_id).all()]
        unions.append({'union_id': u_id, 'categories': cats})
    c, cats = categories_info()
    if unions:
        n = max([u['union_id'] for u in unions]) + 1
    elif [u.union_id for u in CategoryUnions.query.all()]:
        n = max([u.union_id for u in CategoryUnions.query.all()]) + 1
    else:
        n = 1
    return render_template('categories/category_unions.html', unions=unions, cats=cats, n=n)


@app.route('/set_category_union', methods=['POST'])
def set_category_union():
    union_id = request.form['union_id']
    cat_id = request.form['cat_id']
    if union_id in [u.union_id for u in CategoryUnions.query.all()]:
        if cat_id not in [u.cat_id for u in CategoryUnions.query.filter(CategoryUnions.union_id == union_id).all()]:
            u = CategoryUnions(year=curr_year, union_id=union_id, cat_id=cat_id)
            db.session.add(u)
            db.session.commit()
    else:
        u = CategoryUnions(year=curr_year, union_id=union_id, cat_id=cat_id)
        db.session.add(u)
        db.session.commit()
    return redirect(url_for('.category_unions'))


@app.route('/clear_union/<union_id>/<cat_id>')
def clear_union(union_id, cat_id):
    union_id = int(union_id)
    cat_id = int(cat_id)
    if union_id in [u.union_id for u in CategoryUnions.query.all()] and cat_id in [u.cat_id for u in
                                                                                   CategoryUnions.query.filter(
                                                                                       CategoryUnions.union_id == union_id).all()]:
        to_del = db.session.query(CategoryUnions).filter(CategoryUnions.union_id == union_id) \
            .filter(CategoryUnions.cat_id == cat_id).first()
        db.session.delete(to_del)
        db.session.commit()
    return redirect(url_for('.category_unions'))


@app.route('/add_cities')
def add_cities():
    return render_template('works/add_cities.html')


@app.route('/many_cities', methods=['POST'])
def many_cities():
    text = request.form['text']
    lines = text.split('\n')
    for line in lines:
        if line != '':
            info = line.split('\t')
            if info[0] == '':
                country = None
            else:
                country = info[0]
            if info[1] == '':
                region = None
            else:
                region = info[1]
            if info[2] == '':
                area = None
            else:
                area = info[2]
            if info[3] == '':
                city = None
            else:
                city = info[3]
            if info[4] == '':
                timeshift = None
            else:
                timeshift = int(info[4])

            if city not in [c.city for c in Cities.query.all()]:
                city_add = Cities(country, region, area, city, timeshift)
                db.session.add(city_add)
            else:
                cit = db.session.query(Cities).filter(Cities.city == city)
                if region not in [c.region for c in cit.all()]:
                    city_add = Cities(country, region, area, city, timeshift)
                    db.session.add(city_add)
                elif area not in [c.area for c in cit.all()]:
                    city_add = Cities(country, region, area, city, timeshift)
                    db.session.add(city_add)
                else:
                    db.session.query(Cities).filter(Cities.region == region
                                                    ).filter(Cities.area == area
                                                             ).filter(Cities.city == city
                                                                      ).update({Cities.msk_time_shift: timeshift})
            db.session.commit()
    return redirect(url_for('.add_cities'))


@app.route('/reported/<cat_id>/<work_id>/<action>')
def reported(cat_id, work_id, action):
    if action == 'check':
        report = True
    else:
        report = False
    db.session.query(Works).filter(Works.work_id == work_id).update({Works.reported: report})
    db.session.commit()
    return redirect(url_for('.reports_order', cat_id=cat_id))


@app.route('/search_participant', defaults={'query': 'sear'})
@app.route('/search_participant/<query>')
def search_participant(query):
    if isinstance(query, str) and '%' in query:
        query = unquote(query)
    access = check_access(3)
    if access is not True:
        return access
    response = {'type': None, 'value': query}
    if query:
        try:
            qu = int(query)
            if len(query) == 6 and qu:
                response = {'type': 'work', 'works': work_info(int(query), w_payment_info=True, appl_info=True,
                                                               cat_info=True)}
            elif len(query) == 5:
                response = {'type': 'appl', 'value': application_2_tour(int(query))}
            else:
                response = {'type': None}
        except ValueError:
            if query == 'sear':
                response = 'search'
            else:
                parts = [u.participant_id for u in ParticipantsApplied.query.all()
                         if query.lower() == u.last_name.lower()[:len(query)]]
                parts.extend([u.participant_id for u in ParticipantsApplied.query.all()
                              if query.lower() == u.first_name.lower()[:len(query)]])
                parts.extend([u.participant_id for u in ParticipantsApplied.query.all()
                              if query.lower() == u.patronymic_name.lower()[:len(query)]])
                p = []
                for part in parts:
                    appl = ParticipantsApplied.query.filter(ParticipantsApplied.participant_id == part).first().appl_id
                    if appl not in [pa['id'] for pa in p]:
                        partic = application_2_tour(appl)
                        p.append(partic)
                response = {'type': 'appls', 'value': p}

                works = [u.work_id for u in Works.query.all()
                         if query.lower() == u.author_1_name.lower()[:len(query)]]
                works.extend([u.work_id for u in Works.query.all()
                              if u.author_2_name and query.lower() == u.author_2_name.lower()[:len(query)]])
                works.extend([u.work_id for u in Works.query.all()
                              if u.author_3_name and query.lower() == u.author_3_name.lower()[:len(query)]])
                w = [work_info(wo, w_payment_info=True, appl_info=True, cat_info=True) for wo in works]
                if response['type']:
                    response['works'] = w
                else:
                    response = {'type': 'name', 'value': w}
                if not parts and not works:
                    response = {'type': None, 'value': query}
    else:
        response = {'type': None, 'value': query}
    return render_template('participants_and_payment/search_participant.html', response=response)


@app.route('/searching_participant', methods=['GET'])
def searching_participant():
    renew_session()
    query = request.values.get('query', str)
    return redirect(url_for('.search_participant', query=query))


@app.route('/searching_payment', methods=['GET'])
def searching_payment():
    renew_session()
    query = request.values.get('query', str)
    return redirect(url_for('.id_payments', mode=query, page=1, length=30, query=query))


@app.route('/unpayed')
def unpayed():
    access = check_access(8)
    if access is not True:
        return access
    applied = [work_info(w.work_id, w_payment_info=True, cat_info=True, additional_info=True)
               for w in AppliedForOnline.query.all()
               if w.work_id not in
               [p.participant for p in PaymentRegistration.query.filter(PaymentRegistration.for_work == 1).all()]]
    unpayed = [w for w in applied if w['payed'] is False]

    file = [{'Номер работы': w['work_id'], 'Название работы': w['work_name'], 'Авторы': w['authors'],
             'Название секции': w['cat_name'], 'Сумма оргвзноса': str(w['fee']) + ' р.', 'e-mail': w['email']}
            for w in unpayed]

    df = pd.DataFrame(data=file)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/unpayed.xlsx') as writer:
        df.to_excel(writer, sheet_name='Нет оплаты')
    return render_template('online_reports/unpayed.html', unpayed=unpayed)


@app.route('/download_unpayed')
def download_unpayed():
    return send_file('static/files/generated_files/unpayed.xlsx', as_attachment=True)


@app.route('/works_participated')
def works_participated():
    access = check_access(5)
    if access is not True:
        return access
    wks = [w.work_id for w in ParticipatedWorks.query.all()]
    works = [work_info(w) for w in wks if str(w)[:2] == str(curr_year)[-2:]]
    for work in works:
        work['link_name'] = work['work_name'].strip('?')
    return render_template('works/works_participated.html', works=works)


@app.route('/delete_participated_work/<work_id>')
def delete_participated_work(work_id):
    to_del = db.session.query(ParticipatedWorks).filter(ParticipatedWorks.work_id == int(work_id)).first()
    db.session.delete(to_del)
    db.session.commit()
    return redirect(url_for('.works_participated'))


@app.route('/online_participants', defaults={'length': 30, 'page': 1})
@app.route('/online_participants/<length>/<page>')
def online_participants(length, page):
    access = check_access(3)
    if access is not True:
        return access
    wks = [w.work_id for w in AppliedForOnline.query.all()]
    applied = len(wks)
    n, data = make_pages(length, wks, page)
    works = [work_info(w, reports_info=True) for w in data if str(w)[:2] == str(curr_year)[-2:]]
    for work in works:
        work['link_name'] = work['work_name'].strip('?')
    reported = len([w.work_id for w in AppliedForOnline.query
                   .join(Works, AppliedForOnline.work_id == Works.work_id).filter(Works.reported == 1).all()])
    return render_template('online_reports/online_participants.html', works=works, pages=n, page=page,
                           length=length, link='online_participants', applied=applied, reported=reported)


@app.route('/download_online_participants')
def download_online_participants():
    works = [work_info(w.work_id, reports_info=True, w_payment_info=True, cat_info=True, additional_info=True)
             for w in AppliedForOnline.query.all()
             if str(w.work_id)[:2] == str(curr_year)[-2:]]
    file = [{'Номер работы': w['work_id'], 'Название работы': w['work_name'], 'Авторы': w['authors'],
             'Руководитель': w['supervisor'], 'email': w['email'],
             'Название секции': w['cat_name'], 'Оргвзнос оплачен': w['payed'], 'Выступил': w['reported']}
            for w in works]

    df = pd.DataFrame(data=file)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/online_participants.xlsx') as writer:
        df.to_excel(writer, sheet_name='Участники онлайна')
    return send_file('static/files/generated_files/online_participants.xlsx', as_attachment=True)


@app.route('/download_online_participants_html')
def download_online_participants_html():
    c, cats = categories_info()
    cat_works = [(w.cat_id, w.work_id) for w in
                 WorkCategories.query.join(AppliedForOnline, WorkCategories.work_id == AppliedForOnline.work_id).all()
                 if w.cat_id in [c['id'] for c in cats]]
    for cat in cats:
        cat['online_participants'] = [work_info(w[1], reports_info=True) for w in cat_works if w[0] == cat['id']]

    c_w = [cat for cat in cats if len(cat['online_participants']) > 0]

    with open('static/files/generated_files/online_participants_' + str(curr_year) + '.html', 'w',
              encoding='utf-8') as f:
        f.write(render_template('online_reports/online_participants_html.html', cats=c_w))
    return send_file('static/files/generated_files/online_participants_' + str(curr_year) + '.html', as_attachment=True)


@app.route('/delete_online_participant/<work_id>')
def delete_online_participant(work_id):
    to_del = db.session.query(AppliedForOnline).filter(AppliedForOnline.work_id == int(work_id)).first()
    db.session.delete(to_del)
    db.session.commit()
    return redirect(url_for('.online_participants'))


@app.route('/online_participants_applications/<one_cat>', defaults={'length': 50, 'page': 1})
@app.route('/online_participants_applications/<one_cat>/<length>/<page>')
def online_participants_applications(one_cat, length, page):
    access = check_access(3)
    if access is not True:
        return access
    c, cats = categories_info()
    if one_cat == 'all':
        wks = [w.work_id for w in AppliedForOnline.query
        .join(WorkOrganisations, AppliedForOnline.work_id == WorkOrganisations.work_id)
        .join(WorkStatuses, AppliedForOnline.work_id == WorkStatuses.work_id)
        .join(Works, AppliedForOnline.work_id == Works.work_id)
        .join(OrganisationApplication, WorkOrganisations.organisation_id == OrganisationApplication.organisation_id)
        .join(Applications2Tour, AppliedForOnline.work_id == Applications2Tour.work_id)
        .filter(Works.reported == 1)
        .order_by(Applications2Tour.appl_no)
        .order_by(OrganisationApplication.arrived)
        .order_by(OrganisationApplication.appl_no)
        .order_by(WorkOrganisations.organisation_id)
        .order_by(WorkStatuses.status_id).all()]
        works = [w for w in wks if str(w)[:2] == str(curr_year)[-2:]]
        w_len = len(works)
        n, data = make_pages(length, works, page)
        works = [work_info(w, organisation_info=True, appl_info=True, status_info=True) for w in data]
        one_cat = 'all'
    else:
        if '{' in one_cat:
            one_cat = json.loads(one_cat.replace("'", '"'))
        else:
            cat_id = int(one_cat)
            one_cat = {'cat_id': cat_id, 'short_name': Categories.query.filter(Categories.cat_id == cat_id)
            .first().short_name}
        wks = [w.work_id for w in AppliedForOnline.query
        .join(WorkCategories, AppliedForOnline.work_id == WorkCategories.work_id)
        .join(WorkStatuses, AppliedForOnline.work_id == WorkStatuses.work_id)
        .join(Works, AppliedForOnline.work_id == Works.work_id)
        .filter(WorkCategories.cat_id == one_cat['cat_id'])
        .filter(Works.reported == 1)
        .order_by(WorkStatuses.status_id).all()]
        works = [work_info(w, organisation_info=True, appl_info=True, status_info=True) for w in wks]
        w_len = len(works)
        # works = sorted(works_applied, key=lambda x: x['organisation_id'])
        n = 1
    return render_template('online_reports/online_participants_applications.html', works=works, pages=n, page=page,
                           length=length, link='online_participants_applications/all', cats=cats, one_cat=one_cat,
                           w_len=w_len)


@app.route('/renew_applications/<one_cat>/<q_type>/<q_id>')
def renew_applications(one_cat, q_type, q_id):
    q_id = int(q_id)
    response = json.loads(requests.post(url="https://vernadsky.info/second-tour-requests-json/" + str(curr_year) + "/",
                                        headers=mail_data.headers).text)
    if q_type == 'work':
        for a in response:
            for w in a['works']:
                if int(w['number']) == q_id:
                    application = a
                    break
                else:
                    application = None
            if application:
                break
    elif q_type == 'appl':
        for a in response:
            if int(a['id']) == q_id:
                application = a
                break
            else:
                application = None
    elif q_type == 'org':
        for a in response:
            if int(a['organization']['id']) == q_id:
                application = a
                break
            else:
                application = None
    else:
        application = None

    if application:
        appl_no = int(application['id'])
        organisation_id = int(application['organization']['id'])
        works = [int(w['number']) for w in application['works']]
        arrival = bool(application['arrival'])
        if OrganisationApplication(organisation_id, appl_no, arrival) not in OrganisationApplication.query.all():
            if organisation_id in [o.organisation_id for o in OrganisationApplication.query.all()]:
                db.session.query(OrganisationApplication) \
                    .filter(OrganisationApplication.organisation_id == organisation_id) \
                    .update({OrganisationApplication.appl_no: appl_no,
                             OrganisationApplication.arrived: arrival})
                db.session.commit()
            else:
                o = OrganisationApplication(organisation_id, appl_no, arrival)
                db.session.add(o)
                db.session.commit()
        for work_id in works:
            if Applications2Tour(work_id, appl_no, arrival) not in Applications2Tour.query.all():
                if work_id in [w.work_id for w in Applications2Tour.query.all()]:
                    db.session.query(Applications2Tour) \
                        .filter(Applications2Tour.work_id == work_id) \
                        .update({Applications2Tour.appl_no: appl_no,
                                 Applications2Tour.arrived: arrival})
                    db.session.commit()
                else:
                    a = Applications2Tour(work_id, appl_no, arrival)
                    db.session.add(a)
                    db.session.commit()

    return redirect(url_for('.online_participants_applications', one_cat=one_cat))


@app.route('/renew_organisations/<one_cat>', defaults={'which': 'all'})
@app.route('/renew_organisations/<one_cat>/<which>')
def renew_organisations(one_cat, which):
    if which == 'online':
        wks = [w.work_id for w in AppliedForOnline.query.all()]
        works = [w for w in wks if str(w)[:2] == str(curr_year)[-2:]]
    else:
        wks = [w.work_id for w in Works.query.all()]
        works = [w for w in wks if str(w)[:2] == str(curr_year)[-2:]]
    response = json.loads(requests.post(url="https://vernadsky.info/all-works-json/" + str(curr_year) + "/",
                                        headers=mail_data.headers).text)
    for w in response:
        if int(w['number']) in works:
            organisation = {'work_id': int(w['number']), 'organisation_id': int(w['organization']['id']),
                            'name': w['organization']['name'], 'city': w['organization']['city'],
                            'country': w['organization']['country']}

            if organisation['organisation_id'] in [o.organisation_id for o in Organisations.query.all()]:
                a = Organisations(organisation['organisation_id'], organisation['name'], organisation['city'],
                                  organisation['country'])
                if Organisations.query.filter(Organisations.organisation_id == organisation['organisation_id']) \
                        .first() != a:
                    db.session.query(Organisations).filter(
                        Organisations.organisation_id == organisation['organisation_id']) \
                        .update({Organisations.organisation_id: organisation['organisation_id'],
                                 Organisations.name: organisation['name'],
                                 Organisations.city: organisation['city'],
                                 Organisations.country: organisation['country']})
                    db.session.commit()
            else:
                a = Organisations(organisation['organisation_id'], organisation['name'], organisation['city'],
                                  organisation['country'])
                db.session.add(a)
                db.session.commit()

            if organisation['work_id'] in [w.work_id for w in WorkOrganisations.query.all()]:
                a = WorkOrganisations(organisation['work_id'], organisation['organisation_id'])
                if WorkOrganisations.query.filter(WorkOrganisations.work_id == organisation['work_id']) \
                        .first() != a:
                    db.session.query(WorkOrganisations).filter(WorkOrganisations.work_id == organisation['work_id']) \
                        .update({WorkOrganisations.organisation_id: organisation['organisation_id']})
                    db.session.commit()
            else:
                a = WorkOrganisations(organisation['work_id'], organisation['organisation_id'])
                db.session.add(a)
                db.session.commit()

    if one_cat == 'all':
        length = 30
        page = 1
    else:
        length = 'all'
        page = 1
    return redirect(url_for('.online_participants_applications', one_cat=one_cat, length=length, page=page))


@app.route('/download_online_reported')
def download_online_reported():
    works = [{'Номер работы': w.work_id, 'Название': w.work_name, 'e-mail': w.email} for w
             in Works.query.join(AppliedForOnline, Works.work_id == AppliedForOnline.work_id)
             .filter(Works.reported == 1).all()]
    for w in works:
        if w['Номер работы'] in [wk.participant for wk
                                 in PaymentRegistration.query.filter(PaymentRegistration.for_work == 1).all()]:
            w['Оплата'] = 'Да'
        else:
            w['Оплата'] = 'Нет'
    df = pd.DataFrame(data=works)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/online_reported.xlsx') as writer:
        df.to_excel(writer, sheet_name='Работы, выступившие онлайн')
    return send_file('static/files/generated_files/online_reported.xlsx', as_attachment=True)


# @app.route('/experts/<cat_id>', defaults={'expert_to_edit': None})
# @app.route('/experts/<cat_id>/<expert_to_edit>')
# def experts(cat_id, expert_to_edit):
#     if cat_id != 'all':
#         cat_id = int(cat_id)
#     c, cat = categories_info(cat_id)
#     cats = []
#     if type(cat) != list:
#         cats.append(cat)
#     else:
#         cats = cat

#     all_exps = [{'id': e.expert_id, 'name': e.last_name + ' ' + e.first_name + ' ' + e.patronymic}
#                 for e in Experts.query.filter(Experts.year == curr_year).all()]
#     all_exps = sorted(all_exps, key=lambda e: e['name'])

#     for cat in cats:
#         cat_experts = []
#         exps = [e.expert_id for e in CatExperts.query.filter(CatExperts.cat_id == cat['id']).all()]
#         for e in exps:
#             e_db = db.session.query(Experts).filter(Experts.expert_id == e).first()
#             e_t_db = db.session.query(CatExperts).filter(CatExperts.cat_id == cat['id']) \
#                 .filter(CatExperts.expert_id == e).first()
#             if e_t_db:
#                 if e_t_db.day_1_started is not None:
#                     day_1_start = e_t_db.day_1_started.strftime('%H:%M')
#                 else:
#                     day_1_start = ''
#                 if e_t_db.day_1_finished is not None:
#                     day_1_end = e_t_db.day_1_finished.strftime('%H:%M')
#                 else:
#                     day_1_end = ''
#                 if e_t_db.day_2_started is not None:
#                     day_2_start = e_t_db.day_2_started.strftime('%H:%M')
#                 else:
#                     day_2_start = ''
#                 if e_t_db.day_2_finished is not None:
#                     day_2_end = e_t_db.day_2_finished.strftime('%H:%M')
#                 else:
#                     day_2_end = ''
#                 if e_t_db.day_3_started is not None:
#                     day_3_start = e_t_db.day_3_started.strftime('%H:%M')
#                 else:
#                     day_3_start = ''
#                 if e_t_db.day_3_finished is not None:
#                     day_3_end = e_t_db.day_3_finished.strftime('%H:%M')
#                 else:
#                     day_3_end = ''
#             else:
#                 day_1_start = ''
#                 day_1_end = ''
#                 day_2_start = ''
#                 day_2_end = ''
#                 day_3_start = ''
#                 day_3_end = ''
#             cat_experts.append({'expert_id': e_db.expert_id,
#                                 'last_name': e_db.last_name,
#                                 'first_name': e_db.first_name,
#                                 'patronymic': e_db.patronymic,
#                                 'email': e_db.email,
#                                 'degree': e_db.degree,
#                                 'place_of_work': e_db.place_of_work,
#                                 'day_1_start': day_1_start,
#                                 'day_1_end': day_1_end,
#                                 'day_2_start': day_2_start,
#                                 'day_2_end': day_2_end,
#                                 'day_3_start': day_3_start,
#                                 'day_3_end': day_3_end})
#         cat_experts = sorted(cat_experts, key=lambda e: e['first_name'])
#         cat_experts = sorted(cat_experts, key=lambda e: e['last_name'])
#         cat['experts'] = cat_experts

#         c_dates = {}
#         if cat['id'] in [c.cat_id for c in ReportDates.query.all()]:
#             dates_db = db.session.query(ReportDates).filter(ReportDates.cat_id == cat['id']).first()
#             if dates_db.day_1:
#                 c_dates['d_1'] = days[dates_db.day_1.strftime('%w')] + ', ' + dates_db.day_1.strftime('%d.%m')
#             else:
#                 c_dates['d_1'] = None
#             if dates_db.day_2:
#                 c_dates['d_2'] = days[dates_db.day_2.strftime('%w')] + ', ' + dates_db.day_2.strftime('%d.%m')
#             else:
#                 c_dates['d_2'] = None
#             if dates_db.day_3:
#                 c_dates['d_3'] = days[dates_db.day_3.strftime('%w')] + ', ' + dates_db.day_3.strftime('%d.%m')
#             else:
#                 c_dates['d_3'] = None
#         else:
#             c_dates['day_1'] = None
#             c_dates['day_2'] = None
#             c_dates['day_3'] = None
#             # cat_dates.append(c_dates)
#         cat['c_dates'] = c_dates

#     if expert_to_edit:
#         exp = int(expert_to_edit)
#         e_db = db.session.query(Experts).filter(Experts.expert_id == exp).first()
#         expert_to_edit = {'expert_id': e_db.expert_id,
#                           'last_name': e_db.last_name,
#                           'first_name': e_db.first_name,
#                           'patronymic': e_db.patronymic,
#                           'email': e_db.email,
#                           'degree': e_db.degree,
#                           'place_of_work': e_db.place_of_work}

#     all_days = set(
#         d.day_1 for d in ReportDates.query.all() if d.cat_id in [c['id'] for c in cats] and d.day_1 is not None)
#     all_days.update(
#         set(d.day_2 for d in ReportDates.query.all() if d.cat_id in [c['id'] for c in cats] and d.day_2 is not None))
#     all_days.update(
#         set(d.day_3 for d in ReportDates.query.all() if d.cat_id in [c['id'] for c in cats] and d.day_3 is not None))
#     a_days = sorted(all_days)
#     all_days = [days[a.strftime('%w')] + ', ' + a.strftime('%d.%m') for a in a_days]
#     return render_template('supervisors/experts.html', expert_to_edit=expert_to_edit, cats=cats, all_days=all_days,
#                            all_exps=all_exps, cat_id=cat_id)


# @app.route('/expert_time/<cat_id>', methods=['POST'])
# def expert_time(cat_id):
#     cat_id = int(cat_id)
#     for expert in [e.expert_id for e in CatExperts.query.filter(CatExperts.cat_id == cat_id).all()]:
#         if 'day_1_start/' + str(expert) in request.form.keys() and request.form['day_1_start/' + str(expert)] != '':
#             day_1_start = datetime.datetime.strptime(request.form['day_1_start/' + str(expert)], '%H:%M').time()
#         else:
#             day_1_start = None
#         if 'day_1_end/' + str(expert) in request.form.keys() and request.form['day_1_end/' + str(expert)] != '':
#             day_1_end = datetime.datetime.strptime(request.form['day_1_end/' + str(expert)], '%H:%M').time()
#         else:
#             day_1_end = None
#         if 'day_2_start/' + str(expert) in request.form.keys() and request.form['day_2_start/' + str(expert)] != '':
#             day_2_start = datetime.datetime.strptime(request.form['day_2_start/' + str(expert)], '%H:%M').time()
#         else:
#             day_2_start = None
#         if 'day_2_end/' + str(expert) in request.form.keys() and request.form['day_2_end/' + str(expert)] != '':
#             day_2_end = datetime.datetime.strptime(request.form['day_2_end/' + str(expert)], '%H:%M').time()
#         else:
#             day_2_end = None
#         if 'day_3_start/' + str(expert) in request.form.keys() and request.form['day_3_start/' + str(expert)] != '':
#             day_3_start = datetime.datetime.strptime(request.form['day_3_start/' + str(expert)], '%H:%M').time()
#         else:
#             day_3_start = None
#         if 'day_3_end/' + str(expert) in request.form.keys() and request.form['day_3_end/' + str(expert)] != '':
#             day_3_end = datetime.datetime.strptime(request.form['day_3_end/' + str(expert)], '%H:%M').time()
#         else:
#             day_3_end = None
#         c_e = CatExperts(expert, cat_id, day_1_start, day_1_end, day_2_start, day_2_end, day_3_start, day_3_end)
#         if c_e not in CatExperts.query.all():
#             db.session.query(CatExperts).filter(CatExperts.cat_id == cat_id).filter(CatExperts.expert_id == expert) \
#                 .update({CatExperts.day_1_started: day_1_start,
#                          CatExperts.day_1_finished: day_1_end,
#                          CatExperts.day_2_started: day_2_start,
#                          CatExperts.day_2_finished: day_2_end,
#                          CatExperts.day_3_started: day_3_start,
#                          CatExperts.day_3_finished: day_3_end})
#             db.session.commit()
#     return redirect(url_for('.experts', cat_id=cat_id, expert_to_edit=None))


# @app.route('/save_expert/<cat_id>', methods=['POST'])
# def save_expert(cat_id):
#     cat_id = int(cat_id)
#     last_name = request.form['last_name']
#     first_name = request.form['first_name']
#     patronymic = request.form['patronymic']
#     email = request.form['email']
#     degree = request.form['degree']
#     place_of_work = request.form['place_of_work']
#     expert = Experts(last_name=last_name, first_name=first_name, patronymic=patronymic, email=email, degree=degree,
#                      place_of_work=place_of_work, year=curr_year)
#     if 'expert_id' in request.form.keys():
#         expert_id = int(request.form['expert_id'])
#         if expert not in Experts.query.filter(Experts.expert_id == expert_id).all():
#             db.session.query(Experts).filter(Experts.expert_id == expert_id) \
#                 .update({Experts.last_name: last_name,
#                          Experts.first_name: first_name,
#                          Experts.patronymic: patronymic,
#                          Experts.email: email,
#                          Experts.degree: degree,
#                          Experts.place_of_work: place_of_work,
#                          Experts.year: curr_year})
#             db.session.commit()
#     else:
#         db.session.add(expert)
#         db.session.flush()
#         db.session.commit()
#         expert_id = expert.expert_id
#     cat_expert = CatExperts(expert_id, cat_id, None, None, None, None, None, None)
#     if expert_id in [e.expert_id for e in CatExperts.query.all()]:
#         if cat_id not in [e.cat_id for e in CatExperts.query.filter(CatExperts.expert_id == expert_id).all()]:
#             db.session.query(CatExperts).filter(CatExperts.expert_id == expert_id).update({CatExperts.cat_id: cat_id})
#             db.session.commit()
#     else:
#         db.session.add(cat_expert)
#         db.session.commit()
#     return redirect(url_for('.experts', cat_id=cat_id, expert_to_edit=None))


# @app.route('/add_existing_expert/<cat_id>/<expert_id>')
# def add_existing_expert(cat_id, expert_id):
#     cat_id = int(cat_id)
#     expert_id = int(expert_id)
#     c_e = CatExperts(expert_id, cat_id, None, None, None, None, None, None)
#     if expert_id in [e.expert_id for e in Experts.query.all()]:
#         if expert_id in [e.expert_id for e in CatExperts.query.all()]:
#             if cat_id not in [e.cat_id for e in CatExperts.query.filter(CatExperts.expert_id == expert_id).all()]:
#                 db.session.add(c_e)
#                 db.session.commit()
#         else:
#             db.session.add(c_e)
#             db.session.commit()
#     return redirect(url_for('.experts', cat_id=cat_id, expert_to_edit=None))


# @app.route('/delete_expert/<cat_id>/<expert_id>')
# def delete_expert(cat_id, expert_id):
#     cat_id = int(cat_id)
#     expert_id = int(expert_id)
#     if expert_id in [e.expert_id for e in CatExperts.query.all()]:
#         if cat_id in [e.cat_id for e in CatExperts.query.filter(CatExperts.expert_id == expert_id).all()]:
#             to_del = db.session.query(CatExperts).filter(CatExperts.expert_id == expert_id) \
#                 .filter(CatExperts.cat_id == cat_id).first()
#             db.session.delete(to_del)
#             db.session.commit()
#     return redirect(url_for('.experts', cat_id=cat_id, expert_to_edit=None))


@app.route('/discount_and_participation_mode/<part_id>')
def discount_and_participation_mode(part_id):
    if len(part_id) == 5:
        info = application_2_tour(int(part_id))
        info['type'] = 'application'
    elif len(part_id) == 6:
        info = work_info(int(part_id), w_payment_info=True)
        info['type'] = 'work'
    else:
        return redirect(url_for('.search_participant'))
    return render_template('participants_and_payment/discount_and_participation_mode.html', info=info,
                           full_fee=fee, discounted=tour_fee)


@app.route('/set_fee/<part_id>', methods=['POST'])
def set_fee(part_id):
    if len(part_id) == 6:
        part_fee = int(request.form[str(part_id) + ';fee'])
        part_format = request.form[str(part_id) + ';format']
        if int(part_id) in [p.work_id for p in Discounts.query.all()]:
            db.session.query(Discounts).filter(Discounts.work_id == int(part_id)
                                               ).update({Discounts.payment: part_fee,
                                                         Discounts.participation_format: part_format})
        else:
            discount = Discounts(None, int(part_id), part_fee, part_format)
            db.session.add(discount)
        db.session.commit()
    elif len(part_id) == 5:
        for participant in [p['id'] for p in application_2_tour(part_id)['participants']]:
            part_fee = int(request.form[str(participant) + ';fee'])
            part_format = request.form[str(participant) + ';format']
            if int(participant) in [p.participant_id for p in Discounts.query.all()]:
                db.session.query(Discounts).filter(Discounts.participant_id == int(participant)
                                                   ).update({Discounts.payment: part_fee,
                                                             Discounts.participation_format: part_format})
            else:
                discount = Discounts(int(participant), None, part_fee, part_format)
                db.session.add(discount)
            db.session.commit()
    return redirect(url_for('.search_participant', query=part_id))


@app.route('/load_statement', defaults={'success': False})
@app.route('/load_statement/<success>')
def load_statement(success):
    renew_session()
    return render_template('participants_and_payment/load_statement.html', success=success)


@app.route('/add_bank_statement', methods=['POST'])
def add_bank_statement():
    data = request.files['file'].read().decode('ptcp154')
    lines = data.split('\n')
    statement = []
    existing = [{'d_c': 'C', 'date_oper': datetime.datetime.strftime(p.date, '%d.%m.%Y'), 'number': str(p.order_id),
                 'sum_val': p.debit, 'plat_name': p.organisation, 'plat_inn': str(p.tin),
                 'plat_bic': str(p.bic), 'plat_bank': p.bank_name, 'plat_acc': str(p.account),
                 'text70': p.payment_comment}
                for p in BankStatement.query.all()]
    p_types_existing = {p.payment_id: p.payment_type for p in PaymentTypes.query.all()}
    for line in lines[2:]:
        if line != '':
            sta = {name: value for name, value in zip(lines[0].split('\t'), line.split('\t'))}
            statement.append(sta)
    for payment in statement:
        if payment != {}:
            p = {'d_c': payment['d_c'], 'date_oper': payment['date_oper'], 'number': payment['number'],
                 'sum_val': float(payment['sum_val'].replace(',', '.')), 'plat_name': payment['plat_name'], 'plat_inn':
                     payment['plat_inn'],
                 'plat_bic': payment['plat_bic'], 'plat_bank': payment['plat_bank'], 'plat_acc': payment['plat_acc'],
                 'text70': payment['text70']}
            if p not in existing:
                payment['date_oper'] = datetime.datetime.strptime(payment['date_oper'], '%d.%m.%Y')
                if payment['date_oper'] != datetime.datetime.now().date:
                    if payment['d_c'] == 'C':
                        pay = BankStatement(date=payment['date_oper'], order_id=payment['number'],
                                            debit=float(payment['sum_val'].replace(',', '.')), credit=0,
                                            organisation=payment['plat_name'], tin=payment['plat_inn'],
                                            bic=payment['plat_bic'],
                                            bank_name=payment['plat_bank'], account=payment['plat_acc'],
                                            payment_comment=payment['text70'], alternative=None,
                                            alternative_comment=None)
                        db.session.add(pay)
                        db.session.commit()
                        db.session.flush()
                        payment_id = pay.payment_id
                        if payment_id not in p_types_existing:
                            p_t = PaymentTypes(payment_id, 'Чтения Вернадского')
                            db.session.add(p_t)
                            db.session.commit()
                        else:
                            if p_types_existing[payment_id] != 'Чтения Вернадского':
                                db.session.query(PaymentTypes) \
                                    .filter(PaymentTypes.payment_id == payment_id) \
                                    .update({PaymentTypes.payment_type: 'Чтения Вернадского'})
                                db.session.commit()
                            else:
                                pass
                    # else:
                    #     pay = BankStatement(date=payment['date_oper'], order_id=payment['number'],
                    #                         debit=0, credit=float(payment['sum_val'].replace(',', '.')),
                    #                         organisation=payment['plat_name'], tin=payment['plat_inn'],
                    #                         bic=payment['plat_bic'],
                    #                         bank_name=payment['plat_bank'], account=payment['plat_acc'],
                    #                         payment_comment=payment['text70'], alternative=None, alternative_comment=None)
                    #     db.session.add(pay)
                    #     db.session.commit()
    return redirect(url_for('.load_statement', success=True))


@app.route('/payment_stats')
def payment_stats():
    access = check_access(10)
    if access is not True:
        return access
    statement_db = db.session.query(BankStatement) \
        .join(PaymentTypes, BankStatement.payment_id == PaymentTypes.payment_id)
    clauses = []
    for t in set(p.payment_type for p in PaymentTypes.query.all()):
        t_payments = statement_db.filter(PaymentTypes.payment_type == t).all()
        s = sum([p.debit for p in t_payments])
        if s % 1 == 0:
            s = f'{int(s):,}'.replace(',', ' ')
        else:
            s = f'{s:,}'.replace(',', ' ')
        clauses.append({'name': t, 'sum': s})
    online = [w.work_id for w in AppliedForOnline.query.all()]
    online_sum = 0
    online_unpayed = 0
    reported_unpayed = 0

    for work in online:
        w = Works.query.filter(Works.work_id == work).first()
        reg_tour = w.reg_tour
        reported = w.reported
        if work in [w.work_id for w in Discounts.query.all()]:
            disc = db.session.query(Discounts).filter(Discounts.work_id == work).first()
            w_fee = disc.payment
        elif work in [w.work_id for w in WorksNoFee.query.all()]:
            w_fee = 0
        elif reg_tour is not None:
            w_fee = tour_fee
        else:
            w_fee = fee

        if work in [w.participant for w in PaymentRegistration.query.all()]:
            online_sum += w_fee
        else:
            online_unpayed += w_fee
            if reported is True:
                reported_unpayed += w_fee

    if online_sum % 1 == 0:
        online_sum = f'{int(online_sum):,}'.replace(',', ' ')
    else:
        online_sum = f'{online_sum:,}'.replace(',', ' ')
    if online_unpayed % 1 == 0:
        online_unpayed = f'{int(online_unpayed):,}'.replace(',', ' ')
    else:
        online_unpayed = f'{online_unpayed:,}'.replace(',', ' ')
    if reported_unpayed % 1 == 0:
        reported_unpayed = f'{int(reported_unpayed):,}'.replace(',', ' ')
    else:
        reported_unpayed = f'{reported_unpayed:,}'.replace(',', ' ')

    return render_template('participants_and_payment/payment_stats.html', year=curr_year, clauses=clauses,
                           online_sum=online_sum, online_unpayed=online_unpayed, reported_unpayed=reported_unpayed)


@app.route('/alternative_payments', defaults={'edit': None, 'payment_id': None, 'length': 30, 'page': 1})
@app.route('/alternative_payments/<edit>/<payment_id>/<length>/<page>')
def alternative_payments(edit, payment_id, length, page):
    access = check_access(8)
    if access is not True:
        return access
    length = int(length)
    page = int(page)
    if edit:
        payment_id = int(payment_id)
        if payment_id in [p.payment_id for p in BankStatement.query.filter(BankStatement.alternative == 1).all()]:
            payment_to_edit = payment_info(payment_id)
            date = datetime.datetime.strptime(payment_to_edit['date'], '%d.%m.%Y').strftime('%Y-%m-%d')
            payment_to_edit['date'] = date
            payment_to_edit['debit'] = payment_to_edit['debit'].strip(' р.')
        else:
            payment_to_edit = None
    else:
        payment_to_edit = None
    payments = [p.payment_id for p in BankStatement.query.filter(BankStatement.alternative == 1).all()]
    if edit:
        payments.remove(payment_id)
    n, data = make_pages(length, payments, page)
    payments = statement_info(payments)
    return render_template('participants_and_payment/alternative_payments.html', payments=payments, pages=n, page=page,
                           length=length, link='alternative_payments', payment_to_edit=payment_to_edit)


@app.route('/add_alternative_payment', methods=['POST'])
def add_alternative_payment():
    date = datetime.datetime.strptime(request.form['date'], '%Y-%m-%d').date()
    debit = float(request.form['debit'].replace(',', '.'))
    organisation = request.form['organisation']
    payment_comment = request.form['payment_comment']
    alternative = True
    alternative_comment = request.form['alternative_comment']
    if 'payment_id' in request.form.keys():
        payment_id = int(request.form['payment_id'])
        if payment_id in [p.payment_id for p in BankStatement.query.all()]:
            db.session.query(BankStatement).filter(BankStatement.payment_id == payment_id) \
                .update({BankStatement.date: date,
                         BankStatement.organisation: organisation,
                         BankStatement.payment_comment: payment_comment,
                         BankStatement.alternative: alternative,
                         BankStatement.alternative_comment: alternative_comment})
            db.session.commit()
    else:
        p = BankStatement(date=date, order_id=None, debit=debit, credit=None, organisation=organisation, tin=None,
                          bic=None, bank_name=None, account=None, payment_comment=payment_comment,
                          alternative=alternative, alternative_comment=alternative_comment)
        db.session.add(p)
        db.session.commit()
    return redirect(url_for('.alternative_payments'))


@app.route('/delete_alternative/<payment_id>')
def delete_alternative(payment_id):
    payment_id = int(payment_id)
    if payment_id in [p.payment_id for p in BankStatement.query.filter(BankStatement.alternative == 1).all()]:
        to_del = db.session.query(BankStatement).filter(BankStatement.alternative == 1) \
            .filter(BankStatement.payment_id == payment_id).first()
        db.session.delete(to_del)
        db.session.commit()
    return redirect(url_for('.alternative_payments'))


@app.route('/manage_payments', defaults={'query': 'all', 'length': 30, 'page': 1})
@app.route('/manage_payments/<query>/<length>/<page>')
def manage_payments(query, length, page):
    access = check_access(8)
    if access is not True:
        return access
    if query == 'all':
        q = BankStatement.query.order_by(BankStatement.date.desc()).order_by(BankStatement.order_id.asc()).all()
        query_name = 'Все'
    else:
        q = BankStatement.query \
            .join(PaymentTypes, BankStatement.payment_id == PaymentTypes.payment_id) \
            .order_by(BankStatement.date.desc()).order_by(BankStatement.order_id.asc()) \
            .filter(PaymentTypes.payment_type == query).all()
    payments = [p.payment_id for p in q]
    n, data = make_pages(length, payments, page)
    statement = statement_info(data)
    types = set(p.payment_type for p in PaymentTypes.query.all())
    return render_template('participants_and_payment/manage_payments.html', statement=statement, pages=n, page=page,
                           length=length, link='manage_payments/' + query, types=types, query=query,
                           query_name=query_name)


@app.route('/payment_types', defaults={'length': 30, 'page': 1})
@app.route('/payment_types/<length>/<page>')
def payment_types(length, page):
    access = check_access(8)
    if access is not True:
        return access
    payments = [p.payment_id for p in BankStatement.query
    .order_by(BankStatement.date.desc()).order_by(BankStatement.order_id.asc()).all()]
    n, data = make_pages(length, payments, page)
    statement = statement_info(data)
    p_types = set(p['payment_type'] for p in statement)
    return render_template('participants_and_payment/payment_types.html', statement=statement, pages=n, page=page,
                           length=length, link='payment_types', p_types=p_types)


@app.route('/download_payments/<p_type>')
def download_payments(p_type):
    if p_type == 'all':
        payments = db.session.query(BankStatement).order_by(BankStatement.date).all()
    else:
        payments = db.session.query(BankStatement) \
            .join(PaymentTypes, BankStatement.payment_id == PaymentTypes.payment_id) \
            .filter(PaymentTypes.payment_type == p_type).order_by(BankStatement.date).all()
    statement = [{'ID': p.payment_id, 'Дата': datetime.datetime.strftime(p.date, '%d.%m.%Y'),
                  'Номер платежного поручения': p.order_id, 'Дебит': p.debit, 'Кредит': p.credit,
                  'Плательщик': p.organisation, 'ИНН': p.tin, 'БИК': p.bic, 'Банк отправителя': p.bank_name,
                  'Номер счета': p.account, 'Назначение платежа': p.payment_comment,
                  'Альтернативная оплата': p.alternative, 'Комментарий': p.alternative_comment} for p in payments]

    df = pd.DataFrame(data=statement)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/statement_' + str(curr_year) + '.xlsx') as writer:
        df.to_excel(writer, sheet_name=p_type)
    return send_file('static/files/generated_files/statement_' + str(curr_year) + '.xlsx', as_attachment=True)


@app.route('/set_payment_types', methods=['POST'])
def set_payment_types():
    all_payments = PaymentTypes.query.all()
    payment_ids = [p.payment_id for p in all_payments]
    for payment in [p.payment_id for p in BankStatement.query.all()]:
        if 'payment_type/' + str(payment) in request.form.keys():
            p_type = request.form['payment_type/' + str(payment)]
            dict_type = {'payment_id': payment, 'payment_type': p_type}
            if PaymentTypes(dict_type['payment_id'], dict_type['payment_type']) not in all_payments:
                if dict_type['payment_id'] in payment_ids:
                    db.session.query(PaymentTypes).filter(PaymentTypes.payment_id == dict_type['payment_id']) \
                        .update({PaymentTypes.payment_type: dict_type['payment_type']})
                    db.session.commit()
                else:
                    p = PaymentTypes(dict_type['payment_id'], dict_type['payment_type'])
                    db.session.add(p)
                    db.session.commit()
    return redirect(url_for('.payment_types'))


@app.route('/id_payments', defaults={'mode': 'unset', 'length': 30, 'page': 1})
@app.route('/id_payments/<mode>/<length>/<page>')
def id_payments(mode, length, page):
    access = check_access(8)
    if access is not True:
        return access
    set_payments = [p.payment_id for p in PaymentRegistration.query.all()]
    if mode == 'unset':
        payments = [p.payment_id for p in BankStatement.query
        .join(PaymentTypes, BankStatement.payment_id == PaymentTypes.payment_id)
        .filter(PaymentTypes.payment_type == 'Чтения Вернадского')
        .order_by(BankStatement.date.desc()).order_by(BankStatement.order_id.asc()).all()
                    if p.payment_id not in set_payments]
        p_l = len(payments)
    elif mode == 'all':
        payments = [p.payment_id for p in BankStatement.query
        .join(PaymentTypes, BankStatement.payment_id == PaymentTypes.payment_id)
        .filter(PaymentTypes.payment_type == 'Чтения Вернадского')
        .order_by(BankStatement.date.desc()).order_by(BankStatement.order_id.asc()).all()]
        p_l = len(payments)
    else:
        mode = unquote(mode)
        all_payments = {p.payment_id: (str(p.payment_id) + str(p.debit) + p.organisation + p.payment_comment +
                                       str(p.order_id)).lower().replace(' ', '')
                        for p in BankStatement.query.all()}
        payments = []
        for k, v in all_payments.items():
            if mode.lower().replace(' ', '') in v:
                payments.append(k)
        p_l = len(payments)
    n, data = make_pages(length, payments, page)
    statement = statement_info(data)
    return render_template('participants_and_payment/id_payments.html', statement=statement, pages=n, page=page,
                           length=length, link='id_payments/' + mode, mode=mode, p_l=p_l, fee=fee, tour_fee=tour_fee)


@app.route('/set_payee/<payment_id>', defaults={'payee': None})
@app.route('/set_payee/<payment_id>/<payee>')
def set_payee(payment_id, payee):
    if isinstance(payee, str) and '%' in payee:
        payee = unquote(payee)
    access = check_access(8)
    if access is not True:
        return access
    order_ids = {p.payment_id: p.order_id for p in BankStatement.query.all()}
    payment = payment_info(payment_id)
    if payment['order_id'] is not None:
        del order_ids[payment['payment_id']]
        if payment['order_id'] in order_ids.values():
            for k, v in order_ids.items():
                if v == payment['order_id']:
                    if BankStatement.query.filter(BankStatement.date == payment['date']):
                        double = payment_info(k)
                    else:
                        double = None
                else:
                    double = None
        else:
            double = None
    else:
        double = None
    participant = {'type': None, 'participant': payee}
    if payee is not None:
        payee = payee.strip()
        try:
            payee = int(payee)
            if payee in [p.appl_id for p in ParticipantsApplied.query.all()]:
                participant = {'type': 'appl', 'participant': [application_2_tour(payee)]}
            elif payee in [w.work_id for w in Works.query.all()]:
                participant = {'type': 'work', 'works': work_info(payee, w_payment_info=True, appl_info=True)}
            else:
                participant = {'type': None, 'participant': payee}
        except ValueError:
            # if type(payee) == str and '%' in payee:
            #     payee = unquote(payee)
            parts = [u.participant_id for u in ParticipantsApplied.query.all()
                     if payee.lower() == u.last_name.lower()[:len(payee)]]
            parts.extend([u.participant_id for u in ParticipantsApplied.query.all()
                          if payee.lower() == u.first_name.lower()[:len(payee)]])
            parts.extend([u.participant_id for u in ParticipantsApplied.query.all()
                          if payee.lower() == u.patronymic_name.lower()[:len(payee)]])
            p = []
            for part in parts:
                appl = ParticipantsApplied.query.filter(ParticipantsApplied.participant_id == part).first().appl_id
                if appl not in [pa['id'] for pa in p]:
                    partic = application_2_tour(appl)
                    p.append(partic)
            participant = {'type': 'name', 'participant': p}

            works = [u.work_id for u in Works.query.all()
                     if payee.lower() == u.author_1_name.lower()[:len(payee)]]
            works.extend([u.work_id for u in Works.query.all()
                          if u.author_2_name and payee.lower() == u.author_2_name.lower()[:len(payee)]])
            works.extend([u.work_id for u in Works.query.all()
                          if u.author_3_name and payee.lower() == u.author_3_name.lower()[:len(payee)]])
            w = [work_info(wo, w_payment_info=True, appl_info=True) for wo in works]
            if participant['type']:
                participant['works'] = w
            else:
                participant = {'type': 'name', 'works': w}
            if not parts and not works:
                participant = {'type': None, 'participant': payee}
    else:
        participant = {'type': None, 'participant': None}
    p_types = set(p.payment_type for p in PaymentTypes.query.all())
    return render_template('participants_and_payment/set_payee.html', payment=payment, participant=participant,
                           query=payee, p_types=p_types, double=double)


@app.route('/application_payment/<payment_id>', methods=['GET'], defaults={'payee': None})
@app.route('/application_payment/<payment_id>/<payee>')
def application_payment(payment_id, payee):
    if payee is None:
        payee = request.values.get('payee', str)
    return redirect(url_for('.set_payee', payment_id=payment_id, payee=payee))


@app.route('/check_payees/<payment_id>/<appl>')
def check_payees(payment_id, appl):
    access = check_access(8)
    if access is not True:
        return access
    payment = payment_info(payment_id)
    application = application_2_tour(appl)
    return render_template('participants_and_payment/check_payees.html', payment=payment, appl=application)


@app.route('/set_payment/<payment_id>/<payee>', methods=['POST'])
def set_payment(payment_id, payee):
    if len(payee) == 6:
        for_work = True
        participant = int(payee)
        # if str(participant) not in request.form.keys():
        #     if participant in [p.participant for p in PaymentRegistration.query.all()]:
        #         if PaymentRegistration.query.filter(PaymentRegistration.participant == participant
        #                                             ).first().payment_id == int(payment_id):
        #             PaymentRegistration.query.filter(PaymentRegistration.participant == participant).delete()
        #             db.session.commit()
        # else:
        if len(request.form) > 0:
            data = request.form[str(participant)]
        else:
            data = 'off'
        if data == 'on':
            if participant in [p.participant for p in PaymentRegistration.query.all()]:
                if payment_id not in [p.payment_id for p in
                                      PaymentRegistration.query.filter(PaymentRegistration.participant ==
                                                                       participant).all()]:
                    payment = PaymentRegistration(payment_id, participant, for_work)
                    db.session.add(payment)
                    db.session.commit()
            else:
                payment = PaymentRegistration(payment_id, participant, for_work)
                db.session.add(payment)
                db.session.commit()
        else:
            PaymentRegistration.query.filter(PaymentRegistration.participant == participant).delete()
            db.session.commit()
    elif len(payee) == 5:
        for_work = False
        participants = [p.participant_id for p
                        in ParticipantsApplied.query.filter(ParticipantsApplied.appl_id == int(payee)).all()]
        for participant in participants:
            if str(participant) not in request.form.keys():
                if participant in [p.participant for p in PaymentRegistration.query.all()]:
                    if PaymentRegistration.query.filter(PaymentRegistration.participant == participant
                                                        ).first().payment_id == int(payment_id):
                        PaymentRegistration.query.filter(PaymentRegistration.participant == participant).delete()
                        db.session.commit()
            else:
                data = request.form[str(participant)]
                if data == 'on':
                    if participant in [p.participant for p in PaymentRegistration.query.all()]:
                        if payment_id not in [p.payment_id for p in
                                              PaymentRegistration.query.filter(PaymentRegistration.participant ==
                                                                               participant).all()]:
                            payment = PaymentRegistration(payment_id, participant, for_work)
                            db.session.add(payment)
                            db.session.commit()
                    else:
                        payment = PaymentRegistration(payment_id, participant, for_work)
                        db.session.add(payment)
                        db.session.commit()
                else:
                    if PaymentRegistration.query.filter(PaymentRegistration.participant == participant
                                                        ).first().payment_id == int(payment_id):
                        PaymentRegistration.query.filter(PaymentRegistration.participant == participant).delete()
                        db.session.commit()
    return redirect(url_for('.id_payments'))


@app.route('/confirm_delete/<specify>/<del_id>/<url>')
def confirm_delete(specify, del_id, url):
    back = request.referrer
    return render_template('confirm_delete.html', specify=specify, del_id=del_id, url=url, back=back)


@app.route('/delete_payment/<del_id>')
def delete_payment(del_id):
    BankStatement.query.filter(BankStatement.payment_id == del_id).delete()
    db.session.commit()
    return redirect(url_for('.id_payments'))


@app.route('/reset_payment_type/<payment_id>/<payment_type>')
def reset_payment_type(payment_id, payment_type):
    payment_id = int(payment_id)
    if payment_id in [p.payment_id for p in PaymentTypes.query.all()]:
        db.session.query(PaymentTypes).filter(PaymentTypes.payment_id == payment_id) \
            .update({PaymentTypes.payment_type: payment_type})
        db.session.commit()
    else:
        p = PaymentTypes(payment_id, payment_type)
        db.session.add(p)
        db.session.commit()
    return redirect(url_for('.id_payments'))


@app.route('/add_emails', defaults={'success': None})
@app.route('/add_emails/<success>')
def add_emails(success):
    access = check_access(8)
    if access is not True:
        return access
    return render_template('online_reports/add_emails.html', success=success)


@app.route('/save_emails', methods=['POST'])
def save_emails():
    data = request.files['file'].read().decode('mac_cyrillic').replace('\xa0', ' ')
    lines = data.split('\n')
    mail_data = []
    all_works = [w.work_id for w in Works.query.all()]
    for line in lines[1:]:
        if line != '':
            sta = {name.strip().strip('\r'): value.strip().strip('\r')
                   for name, value in zip(lines[0].split('\t'), line.split('\t'))}
            mail_data.append(sta)
    for work in mail_data:
        if work['work_id'] != '':
            work_id = int(work['work_id'])
            if work_id in all_works:
                indices = ['email' + str(i) for i in range(1, len(work))]
                mls = [work[ind] for ind in indices if work[ind] != '' and work[ind] is not None]
                work_mail = Works.query.filter(Works.work_id == work_id).first().email
                mls.append(work_mail)
                mails = set(mls)
                for mail in mails:
                    if mail in [m.email for m in Mails.query.all()]:
                        mail_id = Mails.query.filter(Mails.email == mail).first().mail_id
                    else:
                        m = Mails(mail)
                        db.session.add(m)
                        db.session.flush()
                        db.session.commit()
                        mail_id = m.mail_id
                    w_m = WorkMail(work_id, mail_id, False)
                    if work_id not in [w.work_id for w in WorkMail.query.all()]:
                        db.session.add(w_m)
                        db.session.commit()
                    elif mail_id not in [w.mail_id for w in WorkMail.query.filter(WorkMail.work_id == work_id).all()]:
                        db.session.add(w_m)
                        db.session.commit()
    success = True
    return redirect(url_for('.add_emails', success=success))


@app.route('/load_diplomas', defaults={'success': None})
@app.route('/load_diplomas/<success>')
def load_diplomas(success):
    access = check_access(8)
    if access is not True:
        return access
    return render_template('online_reports/load_diplomas.html', success=success)


@app.route('/save_diplomas', methods=['POST'])
def save_diplomas():
    if not os.path.isdir('static/files/uploaded_files'):
        os.mkdir('static/files/uploaded_files')
    if not os.path.isdir('static/files/uploaded_files/diplomas_' + str(curr_year)):
        os.mkdir('static/files/uploaded_files/diplomas_' + str(curr_year))
    files = request.files.getlist("diplomas")
    for file in files:
        file.save(os.path.join('static/files/uploaded_files/diplomas_' + str(curr_year), file.filename))
    success = True


#     return redirect(url_for('.load_diplomas', success=success))
#
#
# @app.route('/save_diplomas_check')
# def save_diplomas():
#     if not os.path.isdir('static/files/uploaded_files'):
#         os.mkdir('static/files/uploaded_files')
#     if not os.path.isdir('static/files/uploaded_files/diplomas_' + str(curr_year)):
#         os.mkdir('static/files/uploaded_files/diplomas_' + str(curr_year))
#     dirpath, dirnames, filenames = os.walk('static/files/uploaded_files/diplomas_' + str(curr_year))
#     for file in files:
#         file.save(os.path.join('static/files/uploaded_files/diplomas_' + str(curr_year), file.filename))
#     success = True
#     return redirect(url_for('.load_diplomas', success=success))


@app.route('/send_diplomas', defaults={'cat_id': 'first', 'wrong': None})
@app.route('/send_diplomas/<cat_id>/<wrong>')
def send_diplomas(cat_id, wrong):
    if cat_id == 'first':
        cat_id = min([c.cat_id for c in Categories.query.filter(Categories.year == curr_year).all()])
    c, cats = categories_info()
    for cat in cats:
        if cat['id'] == int(cat_id) or cat_id == 'all':
            cat['works'] = [work_info(work_id=w.work_id, mail_info=True, w_payment_info=True) for w in Works.query
            .join(WorkCategories, Works.work_id == WorkCategories.work_id).filter(WorkCategories.cat_id == cat['id'])
            .filter(Works.reported == 1).all()]
    if cat_id != 'all':
        cat_id = int(cat_id)
    return render_template('online_reports/send_diplomas.html', cats=cats, cat_id=cat_id, wrong=wrong)


@app.route('/send_left_diplomas', defaults={'cat_id': 'first', 'wrong': None})
@app.route('/send_left_diplomas/<cat_id>/<wrong>')
def send_left_diplomas(cat_id, wrong):
    if cat_id == 'first':
        cat_id = min([c.cat_id for c in Categories.query.filter(Categories.year == curr_year).all()])
    c, cats = categories_info()
    for cat in cats:
        if cat['id'] == int(cat_id) or cat_id == 'all':
            cat['works'] = [work_info(work_id=w.work_id, mail_info=True, w_payment_info=True) for w in Works.query
            .join(WorkCategories, Works.work_id == WorkCategories.work_id)
            .join(ParticipatedWorks, Works.work_id == ParticipatedWorks.work_id)
            .filter(WorkCategories.cat_id == cat['id']).all()]
    if cat_id != 'all':
        cat_id = int(cat_id)
    return render_template('online_reports/send_left_diplomas.html', cats=cats, cat_id=cat_id, wrong=wrong)


@app.route('/sending_diplomas/<send_type>/<w_c_id>')
def sending_diplomas(send_type, w_c_id):
    if send_type == 'cat':
        cat_id = int(w_c_id)
        works = [w.work_id for w in Works.query
        .join(WorkCategories, Works.work_id == WorkCategories.work_id).filter(WorkCategories.cat_id == cat_id)
        .filter(Works.reported == 1).all()]
    else:
        work_id = int(w_c_id)
        works = [w.work_id for w in Works.query.filter(Works.work_id == work_id)
        .filter(Works.reported == 1).all()]
        cat_id = WorkCategories.query.filter(WorkCategories.work_id == work_id).first().cat_id

    payed = [p.participant for p in PaymentRegistration.query.all()]
    payed.extend([w.work_id for w in WorksNoFee.query.all()])
    payed.extend([w.work_id for w in Discounts.query.filter(Discounts.payment == 0).all()])

    dir = 'static/files/uploaded_files/diplomas_online_' + str(curr_year) + '/'

    service = get_service()

    for w_id in works:
        # try:
        if w_id in payed:
            files = [f for f in os.listdir(dir) if os.path.isfile(os.path.join(dir, f)) if f[:6] == str(w_id)]
            if files:
                mails = [(m.mail_id, m.email) for m in Mails.query.join(WorkMail, Mails.mail_id == WorkMail.mail_id)
                .filter(WorkMail.work_id == w_id).filter(WorkMail.sent == 0).all()]
                if mails:
                    for mail_record in mails:
                        mail_id, recipient_email = mail_record
                        if recipient_email not in ('0', 0, ''):
                            # Формируем attachments_list
                            attachments_list = []
                            for f in files:
                                fi = os.path.join(dir, f)
                                with app.open_resource(fi) as file:
                                    file_data = file.read()
                                attachments_list.append({
                                    'filename': os.path.basename(fi),
                                    'data': file_data
                                })

                            # Формируем письмо
                            html_body = render_template('diplomas_mail.html', work_id=w_id)
                            subject = 'Наградные документы ' + str(w_id)
                            sender = 'info@vernadsky.info'

                            message = create_message_with_attachments(
                                sender=sender,
                                to=recipient_email,
                                subject=subject,
                                html_body=html_body,
                                attachments=attachments_list
                            )

                            # Отправляем
                            send_message(service, "me", message)

                            # Обновляем WorkMail.sent
                            db.session.query(WorkMail) \
                                .filter(WorkMail.work_id == w_id) \
                                .filter(WorkMail.mail_id == mail_id) \
                                .update({WorkMail.sent: True})
                            db.session.commit()
                # if mails:
                #     for a in mails:
                #         m = a[1]
                #         if m != 0 and m != '0' and m != '':
                #             attachments = []
                #             for f in files:
                #                 fi = dir + '/' + f
                #                 with app.open_resource(fi) as file:
                #                     attachments.append(Attachment(filename=os.path.basename(fi),
                #                                                   content_type=mimetypes.guess_type(fi)[0],
                #                                                   data=file.read()))
                #
                #             msg = Message(subject='Наградные документы ' + str(w_id),
                #                           html=render_template('diplomas_mail.html', work_id=w_id),
                #                           attachments=attachments,
                #                           sender=('Команда Конкурса им. В. И. Вернадского', 'team@vernadsky.info'),
                #                           recipients=[m])
                #             mail.send(msg)
                #
                #             db.session.query(WorkMail).filter(WorkMail.work_id == w_id).filter(
                #                 WorkMail.mail_id == a[0]) \
                #                 .update({WorkMail.sent: True})
                #             db.session.commit()
                            # if w_id in [w.work_id for w in Diplomas.query.all()]:
                            #     to_del = db.session.query(Diplomas).filter(Diplomas.work_id == w_id).first()
                            #     db.session.delete(to_del)
                            #     db.session.commit()
            # else:
            #     if w_id not in [w.work_id for w in Diplomas.query.all()]:
            #         a = Diplomas(w_id, False)
            #         db.session.add(a)
            #         db.session.commit()
            #     elif Diplomas.query.filter(Diplomas.work_id == w_id).first().diplomas:
            #         db.session.query(Diplomas).filter(Diplomas.work_id == w_id).update({Diplomas.diplomas: False})
            #         db.session.commit()
        # except Exception:
        #     return redirect(url_for('.send_diplomas', cat_id=cat_id, wrong=True))
    return redirect(url_for('.send_diplomas', cat_id=cat_id, wrong=False))


@app.route('/sending_left_diplomas/<send_type>/<w_c_id>')
def sending_left_diplomas(send_type, w_c_id):
    if send_type == 'cat':
        cat_id = int(w_c_id)
        works = [work_info(work_id=w.work_id, mail_info=True, w_payment_info=True) for w in Works.query
        .join(WorkCategories, Works.work_id == WorkCategories.work_id)
        .join(ParticipatedWorks, Works.work_id == ParticipatedWorks.work_id)
        .filter(WorkCategories.cat_id == cat_id).all()]
    else:
        work_id = int(w_c_id)
        works = [w.work_id for w in Works.query.filter(Works.work_id == work_id).all()]
        cat_id = WorkCategories.query.filter(WorkCategories.work_id == work_id).first().cat_id

    dir = 'static/files/uploaded_files/diplomas_' + str(curr_year) + '/'

    service = get_service()

    for w_id in works:
        # try:
        files = [f for f in os.listdir(dir) if os.path.isfile(os.path.join(dir, f)) if f[:6] == str(w_id)]
        if files:
            mails = [m.email for m in Mails.query.join(WorkMail, Mails.mail_id == WorkMail.mail_id)
            .filter(WorkMail.work_id == w_id).filter(WorkMail.sent == 0).all()]
            if mails:
                for m in mails:
                    if m not in (0, '0', ''):
                        attachments_list = []
                        for f in files:
                            fi = os.path.join(dir, f)
                            with app.open_resource(fi) as file:
                                file_data = file.read()
                            attachments_list.append({
                                'filename': os.path.basename(fi),
                                'data': file_data
                            })

                        html_body = render_template('diplomas_mail.html', work_id=w_id)
                        subject = 'Наградные документы ' + str(w_id)
                        sender = 'info@vernadsky.info'

                        message = create_message_with_attachments(
                            sender=sender,
                            to=m,
                            subject=subject,
                            html_body=html_body,
                            attachments=attachments_list,
                            bcc='info@vernadsky.info',
                            reply_to='info@vernadsky.info'
                        )
                        send_message(service, "me", message)

                        # Обновляем WorkMail
                        a = [x.mail_id for x in WorkMail.query.filter(WorkMail.work_id == w_id).all()]
                        for b in a:
                            db.session.query(WorkMail).filter(WorkMail.work_id == w_id) \
                                .filter(WorkMail.mail_id == b) \
                                .update({WorkMail.sent: True})
                            db.session.commit()

            # if mails:
            #     for m in mails:
            #         if m != 0 and m != '0' and m != '':
            #             attachments = []
            #             for f in files:
            #                 fi = dir + '/' + f
            #                 with app.open_resource(fi) as file:
            #                     attachments.append(Attachment(filename=os.path.basename(fi),
            #                                                   content_type=mimetypes.guess_type(fi)[0],
            #                                                   data=file.read()))
            #
            #             msg = Message(subject='Наградные документы ' + str(w_id),
            #                           html=render_template('diplomas_mail.html', work_id=w_id),
            #                           attachments=attachments,
            #                           sender=('Команда Конкурса им. В. И. Вернадского', 'team@vernadsky.info'),
            #                           recipients=[m],
            #                           bcc=['info@vernadsky.info'],
            #                           reply_to='info@vernadsky.info')
            #             mail.send(msg)
            #             a = [a.mail_id for a in WorkMail.query.filter(WorkMail.work_id == w_id).all()]
            #             for b in a:
            #                 db.session.query(WorkMail).filter(WorkMail.work_id == w_id).filter(
            #                     WorkMail.mail_id == b) \
            #                     .update({WorkMail.sent: True})
            #                 db.session.commit()
                        # if w_id in [w.work_id for w in Diplomas.query.all()]:
                        #     to_del = db.session.query(Diplomas).filter(Diplomas.work_id == w_id).first()
                        #     db.session.delete(to_del)
                        #     db.session.commit()
            # else:
            #     if w_id not in [w.work_id for w in Diplomas.query.all()]:
            #         a = Diplomas(w_id, False)
            #         db.session.add(a)
            #         db.session.commit()
            #     elif Diplomas.query.filter(Diplomas.work_id == w_id).first().diplomas:
            #         db.session.query(Diplomas).filter(Diplomas.work_id == w_id).update({Diplomas.diplomas: False})
            #         db.session.commit()
        # except Exception:
        #     return redirect(url_for('.send_diplomas', cat_id=cat_id, wrong=True))
    return redirect(url_for('.send_left_diplomas', cat_id=cat_id, wrong=False))


@app.route('/volunteer_tasks/', defaults={'task_id': ''})
@app.route('/volunteer_tasks/<task_id>')
def volunteer_tasks(task_id):
    access = check_access(8)
    if access is not True:
        return access
    tasks = [{'id': t.task_id,
              'task_name': t.task_name,
              'location': t.location,
              'address': t.address,
              'description': t.description,
              'task_date': days[t.start_time.strftime('%w')] + ', ' + t.start_time.strftime('%d') + ' ' +
                           months_full[t.start_time.strftime('%m')],
              'start_time': datetime.datetime.strftime(t.start_time, '%H:%M'),
              'end_time': datetime.datetime.strftime(t.end_time, '%H:%M'),
              'volunteers_required': t.volunteers_required}
             for t in VolunteerTasks.query.filter(VolunteerTasks.year == curr_year)
             .order_by(VolunteerTasks.start_time).all()]
    if task_id != '':
        t = VolunteerTasks.query.filter(VolunteerTasks.task_id == task_id).first()
        to_edit = {'id': t.task_id,
                   'task_name': t.task_name,
                   'location': t.location,
                   'address': t.address,
                   'description': t.description,
                   'task_date': t.start_time.strftime('%Y-%m-%d'),
                   'start_time': datetime.datetime.strftime(t.start_time, '%H:%M'),
                   'end_time': datetime.datetime.strftime(t.end_time, '%H:%M'),
                   'volunteers_required': t.volunteers_required}
    else:
        to_edit = {}
    return render_template('application management/volunteer_tasks.html', tasks=tasks, to_edit=to_edit)


@app.route('/save_volunteer_task', methods=['POST'])
def save_volunteer_task():
    if 'task_id' in request.form.keys() and request.form['task_id']:
        task_id = request.form['task_id']
    else:
        task_id = None
    task_name = request.form['task_name']
    location = request.form['location']
    address = request.form['address']
    description = request.form['description']
    task_date = request.form['task_date']
    start_time = request.form['start_time']
    end_time = request.form['end_time']
    volunteers_required = request.form['volunteers_required']
    start = datetime.datetime.strptime(task_date + ' ' + start_time, '%Y-%m-%d %H:%M')
    end = datetime.datetime.strptime(task_date + ' ' + end_time, '%Y-%m-%d %H:%M')
    if task_id:
        db.session.query(VolunteerTasks).filter(VolunteerTasks.task_id == int(task_id)) \
            .update({VolunteerTasks.task_name: task_name,
                     VolunteerTasks.location: location,
                     VolunteerTasks.address: address,
                     VolunteerTasks.description: description,
                     VolunteerTasks.start_time: start,
                     VolunteerTasks.end_time: end,
                     VolunteerTasks.volunteers_required: volunteers_required})
    else:
        task = VolunteerTasks(task_name, location, address, description, start, end, volunteers_required, curr_year)
        db.session.add(task)
    db.session.commit()
    return redirect(url_for('.volunteer_tasks'))


@app.route('/school_classes', defaults={'class_id': ''})
@app.route('/school_classes/<class_id>')
def school_classes(class_id):
    access = check_access(8)
    if access is not True:
        return access
    sch_classes = [{'class_id': c.class_id,
                    'class_name': c.class_name,
                    'school': c.school} for c in SchoolClasses.query.filter(SchoolClasses.year == curr_year)
                   .filter(SchoolClasses.class_type == 'class').all()]
    s_cl = sorted(sch_classes, key=lambda x: x['class_name'])
    sch_classes = sorted(s_cl, key=lambda x: x['school'])
    if class_id != '':
        c = SchoolClasses.query.filter(SchoolClasses.class_id == int(class_id)).first()
        to_edit = {'class_id': c.class_id, 'class_name': c.class_name, 'school': c.school}
    else:
        to_edit = {}
    return render_template('application management/school_classes.html', sch_classes=sch_classes, to_edit=to_edit)


@app.route('/add_classes', methods=['POST'])
def add_classes():
    if 'class_id' in request.form.keys() and request.form['class_id']:
        class_id = request.form['class_id']
    else:
        class_id = None
    school = request.form['school']
    class_name = request.form['class_name']
    if class_id:
        db.session.query(SchoolClasses).filter(SchoolClasses.class_id == int(class_id)) \
            .update({SchoolClasses.school: school,
                     SchoolClasses.class_name: class_name,
                     SchoolClasses.class_type: 'class'})
    else:
        school_class = SchoolClasses(school, class_name, curr_year, 'class')
        db.session.add(school_class)
    db.session.commit()
    return redirect(url_for('.school_classes'))


@app.route('/my_volunteer_tasks')
def my_volunteer_tasks():
    access = check_access(2)
    if access is not True:
        return access
    user_id = int(session['user_id'])
    profile = Profile.query.filter(Profile.user_id == user_id).first()
    involved = profile.involved
    if user_id in [u.user_id for u in StudentClass.query.filter(StudentClass.year == curr_year).all()]:
        class_id = StudentClass.query.filter(StudentClass.year == curr_year) \
            .filter(StudentClass.user_id == user_id).first().class_id
    else:
        class_id = None

    if involved == '1553':
        s_cl = [{'class_id': c.class_id, 'class_name': c.class_name, 'school': c.school}
                for c in SchoolClasses.query.filter(SchoolClasses.year == curr_year)
                .filter(SchoolClasses.school == '1553').filter(SchoolClasses.class_type == 'class').all()]
        sch_class = sorted(s_cl, key=lambda x: x['class_name'])

    elif involved == 'MSU_School':
        s_cl = [{'class_id': c.class_id, 'class_name': c.class_name, 'school': c.school}
                for c in SchoolClasses.query.filter(SchoolClasses.year == curr_year)
                .filter(SchoolClasses.school == 'MSU_School').filter(SchoolClasses.class_type == 'class').all()]
        sch_class = sorted(s_cl, key=lambda x: x['class_name'])
    else:
        sch_class = None

    my_tasks_db = db.session.query(VolunteerAssignment).filter(
        VolunteerAssignment.user_id == int(session['user_id'])).all()
    permitted = []
    prohibited = []
    pending = []
    unresolved = []
    for task in my_tasks_db:
        if task.permitted == 'yes':
            permitted.append(task.task_id)
        elif task.permitted == 'no':
            prohibited.append(task.task_id)
        elif task.permitted == 'pending':
            pending.append(task.task_id)
        else:
            unresolved.append(task.task_id)
    tasks = [{'id': t.task_id,
              'task_name': t.task_name,
              'location': t.location,
              'address': t.address,
              'description': t.description,
              'real_date': t.start_time,
              'task_date': days[t.start_time.strftime('%w')] + ', ' + t.start_time.strftime('%d') + ' ' +
                           months_full[t.start_time.strftime('%m')],
              'start_time': datetime.datetime.strftime(t.start_time, '%H:%M'),
              'end_time': datetime.datetime.strftime(t.end_time, '%H:%M'),
              'volunteers_required': t.volunteers_required}
             for t in VolunteerTasks.query.filter(VolunteerTasks.year == curr_year)
             .order_by(VolunteerTasks.start_time).all()]
    task_days = [task['real_date'] for task in tasks]
    day_tasks = {day.date(): [] for day in sorted(list(set(task_days)))}
    for task in tasks:
        if task['id'] in permitted:
            task['applied'] = True
            task['permitted'] = 'yes'
        elif task['id'] in prohibited:
            task['applied'] = True
            task['permitted'] = 'no'
        elif task['id'] in pending:
            task['applied'] = True
            task['permitted'] = 'pending'
        elif task['id'] in unresolved:
            task['applied'] = True
            task['permitted'] = None
        else:
            task['applied'] = False
        day_tasks[task['real_date'].date()].append(task)
    return render_template('application management/my_volunteer_tasks.html', involved=involved, sch_class=sch_class,
                           class_id=class_id, tasks=day_tasks)


@app.route('/set_class', methods=['POST'])
def set_class():
    class_id = int(request.form['class_id'])
    class_name = SchoolClasses.query.filter(SchoolClasses.class_id == class_id).first().class_name
    n = len(class_name)
    while n >= 0:
        try:
            class_num = int(class_name[0:n])
            break
        except BaseException:
            n -= 1
    user_id = int(session['user_id'])
    if user_id in [u.user_id for u in StudentClass.query.filter(StudentClass.year == curr_year).all()]:
        db.session.query(StudentClass).filter(StudentClass.year == curr_year).filter(StudentClass.user_id == user_id) \
            .update({StudentClass.class_id: class_id})
    else:
        u_c = StudentClass(user_id, class_id, curr_year)
        db.session.add(u_c)
    db.session.commit()
    db.session.query(Profile).filter(Profile.user_id == user_id).update({Profile.grade: class_num})
    db.session.commit()
    return redirect(url_for('.my_volunteer_tasks'))


@app.route('/pick_task/<task_id>/<action>')
def pick_task(task_id, action):
    task_id = int(task_id)
    user_id = int(session['user_id'])
    assignments = [(a.user_id, a.task_id) for a in VolunteerAssignment.query.all()]
    to_assign = (user_id, task_id)
    if to_assign not in assignments and action == 'pick':
        ass = VolunteerAssignment(user_id, task_id, None, None)
        db.session.add(ass)
        db.session.commit()
    elif to_assign in assignments and action == 'delete':
        db.session.query(VolunteerAssignment).filter(VolunteerAssignment.user_id == user_id) \
            .filter(VolunteerAssignment.task_id == task_id).delete()
        db.session.commit()
    else:
        pass
    if user_id in [u.user_id for u in StudentClass.query.filter(StudentClass.year == curr_year).all()]:
        pass
    else:
        profile = Profile.query.filter(Profile.user_id == user_id).first()
        involved = profile.involved
        class_num = profile.grade
        if involved == '1553':
            try:
                class_db = SchoolClasses.query.filter(SchoolClasses.year == curr_year).filter(SchoolClasses.school == '1553')\
                    .filter(SchoolClasses.class_type == 'class').filter(SchoolClasses.class_name == str(class_num)).first()
                u_c = StudentClass(user_id, class_db.class_id, curr_year)
                db.session.add(u_c)
                db.session.commit()
            except Exception:
                pass
        else:
            pass
    return redirect(url_for('.my_volunteer_tasks'))


@app.route('/volunteer_applications/', defaults={'view': 'all'})
@app.route('/volunteer_applications/<view>')
def volunteer_applications(view):
    access = check_access(7)
    if access is not True:
        return access
    tasks = [{'id': t.task_id,
              'task_name': t.task_name,
              'location': t.location,
              'address': t.address,
              'description': t.description,
              'task_date': days[t.start_time.strftime('%w')] + ', ' + t.start_time.strftime('%d') + ' ' +
                           months_full[t.start_time.strftime('%m')],
              'start_time': datetime.datetime.strftime(t.start_time, '%H:%M'),
              'end_time': datetime.datetime.strftime(t.end_time, '%H:%M'),
              'volunteers_required': t.volunteers_required}
             for t in VolunteerTasks.query.filter(VolunteerTasks.year == curr_year)
             .order_by(VolunteerTasks.start_time).all()]
    if (session['tutor'] is True and session['type'] not in ['admin', 'org', 'manager']) or view == 'tutor':
        volunteers = set(v.user_id for v in VolunteerAssignment.query
                         .join(VolunteerTasks, VolunteerAssignment.task_id == VolunteerTasks.task_id)
                         .join(StudentClass, VolunteerAssignment.user_id == StudentClass.user_id)
                         .filter(VolunteerTasks.year == curr_year)
                         .filter(StudentClass.class_id == int(session['class_id'])).all())
    else:
        volunteers = set(v.user_id for v in VolunteerAssignment.query
                         .join(VolunteerTasks, VolunteerAssignment.task_id == VolunteerTasks.task_id)
                         .filter(VolunteerTasks.year == curr_year).all())
    sch_classes = {c.class_id: {'school': c.school, 'class_name': c.class_name}
                   for c in SchoolClasses.query.filter(SchoolClasses.year == curr_year)
                   .filter(SchoolClasses.class_type == 'class').all()}
    school_info = {u.user_id: sch_classes[u.class_id] for u in StudentClass.query.all() if u.user_id in volunteers}
    for u in volunteers:
        if u not in school_info.keys():
            school_info[u] = {'school': '', 'class_name': ''}
    user_info = {u.user_id: {'user_id': u.user_id,
                             'name': u.last_name + ' ' + u.first_name,
                             'school': school_info[u.user_id]['school'],
                             'class_name': school_info[u.user_id]['class_name']}
                 for u in Users.query.all() if u.user_id in volunteers}
    if (session['tutor'] is True and session['type'] not in ['admin', 'org', 'manager']) or view == 'tutor':
        t_list = {t['id']: [{'user_id': u.user_id, 'permitted': u.permitted,
                             'permitter': u.permitter_id} for u in VolunteerAssignment.query
                            .join(StudentClass, VolunteerAssignment.user_id == StudentClass.user_id)
                            .filter(VolunteerAssignment.task_id == t['id'])
                            .filter(StudentClass.class_id == int(session['class_id'])).all()] for t in tasks}
    else:
        t_list = {t['id']: [{'user_id': u.user_id, 'permitted': u.permitted,
                             'permitter': u.permitter_id} for u in VolunteerAssignment.query
                            .filter(VolunteerAssignment.task_id == t['id']).all()] for t in tasks}
    for task in tasks:
        vols = []
        for u in t_list[task['id']]:
            u.update(user_info[u['user_id']])
            if u['permitter'] is not None:
                permitter = get_user_info(u['permitter'])
                u['permitter'] = permitter
            vols.append(u)
        task['volunteers_list'] = sorted(vols, key=lambda x: x['name'])
        task['vols_got'] = len(task['volunteers_list'])
    v_t = []
    for t in tasks:
        v_t.extend([u['user_id'] for u in t['volunteers_list']])
    vol_with_tasks = len(set(v_t))
    return render_template('application management/volunteer_applications.html', tasks=tasks, year=curr_year,
                           vol_with_tasks=vol_with_tasks, view=view)


@app.route('/download_volunteer_applications')
def download_volunteer_applications():
    access = check_access(7)
    if access is not True:
        return access

    dnl = []

    tasks = [{'id': t.task_id,
              'location': t.location + ' (' + t.address + ')',
              'task_date': t.start_time.strftime('%w'),
              'start_time': t.start_time,
              'end_time': t.end_time}
             for t in VolunteerTasks.query.filter(VolunteerTasks.year == curr_year)
             .order_by(VolunteerTasks.start_time).all()]
    volunteers = set(v.user_id for v in VolunteerAssignment.query
                     .join(VolunteerTasks, VolunteerAssignment.task_id == VolunteerTasks.task_id)
                     .filter(VolunteerTasks.year == curr_year).all())
    sch_classes = {c.class_id: {'school': c.school, 'class_name': c.class_name}
                   for c in SchoolClasses.query.filter(SchoolClasses.year == curr_year)
                   .filter(SchoolClasses.class_type == 'class').all()}
    school_info = {u.user_id: sch_classes[u.class_id] for u in StudentClass.query.all() if u.user_id in volunteers}
    for u in volunteers:
        if u not in school_info.keys():
            school_info[u] = {'school': '', 'class_name': ''}
    user_info = {u.user_id: {'user_id': u.user_id,
                             'name': u.last_name + ' ' + u.first_name + ' ' + u.patronymic,
                             'school': school_info[u.user_id]['school'],
                             'class_name': school_info[u.user_id]['class_name']}
                 for u in Users.query.all() if u.user_id in volunteers}
    t_list = {t['id']: [{'user_id': u.user_id, 'permitted': u.permitted,
                         'permitter': u.permitter_id} for u in VolunteerAssignment.query
                        .filter(VolunteerAssignment.task_id == t['id']).all()] for t in tasks}
    for task in tasks:
        vols = []
        for u in t_list[task['id']]:
            u.update(user_info[u['user_id']])
            vols.append(u)
        task['volunteers_list'] = sorted(vols, key=lambda x: x['name'])
        for v in task['volunteers_list']:
            if v['school'] == 'MSU_School':
                if v['class_name'][:1] == '10':
                    lesson_time = MSU_lessons_10
                else:
                    lesson_time = MSU_lessons

                if task['task_date'] == 5:
                    s = task['start_time'] - datetime.timedelta(hours=1, minutes=30)
                    e = task['end_time'] + datetime.timedelta(hours=1, minutes=30)
                else:
                    s = task['start_time'] - datetime.timedelta(hours=1)
                    e = task['end_time'] + datetime.timedelta(hours=1)
                first_l = 0
                last_l = 0
                d = s.date() - lesson_time[1]['start'].date()
                for k, q in lesson_time.items():
                    if q['end'] + d <= s or k == 1:
                        first_l = k
                    if q['start'] + d < e:
                        last_l = k

                line = {'day': days[task['task_date']], 'permitted': v['permitted'],
                        'name': v['name'] + ' (' + v['class_name'] + ')',
                        'exit': datetime.datetime.strftime(s, '%d.%m.%Y %H:%M'),
                        'return': datetime.datetime.strftime(e, '%d.%m.%Y %H:%M'),
                        'lessons': str(first_l) + ' - ' + str(last_l), 'location': task['location']}
                dnl.append(line)

    df = pd.DataFrame(data=dnl)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/volunteer_tasks.xlsx') as writer:
        df.to_excel(writer, sheet_name='Задачи волонтеров')
    return send_file('static/files/generated_files/volunteer_tasks.xlsx', as_attachment=True)


@app.route('/approve_volunteer/<task_id>/<user_id>/<approval>/<view>')
def approve_volunteer(task_id, user_id, approval, view):
    task_id = int(task_id)
    user_id = int(user_id)
    permitter_id = int(session['user_id'])
    try:
        VolunteerAssignment.query.filter(VolunteerAssignment.user_id == user_id) \
            .filter(VolunteerAssignment.task_id == task_id).update({VolunteerAssignment.permitted: approval,
                                                                    VolunteerAssignment.permitter_id: permitter_id})
        db.session.commit()
    except BaseException:
        a = VolunteerAssignment(user_id, task_id, approval, permitter_id)
        db.session.add(a)
        db.session.commit()
    return redirect(url_for('.volunteer_applications', view=view))


@app.route('/download_team_applicants')
def download_team_applicants():
    prof_info = {p.user_id: {'user_id': p.user_id, 'occupation': p.occupation, 'place_of_work': p.place_of_work,
                             'involved': p.involved, 'grade': p.grade, 'year': p.year, 'vk': 'vk.com/' + p.vk,
                             'tg': 'https://t.me/' + p.telegram, 'vernadsky_username': p.vernadsky_username}
                 for p in Profile.query.join(Application, Profile.user_id == Application.user_id)
                 .filter(Application.year == curr_year).all()}
    cats = {c.cat_id: c.short_name for c in Categories.query.filter(Categories.year == curr_year).all()}
    appl = {}
    for a in Application.query.filter(Application.year == curr_year):
        appl[a.user_id] = {'role': a.role, 'any_category': a.any_category}
        if a.category_1 != 'None':
            appl[a.user_id]['cat_1'] = cats[a.category_1]
        else:
            appl[a.user_id]['cat_1'] = ''
        if a.category_2 != 'None':
            appl[a.user_id]['cat_2'] = cats[a.category_2]
        else:
            appl[a.user_id]['cat_2'] = ''
        if a.category_3 != 'None':
            appl[a.user_id]['cat_3'] = cats[a.category_3]
        else:
            appl[a.user_id]['cat_3'] = ''
    sch_class = {s.class_id: s.class_name for s in SchoolClasses.query.filter(SchoolClasses.year == curr_year)}
    user_class = {s.user_id: sch_class[s.class_id]
                  for s in StudentClass.query.filter(StudentClass.year == curr_year).all()}
    user_class.update({u: '' for u in prof_info.keys() if u not in user_class})
    applicants = [{'user_id': u.user_id, 'last_name': u.last_name, 'first_name': u.first_name,
                   'patronymic': u.patronymic, 'email': u.email,
                   'tel': u.tel, 'occupation': prof_info[u.user_id]['occupation'],
                   'place_of_work': prof_info[u.user_id]['place_of_work'], 'involved': prof_info[u.user_id]['involved'],
                   'grade': prof_info[u.user_id]['grade'], 'year': prof_info[u.user_id]['year'],
                   'class_name': user_class[u.user_id],
                   'vk': 'vk.com/' + prof_info[u.user_id]['vk'],
                   'tg': 'https://t.me/' + prof_info[u.user_id]['tg'],
                   'vernadsky_username': prof_info[u.user_id]['vernadsky_username'],
                   'role': appl[u.user_id]['role'], 'cat_1': appl[u.user_id]['cat_1'],
                   'cat_2': appl[u.user_id]['cat_2'], 'cat_3': appl[u.user_id]['cat_3'],
                   'any_category': appl[u.user_id]['any_category']} for u in
                  Users.query.join(Application, Users.user_id == Application.user_id)
                  .filter(Application.year == curr_year).all()]

    df = pd.DataFrame(data=applicants)
    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    with pd.ExcelWriter('static/files/generated_files/team_applicants.xlsx') as writer:
        df.to_excel(writer, sheet_name='Заявки ' + str(curr_year) + ' года')
    return send_file('static/files/generated_files/team_applicants.xlsx', as_attachment=True)


@app.route('/upload_school_schedule')
def upload_school_schedule():
    return render_template('application management/upload_school_schedule.html')


@app.route('/save_schedule', methods=['POST'])
def save_schedule():
    data = request.files['file'].read().decode('mac_cyrillic').replace('\r', '')
    lines = data.split('\n')
    study_groups = {g.class_name: g.class_id for g in SchoolClasses.query.filter(SchoolClasses.year == curr_year).all()}
    schedule = []
    for line in lines[2:]:
        if line != '':
            sta = {name: value for name, value in zip(lines[0].split('\t'), line.split('\t'))}
            if sta['group'] in study_groups:
                group_id = study_groups[sta['group']]
            else:
                if 'Английский язык' in sta['group']:
                    group_type = 'eng'
                elif 'язык' in sta['group']:
                    group_type = '2l'
                else:
                    group_type = 'profile'
                cl = SchoolClasses('MSU_School', sta['group'], curr_year, group_type)
                db.session.add(cl)
                db.session.commit()
                db.session.flush()
                group_id = cl.class_id
                study_groups[cl.class_name] = group_id
            less = LessonSchedule(int(sta['day']), int(sta['lesson']), sta['name'], curr_year)
            db.session.add(less)
            db.session.commit()
            db.session.flush()
            lesson_id = less.lesson_id
            less_gr = LessonGroup(lesson_id, group_id)
            db.session.add(less_gr)
            db.session.commit()
            schedule.append(sta)
    return redirect(url_for('.upload_school_schedule'))


@app.route('/get_runner')
def get_runner():
    groups = [{'type': p.class_type, 'class_id': p.class_id, 'class_name': p.class_name}
              for p in SchoolClasses.query.filter(SchoolClasses.year == curr_year)
              .filter(SchoolClasses.school == 'MSU_School').all()]
    profiles = []
    eng = []
    l2 = []
    for g in groups:
        if g['type'] == 'profile' or g['class_name'][0] == '8' or g['class_name'][0] == '9':
            profiles.append({'class_id': g['class_id'], 'class_name': g['class_name']})
        elif g['type'] == 'eng':
            eng.append({'class_id': g['class_id'], 'class_name': g['class_name']})
        elif g['type'] == '2l':
            l2.append({'class_id': g['class_id'], 'class_name': g['class_name']})
    pr = sorted(profiles, key=lambda x: x['class_name'])
    e = sorted(eng, key=lambda x: x['class_name'])
    l2s = sorted(l2, key=lambda x: x['class_name'])

    user_groups = [g.class_id
                   for g in StudentGroup.query.filter(StudentGroup.user_id == int(session['user_id'])).all()]
    return render_template('application management/get_runner.html', profiles=pr, eng=e, l2=l2s,
                           user_groups=user_groups)


@app.route('/my_study_groups', methods=['POST'])
def my_study_groups():
    profile = int(request.form['profile'])
    eng = int(request.form['eng'])
    l2 = int(request.form['l2'])
    user_id = int(session['user_id'])
    st_pr = StudentGroup(user_id, profile, 'profile')
    st_en = StudentGroup(user_id, eng, 'eng')
    st_l2 = StudentGroup(user_id, l2, '2l')

    if 'profile' in [g.group_type for g in StudentGroup.query.all()]:
        if user_id not in [u.user_id for u in StudentGroup.query.filter(StudentGroup.group_type == 'profile').all()]:
            db.session.add(st_pr)
            db.session.commit()
        else:
            db.session.query(StudentGroup).filter(StudentGroup.user_id == user_id) \
                .filter(StudentGroup.group_type == 'profile').update({StudentGroup.class_id: profile})
            db.session.commit()
    else:
        db.session.add(st_pr)
        db.session.commit()

    if 'eng' in [g.group_type for g in StudentGroup.query.all()]:
        if user_id not in [u.user_id for u in StudentGroup.query.filter(StudentGroup.group_type == 'eng').all()]:
            db.session.add(st_en)
            db.session.commit()
        else:
            db.session.query(StudentGroup).filter(StudentGroup.user_id == user_id) \
                .filter(StudentGroup.group_type == 'eng').update({StudentGroup.class_id: eng})
            db.session.commit()
    else:
        db.session.add(st_en)
        db.session.commit()

    if '2l' in [g.group_type for g in StudentGroup.query.all()]:
        if user_id not in [u.user_id for u in StudentGroup.query.filter(StudentGroup.group_type == '2l').all()]:
            db.session.add(st_l2)
            db.session.commit()
        else:
            db.session.query(StudentGroup).filter(StudentGroup.user_id == user_id) \
                .filter(StudentGroup.group_type == '2l').update({StudentGroup.class_id: l2})
            db.session.commit()
    else:
        db.session.add(st_l2)
        db.session.commit()
    return redirect(url_for('.get_runner'))


@app.route('/download_runner')
def download_runner():
    u = db.session.query(Users).filter(Users.user_id == int(session['user_id'])).first()
    user_id = int(session['user_id'])

    name = u.last_name + ' ' + u.first_name + ' ' + u.patronymic

    gr = db.session.query(StudentGroup).filter(StudentGroup.user_id == user_id)

    prof = gr.filter(StudentGroup.group_type == 'profile').first().class_id
    eng = gr.filter(StudentGroup.group_type == 'eng').first().class_id
    l2 = gr.filter(StudentGroup.group_type == '2l').first().class_id

    my_lessons = [le.lesson_id for le in LessonGroup.query.filter(LessonGroup.class_id == prof).all()]
    my_lessons.extend([le.lesson_id for le in LessonGroup.query.filter(LessonGroup.class_id == eng).all()])
    my_lessons.extend([le.lesson_id for le in LessonGroup.query.filter(LessonGroup.class_id == l2).all()])

    cl = SchoolClasses.query.filter(SchoolClasses.class_id == prof).first().class_name

    if cl[:1] == '10':
        lesson_time = MSU_lessons_10
    else:
        lesson_time = MSU_lessons

    tasks = [{'id': t.task_id,
              'task_name': t.task_name,
              'location': t.location,
              'address': t.address,
              'description': t.description,
              'real_date': t.start_time,
              'task_date': days_full[t.start_time.strftime('%w')] + ', ' + t.start_time.strftime('%d') + ' ' +
                           months_full[t.start_time.strftime('%m')],
              'start_time': datetime.datetime.strftime(t.start_time, '%H:%M'),
              'end_time': datetime.datetime.strftime(t.end_time, '%H:%M'),
              'start': t.start_time,
              'end': t.end_time,
              'volunteers_required': t.volunteers_required}
             for t in VolunteerTasks.query.filter(VolunteerTasks.year == curr_year)
             .order_by(VolunteerTasks.start_time).all()]

    ta = []

    for task in tasks:
        if 'secretary' in session.keys():
            if task['real_date'].strftime('%w') in ['2', '5']:
                if task['description'] == 'secretary':
                    secr_secr = True
                else:
                    secr_secr = False
            elif task['real_date'].strftime('%w') == '3':
                secr_secr = True
            else:
                if task['description'] == 'secretary':
                    secr_secr = False
                else:
                    secr_secr = True
        else:
            if task['description'] == 'secretary':
                secr_secr = False
            else:
                secr_secr = True

        if secr_secr is True:
            ta.append(task)

    for task in ta:
        if datetime.datetime.strftime(task['real_date'], '%w') == 5:
            s = task['start'] - datetime.timedelta(hours=1, minutes=30)
            e = task['end'] + datetime.timedelta(hours=1, minutes=30)
        else:
            s = task['start'] - datetime.timedelta(hours=1)
            e = task['end'] + datetime.timedelta(hours=1)
        task['first_l'] = 0
        task['last_l'] = 0
        d = s.date() - lesson_time[1]['start'].date()
        for k, v in lesson_time.items():
            if v['end'] + d <= s or k == 1:
                task['first_l'] = k
            if v['start'] + d < e:
                task['last_l'] = k

    task_days = [task['real_date'] for task in ta]
    day_tasks = {day.date(): [] for day in sorted(list(set(task_days)))}
    t = [d for d in day_tasks.keys()]

    for task in ta:
        day_tasks[task['real_date'].date()].append(task)

    sched = LessonSchedule.query.filter(LessonSchedule.year == curr_year)
    vol_days = set(datetime.datetime.strftime(task['real_date'], '%w') for task in tasks)
    schedule = {day: sorted([{'lesson_no': lesson.lesson_no, 'lesson_name': lesson.lesson_name}
                             for lesson in sched.filter(LessonSchedule.weekday == day)
                             if lesson.lesson_id in my_lessons], key=lambda x: x['lesson_no'])
                for day in vol_days}

    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.6)
        section.bottom_margin = Cm(0.6)
        section.left_margin = Cm(0.6)
        section.right_margin = Cm(0.6)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    i = 0
    while i < len(t):
        d_tasks = day_tasks[t[i]]
        if len(schedule[datetime.datetime.strftime(d_tasks[0]['real_date'], '%w')]) > 0:
            a = document.add_paragraph()
            a.alignment = 1
            p = a.add_run('Пропуск уроков обучающимся: ' + name + ' (' + cl + ')' + '\n')
            p.bold = True
            pp = a.add_run('(волонтер Чтений им. В. И. Вернадского)')
            a.paragraph_format.space_after = Pt(6)

            b = document.add_paragraph()
            b.alignment = 1
            ppp = b.add_run(d_tasks[0]['task_date'])
            ppp.bold = True

            font = ppp.font
            font.name = 'Calibri'
            font.size = Pt(14)

            font = p.font
            font.name = 'Calibri'
            font.size = Pt(16)

            table = document.add_table(cols=4, rows=1)
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells = table.rows[0].cells

            hdr_cells[0].paragraphs[0].add_run('Место').bold = True
            hdr_cells[0].width = Cm(7.2)
            hdr_cells[1].paragraphs[0].add_run('Задача волонтера').bold = True
            hdr_cells[1].width = Cm(7.4)
            hdr_cells[2].paragraphs[0].add_run('Время работы/ уроки').bold = True
            hdr_cells[2].width = Cm(3.4)
            hdr_cells[3].paragraphs[0].add_run('Выбор задачи').bold = True
            hdr_cells[3].width = Cm(1.8)

            for task in d_tasks:
                if task['first_l'] != 0 and task['last_l'] != 0:
                    missing_lessons = '\n(' + str(task['first_l']) + '-' + str(task['last_l']) + ' урок)'
                elif task['first_l'] != 0:
                    missing_lessons = '\n(' + str(task['first_l']) + ' урок)'
                elif task['last_l'] != 0:
                    missing_lessons = '\n(' + str(task['last_l']) + ' урок)'
                else:
                    missing_lessons = ''

                row_cells = table.add_row().cells
                row_cells[0].text = task['location'] + '\n(' + task['address'] + ')'
                row_cells[0].width = Cm(7.2)
                row_cells[1].text = task['task_name']
                row_cells[1].width = Cm(7.4)
                row_cells[2].text = task['start_time'] + ' – ' + task['end_time'] + missing_lessons
                row_cells[2].width = Cm(3.4)
                row_cells[3].width = Cm(1.8)

            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row.height = Cm(1.3)

            # document.add_paragraph()
            document.add_paragraph()

            table = document.add_table(cols=3, rows=1)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            a = table.cell(0, 0)
            b = table.cell(0, 1)
            A = a.merge(b)

            hdr_cells[0].paragraphs[0].add_run('Урок').bold = True
            hdr_cells[0].width = Cm(1)
            hdr_cells[1].width = Cm(8.9)
            hdr_cells[2].paragraphs[0].add_run(
                'Пропуск согласован' + '\n' + '(подпись, задание при наличии)').bold = True
            hdr_cells[2].width = Cm(8.9)

            for one_lesson in schedule[datetime.datetime.strftime(d_tasks[0]['real_date'], '%w')]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(one_lesson['lesson_no'])
                row_cells[0].width = Cm(1)
                row_cells[1].text = one_lesson['lesson_name']
                row_cells[1].width = Cm(8.9)
                row_cells[2].width = Cm(8.9)

            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row.height = Cm(1.3)

            if i < len(day_tasks) - 1:
                document.add_page_break()
        i += 1

    if not os.path.isdir('static/files/generated_files'):
        os.mkdir('static/files/generated_files')
    document.save('static/files/generated_files/' + u.last_name + ' бегунок ЧВ.docx')
    return send_file('static/files/generated_files/' + u.last_name + ' бегунок ЧВ.docx', as_attachment=True)


# БАЗА ЗНАНИЙ

@app.route('/knowledge_main')
def knowledge_main():
    access = check_access(8)
    if access is not True:
        return access
    now = datetime.datetime.now().date()
    date = days_full[now.strftime('%w')] + ', ' + now.strftime('%d') + ' ' + months_full[
        now.strftime('%m')] + ' ' + now.strftime('%Y')
    return render_template('knowledge-main.html', date=date)


@app.route('/invoice')
def invoice():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/invoice.html')


@app.route('/contact')
def contact():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/contact.html')


@app.route('/email')
def email():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/email.html')


@app.route('/email_schedule')
def email_schedule():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/email_schedule.html')


@app.route('/phone_schedule')
def phone_schedule():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/phone_schedule.html')


@app.route('/working_programme')
def working_programme():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/working_programme.html')


@app.route('/online_additional_contest')
def online_additional_contest():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/online_additional_contest.html')


@app.route('/consult_works')
def consult_works():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/consult_works.html')


@app.route('/vernadsky_olympiade')
def vernadsky_olympiade():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/vernadsky_olympiade.html')


@app.route('/general_info')
def general_info():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/general_info.html')


@app.route('/frequent_actions')
def frequent_actions():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/frequent_actions.html')


@app.route('/registration_on_site')
def registration_on_site():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/registration_on_site.html')


@app.route('/attach_work')
def attach_work():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/attach_work.html')


@app.route('/approve_for_2_tour')
def approve_for_2_tour():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/approve_for_2_tour.html')


@app.route('/approve_for_1_tour')
def approve_for_1_tour():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/approve_for_1_tour.html')


# ОРГКОМИТЕТ

# @app.route('/contact_team')
# def contact_team():
#     if check_access(url='/invoice') < 8:
#         return redirect(url_for('.no_access'))
#     return render_template('knowledge/org/contact_team.html')

# s
@app.route('/bank_details')
def bank_details():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/bank_details.html')


@app.route('/banks_and_payments')
def banks_and_payments():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/banks_and_payments.html')


@app.route('/guarantee_letters')
def guarantee_letters():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/guarantee_letters.html')


@app.route('/creativity_contest')
def creativity_contest():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/creativity_contest.html')


@app.route('/session_shedule')
def session_shedule():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/session_shedule.html')


@app.route('/apply_2_tour')
def apply_2_tour():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/apply_2_tour.html')


@app.route('/programme_grid')
def programme_grid():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/programme_grid.html')


@app.route('/feedback')
def feedback():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/feedback.html')


@app.route('/movement_projects')
def movement_projects():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/movement_projects.html')


@app.route('/working_resources')
def working_resources():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/working_resources.html')


@app.route('/apply_for_participant')
def apply_for_participant():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/apply_for_participant.html')


@app.route('/contest_calendar')
def contest_calendar():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/contest_calendar.html')


@app.route('/apply_1_tour')
def apply_1_tour():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/apply_1_tour.html')


@app.route('/faq')
def faq():
    access = check_access(3)
    if access is not True:
        return access
    return render_template('knowledge/org/FAQ.html')


@app.route('/tour_2')
def tour_2():
    access = check_access(8)
    if access is not True:
        return access
    return render_template('knowledge/org/tour_2.html')


@app.route('/secretary_knowledge')
def secretary_knowledge():
    access = check_access(5)
    if access is not True:
        return access
    return render_template('secretary_knowledge.html')


# # ЯИССЛЕДОВАТЕЛЬ
# @app.route('/yais_main')
# def yais_main():
#     return render_template('ya_issledovatel/yais_main.html')
#
#
# @app.route('/yais_categories', defaults={'cat_to_edit': None})
# @app.route('/yais_categories/<cat_to_edit>')
# def yais_categories(cat_to_edit):
#     access = check_access(7)
#     if access is not True:
#         return access
#     if cat_to_edit:
#         cat_to_edit = int(cat_to_edit)
#         e_db = db.session.query(YaisCategories).filter(YaisCategories.cat_id == cat_to_edit).first()
#         cat_to_edit = {'cat_id': e_db.cat_id,
#                        'cat_name': e_db.cat_name,
#                        'cat_short_name': e_db.cat_short_name}
#     cats = [{'cat_id': c.cat_id, 'cat_name': c.cat_name, 'cat_short_name': c.cat_short_name}
#             for c in YaisCategories.query.filter(YaisCategories.year == curr_year).all()]
#     return render_template('ya_issledovatel/yais_categories.html', cat_to_edit=cat_to_edit, cats=cats)
#
#
# @app.route('/yais_add_category', methods=['POST'])
# def yais_add_category():
#     cat_name = request.form['cat_name']
#     cat_short_name = request.form['cat_short_name']
#     year = curr_year
#     category = YaisCategories(cat_name=cat_name, cat_short_name=cat_short_name, year=year)
#     if 'cat_id' in request.form.keys():
#         cat_id = int(request.form['cat_id'])
#         if category not in YaisCategories.query.filter(YaisCategories.cat_id == cat_id).all():
#             db.session.query(YaisCategories).filter(YaisCategories.cat_id == cat_id) \
#                 .update({YaisCategories.cat_name: cat_name,
#                          YaisCategories.cat_short_name: cat_short_name,
#                          YaisCategories.year: year})
#             db.session.commit()
#     else:
#         db.session.add(category)
#         db.session.commit()
#     return redirect(url_for('.yais_categories'))
#
#
# @app.route('/yais_delete_cat/<cat_id>')
# def yais_delete_cat(cat_id):
#     cat_id = int(cat_id)
#     if cat_id in [c.cat_id for c in YaisCategories.query.all()]:
#         to_del = db.session.query(YaisCategories).filter(YaisCategories.cat_id == cat_id).first()
#         db.session.delete(to_del)
#         db.session.commit()
#     return redirect(url_for('.yais_categories'))
#
#
# @app.route('/load_registration', defaults={'success': None})
# @app.route('/load_registration/<success>')
# def load_registration(success):
#     access = check_access(7)
#     if access is not True:
#         return access
#     return render_template('ya_issledovatel/yais_load_registration.html', success=success)
#
#
# @app.route('/add_registration', methods=['POST'])
# def add_registration():
#     data = request.files['file'].read().decode('mac_cyrillic').replace('\xa0', ' ')
#     lines = data.split('\n')
#     reg_data = []
#     for line in lines[1:]:
#         if line != '':
#             sta = {name.strip().strip('\r'): value.strip().strip('\r')
#                    for name, value in zip(lines[0].split('\t'), line.split('\t'))}
#             reg_data.append(sta)
#     for reg in reg_data:
#         if reg != {}:
#             organ = reg['ОО'].split(',')
#             org = [o.strip() for o in organ]
#             city = org[-1].strip('"')
#             w = YaisWorks(title=reg['Тема'])
#             if reg['Тема'] not in [w.title for w in YaisWorks.query.all()]:
#                 db.session.add(w)
#                 db.session.flush()
#                 db.session.commit()
#                 work_id = w.work_id
#             else:
#                 work_id = YaisWorks.query.filter(YaisWorks.title == reg['Тема']).first().work_id
#
#             name = reg['ФИ автора(ов)'].split()
#             if type(name) == list:
#                 if len(name) > 2:
#                     patronymic = name[2].strip()
#                 else:
#                     patronymic = ''
#                 if len(name) > 1:
#                     first_name = name[1].strip()
#                 else:
#                     first_name = ''
#                 last_name = name[0].strip()
#             else:
#                 last_name = name
#                 first_name = ''
#                 patronymic = ''
#             a = YaisAuthors(last_name=last_name, first_name=first_name, patronymic=patronymic, city=city)
#             if last_name not in [a.last_name for a in YaisAuthors.query.all()]:
#                 db.session.add(a)
#                 db.session.flush()
#                 db.session.commit()
#                 author_id = a.author_id
#             elif first_name not in [a.first_name for a in YaisAuthors.query
#                     .filter(YaisAuthors.last_name == last_name).all()]:
#                 db.session.add(a)
#                 db.session.flush()
#                 db.session.commit()
#                 author_id = a.author_id
#             elif patronymic not in [a.patronymic for a in YaisAuthors.query
#                     .filter(YaisAuthors.last_name == last_name).filter(YaisAuthors.first_name == first_name).all()]:
#                 db.session.add(a)
#                 db.session.flush()
#                 db.session.commit()
#                 author_id = a.author_id
#             elif city not in [a.city for a in YaisAuthors.query
#                     .filter(YaisAuthors.last_name == last_name).filter(YaisAuthors.first_name == first_name)
#                     .filter(YaisAuthors.patronymic == patronymic).all()]:
#                 db.session.add(a)
#                 db.session.flush()
#                 db.session.commit()
#                 author_id = a.author_id
#             else:
#                 author_id = YaisAuthors.query.filter(YaisAuthors.last_name == last_name) \
#                     .filter(YaisAuthors.first_name == first_name) \
#                     .filter(YaisAuthors.patronymic == patronymic).first().author_id
#
#             name = reg['Руководитель'].split()
#             if type(name) == list:
#                 if len(name) > 2:
#                     patronymic = name[2].strip()
#                 else:
#                     patronymic = ''
#                 if len(name) > 1:
#                     first_name = name[1].strip()
#                 else:
#                     first_name = ''
#                 last_name = name[0].strip()
#             else:
#                 last_name = name
#                 first_name = ''
#                 patronymic = ''
#             s = YaisSupervisors(last_name=last_name, first_name=first_name, patronymic=patronymic, city=city)
#             if last_name not in [s.last_name for s in YaisSupervisors.query.all()]:
#                 db.session.add(s)
#                 db.session.flush()
#                 db.session.commit()
#                 supervisor_id = s.supervisor_id
#             elif first_name not in [s.first_name for s in YaisSupervisors.query
#                     .filter(YaisSupervisors.last_name == last_name).all()]:
#                 db.session.add(s)
#                 db.session.flush()
#                 db.session.commit()
#                 supervisor_id = s.supervisor_id
#             elif patronymic not in [s.patronymic for s in YaisSupervisors.query
#                     .filter(YaisSupervisors.last_name == last_name)
#                     .filter(YaisSupervisors.first_name == first_name).all()]:
#                 db.session.add(s)
#                 db.session.flush()
#                 db.session.commit()
#                 supervisor_id = s.supervisor_id
#             elif city not in [s.city for s in YaisSupervisors.query
#                     .filter(YaisSupervisors.last_name == last_name).filter(YaisSupervisors.first_name == first_name)
#                     .filter(YaisSupervisors.patronymic == patronymic).all()]:
#                 db.session.add(s)
#                 db.session.flush()
#                 db.session.commit()
#                 supervisor_id = s.supervisor_id
#             else:
#                 supervisor_id = YaisSupervisors.query.filter(YaisSupervisors.last_name == last_name) \
#                     .filter(YaisSupervisors.first_name == first_name) \
#                     .filter(YaisSupervisors.patronymic == patronymic).first().supervisor_id
#
#             cl = reg['Класс'].split()
#             if cl[1] == 'лет':
#                 age = True
#             else:
#                 age = False
#             class_digit = int(cl[0])
#             cl = YaisClasses(class_digit=class_digit, age=age)
#             if class_digit not in [c.class_digit for c in YaisClasses.query.filter(YaisClasses.age == age).all()]:
#                 db.session.add(cl)
#                 db.session.flush()
#                 db.session.commit()
#                 class_id = cl.class_id
#             else:
#                 class_id = YaisClasses.query.filter(YaisClasses.class_digit == class_digit) \
#                     .filter(YaisClasses.age == age).first().class_id
#
#             cat_id = YaisCategories.query.filter(YaisCategories.cat_short_name == reg['Секция']).first().cat_id
#
#             r = YaisRegions(reg['Регион'])
#             if reg['Регион'] not in [r.region_name for r in YaisRegions.query.all()]:
#                 db.session.add(r)
#                 db.session.flush()
#                 region_id = r.region_id
#             else:
#                 region_id = YaisRegions.query.filter(YaisRegions.region_name == reg['Регион']).first().region_id
#
#             ci = YaisCities(city)
#             if city not in [c.city_name for c in YaisCities.query.all()]:
#                 db.session.add(ci)
#                 db.session.flush()
#                 db.session.commit()
#                 city_id = ci.city_id
#             else:
#                 city_id = YaisCities.query.filter(YaisCities.city_name == city).first().city_id
#
#             organisation = ','.join(org[:-1])
#             org = YaisOrganisations(organisation)
#             if org not in [o.organisation_id for o in YaisOrganisations.query.all()]:
#                 db.session.add(org)
#                 db.session.flush()
#                 db.session.commit()
#                 organisation_id = org.organisation_id
#             else:
#                 organisation_id = YaisCities.query \
#                     .filter(YaisCities.organisation_name == organisation).first().organisation_id
#
#             if city_id not in [c.city_id for c in YaisRegionCities.query.all()]:
#                 reg_city = YaisRegionCities(city_id, region_id)
#                 db.session.add(reg_city)
#                 db.session.commit()
#             if organisation_id not in [o.organisation_id for o in YaisCityOrganisations.query.all()]:
#                 city_org = YaisCityOrganisations(organisation_id, city_id)
#                 db.session.add(city_org)
#                 db.session.commit()
#             if supervisor_id not in [s.supervisor_id for s in YaisSupervisorOrganisation.query.all()]:
#                 supervisor_organisation = YaisSupervisorOrganisation(supervisor_id, organisation_id)
#                 db.session.add(supervisor_organisation)
#                 db.session.commit()
#             else:
#                 db.session.query(YaisSupervisorOrganisation) \
#                     .filter(YaisSupervisorOrganisation.supervisor_id == supervisor_id) \
#                     .update({YaisSupervisorOrganisation.organisation_id: organisation_id})
#             if work_id not in [w.work_id for w in YaisWorkOrganisation.query.all()]:
#                 work_org = YaisWorkOrganisation(work_id, organisation_id)
#                 db.session.add(work_org)
#                 db.session.commit()
#             if author_id not in [a.author_id for a in YaisAuthorClass.query.all()]:
#                 author_class = YaisAuthorClass(author_id, class_id)
#                 db.session.add(author_class)
#                 db.session.commit()
#             if work_id not in [w.work_id for w in YaisWorkCategories.query.all()]:
#                 work_cat = YaisWorkCategories(work_id, cat_id)
#                 db.session.add(work_cat)
#                 db.session.commit()
#             if work_id not in [w.work_id for w in YaisWorkAuthorSupervisor.query.all()]:
#                 work_author_supervisor = YaisWorkAuthorSupervisor(work_id, author_id, supervisor_id)
#                 db.session.add(work_author_supervisor)
#                 db.session.commit()
#             elif supervisor_id not in [w.supervisor_id for w
#                                        in YaisWorkAuthorSupervisor.query
#                                                .filter(YaisWorkAuthorSupervisor.work_id == work_id).all()]:
#                 work_author_supervisor = YaisWorkAuthorSupervisor(work_id, author_id, supervisor_id)
#                 db.session.add(work_author_supervisor)
#                 db.session.commit()
#             elif author_id not in [w.author_id for w
#                                    in YaisWorkAuthorSupervisor.query
#                                            .filter(YaisWorkAuthorSupervisor.work_id == work_id)
#                                            .filter(YaisWorkAuthorSupervisor.supervisor_id == supervisor_id).all()]:
#                 work_author_supervisor = YaisWorkAuthorSupervisor(work_id, author_id, supervisor_id)
#                 db.session.add(work_author_supervisor)
#                 db.session.commit()
#     success = True
#     return redirect(url_for('.load_registration', success=success))
#
#
# @app.route('/yais_id_payments', defaults={'length': 30, 'page': 1})
# @app.route('/yais_id_payments/<length>/<page>')
# def yais_id_payments(length, page):
#     access = check_access(7)
#     if access is not True:
#         return access
#     payments = [p.payment_id for p in BankStatement.query
#     .join(PaymentTypes, BankStatement.payment_id == PaymentTypes.payment_id)
#     .filter(PaymentTypes.payment_type == 'Я - Исследователь')
#     .order_by(BankStatement.date.desc()).order_by(BankStatement.order_id.asc()).all()]
#     n, data = make_pages(length, payments, page)
#     statement = statement_info(data)
#     return render_template('ya_issledovatel/yais_id_payments.html', statement=statement, pages=n, page=page,
#                            length=length, link='yais_id_payments')
#
#
# @app.route('/yais_set_payee/<payment_id>', defaults={'payee': None})
# @app.route('/yais_set_payee/<payment_id>/<payee>')
# def yais_set_payee(payment_id, payee):
#     access = check_access(7)
#     if access is not True:
#         return access
#     payment = payment_info(payment_id)
#     participant = {'type': None, 'participant': payee}
#     if payee is not None:
#         payee = payee.strip()
#         parts = [u.author_id for u in YaisAuthors.query.all()
#                  if payee.lower() in u.last_name.lower()]
#         parts.extend([u.author_id for u in YaisAuthors.query.all()
#                       if payee.lower() in u.first_name.lower()])
#         parts.extend([u.author_id for u in YaisAuthors.query.all()
#                       if payee.lower() in u.patronymic.lower()])
#         p = []
#         for part in parts:
#             w_db = db.session.query(YaisWorks) \
#                 .join(YaisWorkAuthorSupervisor, YaisWorks.work_id == YaisWorkAuthorSupervisor.work_id) \
#                 .filter(YaisWorkAuthorSupervisor.author_id == part).all()
#             for w in w_db:
#                 org = YaisOrganisations.query.join(YaisWorkOrganisation,
#                                                    YaisOrganisations.organisation_id ==
#                                                    YaisWorkOrganisation.organisation_id) \
#                     .filter(YaisWorkOrganisation.work_id == w.work_id).first()
#                 if w.work_id in [wp.work_id for wp in YaisWorkPayment.query.all()]:
#                     payed = True
#                     payment_id = YaisWorkPayment.query.filter(YaisWorkPayment.work_id == w.work_id) \
#                         .first().payment_id
#                 else:
#                     payed = False
#                     payment_id = None
#                 p.append({'work': w.title, 'work_id': w.work_id,
#                           'org_id': org.organisation_id, 'organisation': org.organisation_name,
#                           'payed': payed, 'payment_id': payment_id})
#         for w in p:
#             w['authors'] = [{'author_id': a.author_id,
#                              'author_name': a.last_name + ' ' + a.first_name + ' ' + a.patronymic,
#                              'city': a.city}
#                             for a in YaisAuthors.query
#                             .join(YaisWorkAuthorSupervisor, YaisAuthors.author_id == YaisWorkAuthorSupervisor.author_id)
#                             .filter(YaisWorkAuthorSupervisor.work_id == w['work_id']).all()]
#             w['supervisors'] = [{'supervisor_id': a.supervisor_id,
#                                  'supervisor_name': a.last_name + ' ' + a.first_name + ' ' + a.patronymic}
#                                 for a in YaisSupervisors.query
#                                 .join(YaisWorkAuthorSupervisor, YaisSupervisors.supervisor_id
#                                       == YaisWorkAuthorSupervisor.supervisor_id)
#                                 .filter(YaisWorkAuthorSupervisor.work_id == w['work_id']).all()]
#             for a in w['authors']:
#                 cl_db = YaisClasses.query.join(YaisAuthorClass, YaisClasses.class_id == YaisClasses.class_id) \
#                     .filter(YaisAuthorClass.author_id == a['author_id']).first()
#                 a_class = str(cl_db.class_digit)
#                 if cl_db.age:
#                     a_class += ' лет'
#                 else:
#                     a_class += ' класс'
#                 a['class'] = a_class
#
#         participant = {'type': 'name', 'participant': p}
#
#         if not parts:
#             participant = {'type': None, 'participant': payee}
#     else:
#         participant = {'type': None, 'participant': payee}
#     return render_template('ya_issledovatel/yais_set_payee.html', payment=payment, participant=participant)
#
#
# @app.route('/yais_application_payment/<payment_id>', methods=['GET'], defaults={'payee': None})
# @app.route('/yais_application_payment/<payment_id>/<payee>')
# def yais_application_payment(payment_id, payee):
#     if payee is None:
#         payee = request.values.get('payee', str)
#     return redirect(url_for('.yais_set_payee', payment_id=payment_id, payee=payee))
#
#
# @app.route('/yais_set_payment/<payment_id>/<payee>', methods=['POST'])
# def yais_set_payment(payment_id, payee):
#     participant = int(payee)
#     if str(participant) not in request.form.keys():
#         if participant in [p.work_id for p in YaisWorkPayment.query.all()]:
#             if YaisWorkPayment.query.filter(YaisWorkPayment.work_id == participant) \
#                     .first().payment_id == int(payment_id):
#                 YaisWorkPayment.query.filter(YaisWorkPayment.work_id == participant).delete()
#                 db.session.commit()
#     else:
#         data = request.form[str(participant)]
#         if data == 'on':
#             if participant not in [p.work_id for p in YaisWorkPayment.query.all()]:
#                 payment = YaisWorkPayment(participant, payment_id)
#                 db.session.add(payment)
#                 db.session.commit()
#             else:
#                 YaisWorkPayment.query.filter(YaisWorkPayment.work_id == participant).delete()
#                 db.session.commit()
#     db.session.commit()
#     return redirect(url_for('.yais_id_payments'))
#
#
# @app.route('/yais_find_participant', defaults={'query': 'sear'})
# @app.route('/yais_find_participant/<query>')
# def yais_find_participant(query):
#     access = check_access(3)
#     if access is not True:
#         return access
#     response = {'type': None, 'value': query}
#
#     if query:
#         if query == 'sear':
#             response = 'search'
#         else:
#             query = query.strip()
#             parts = [u.author_id for u in YaisAuthors.query.all()
#                      if query.lower() in u.last_name.lower()]
#             parts.extend([u.author_id for u in YaisAuthors.query.all()
#                           if query.lower() in u.first_name.lower()])
#             parts.extend([u.author_id for u in YaisAuthors.query.all()
#                           if query.lower() in u.patronymic.lower()])
#             p = []
#             for part in parts:
#                 w_db = db.session.query(YaisWorks) \
#                     .join(YaisWorkAuthorSupervisor, YaisWorks.work_id == YaisWorkAuthorSupervisor.work_id) \
#                     .filter(YaisWorkAuthorSupervisor.author_id == part).all()
#                 for w in w_db:
#                     org = YaisOrganisations.query.join(YaisWorkOrganisation,
#                                                        YaisOrganisations.organisation_id ==
#                                                        YaisWorkOrganisation.organisation_id) \
#                         .filter(YaisWorkOrganisation.work_id == w.work_id).first()
#                     if w.work_id in [wp.work_id for wp in YaisWorkPayment.query.all()]:
#                         payed = True
#                         payment_id = YaisWorkPayment.query.filter(YaisWorkPayment.work_id == w.work_id) \
#                             .first().payment_id
#                     else:
#                         payed = False
#                         payment_id = None
#                     cat = YaisCategories.query \
#                         .join(YaisWorkCategories, YaisCategories.cat_id == YaisWorkCategories.cat_id) \
#                         .filter(YaisWorkCategories.work_id == w.work_id).first()
#                     p.append({'work': w.title, 'work_id': w.work_id,
#                               'cat_id': cat.cat_id, 'cat_name': cat.cat_name,
#                               'org_id': org.organisation_id, 'organisation': org.organisation_name,
#                               'payed': payed, 'payment_id': payment_id})
#             for w in p:
#                 w['authors'] = [{'author_id': a.author_id,
#                                  'author_name': a.last_name + ' ' + a.first_name + ' ' + a.patronymic,
#                                  'city': a.city}
#                                 for a in YaisAuthors.query
#                                 .join(YaisWorkAuthorSupervisor,
#                                       YaisAuthors.author_id == YaisWorkAuthorSupervisor.author_id)
#                                 .filter(YaisWorkAuthorSupervisor.work_id == w['work_id']).all()]
#                 w['supervisors'] = [{'supervisor_id': a.supervisor_id,
#                                      'supervisor_name': a.last_name + ' ' + a.first_name + ' ' + a.patronymic}
#                                     for a in YaisSupervisors.query
#                                     .join(YaisWorkAuthorSupervisor, YaisSupervisors.supervisor_id
#                                           == YaisWorkAuthorSupervisor.supervisor_id)
#                                     .filter(YaisWorkAuthorSupervisor.work_id == w['work_id']).all()]
#                 for a in w['authors']:
#                     cl_db = YaisClasses.query.join(YaisAuthorClass, YaisClasses.class_id == YaisClasses.class_id) \
#                         .filter(YaisAuthorClass.author_id == a['author_id']).first()
#                     a_class = str(cl_db.class_digit)
#                     if cl_db.age:
#                         a_class += ' лет'
#                     else:
#                         a_class += ' класс'
#                     a['class'] = a_class
#             response = {'type': 'appls', 'value': p}
#     else:
#         response = {'type': None, 'value': query}
#     return render_template('ya_issledovatel/yais_find_participant.html', response=response)
#
#
# @app.route('/yais_searching_participant', methods=['GET'])
# def yais_searching_participant():
#     renew_session()
#     query = request.values.get('query', str)
#     return redirect(url_for('.yais_find_participant', query=query))
#
#
# @app.route('/yais_check_arrival')
# def yais_check_arrival():
#     access = check_access(3)
#     if access is not True:
#         return access
#     authors = [{'author_id': a.author_id,
#                 'author_name': a.last_name + ' ' + a.first_name + ' ' + a.patronymic,
#                 'city': a.city} for a in YaisAuthors.query.order_by(YaisAuthors.last_name).all()]
#     for a in authors:
#         cl_db = YaisClasses.query.join(YaisAuthorClass, YaisClasses.class_id == YaisClasses.class_id) \
#             .filter(YaisAuthorClass.author_id == a['author_id']).first()
#         a_class = str(cl_db.class_digit)
#         if cl_db.age:
#             a_class += ' лет'
#         else:
#             a_class += ' класс'
#         a['class'] = a_class
#     supervisors = [{'supervisor_id': a.supervisor_id,
#                     'supervisor_name': a.last_name + ' ' + a.first_name + ' ' + a.patronymic,
#                     'city': a.city}
#                    for a in YaisSupervisors.query.order_by(YaisSupervisors.last_name).all()]
#     for a in authors:
#         w_db = db.session.query(YaisWorks) \
#             .join(YaisWorkAuthorSupervisor, YaisWorks.work_id == YaisWorkAuthorSupervisor.work_id) \
#             .filter(YaisWorkAuthorSupervisor.author_id == a['author_id']).first()
#         org = YaisOrganisations.query.join(YaisWorkOrganisation,
#                                            YaisOrganisations.organisation_id ==
#                                            YaisWorkOrganisation.organisation_id) \
#             .filter(YaisWorkOrganisation.work_id == w_db.work_id).first()
#         if w_db.work_id in [wp.work_id for wp in YaisWorkPayment.query.all()]:
#             payed = True
#             payment_id = YaisWorkPayment.query.filter(YaisWorkPayment.work_id == w_db.work_id) \
#                 .first().payment_id
#         else:
#             payed = False
#             payment_id = None
#         cat = YaisCategories.query \
#             .join(YaisWorkCategories, YaisCategories.cat_id == YaisWorkCategories.cat_id) \
#             .filter(YaisWorkCategories.work_id == w_db.work_id).first()
#         a['work'] = w_db.title
#         a['work_id'] = w_db.work_id
#         a['cat_id'] = cat.cat_id
#         a['cat_name'] = cat.cat_name
#         a['org_id'] = org.organisation_id
#         a['organisation'] = org.organisation_name
#         a['payed'] = payed
#         a['payment_id'] = payment_id
#         if a['author_id'] in [au.author_id for au in YaisArrival.query.all()]:
#             a['arrived'] = YaisArrival.query.filter(YaisArrival.author_id == a['author_id']).first().arrived
#         else:
#             a['arrived'] = True
#     for a in supervisors:
#         w_db = db.session.query(YaisWorks) \
#             .join(YaisWorkAuthorSupervisor, YaisWorks.work_id == YaisWorkAuthorSupervisor.work_id) \
#             .filter(YaisWorkAuthorSupervisor.author_id == a['supervisor_id']).first()
#         org = YaisOrganisations.query.join(YaisWorkOrganisation,
#                                            YaisOrganisations.organisation_id ==
#                                            YaisWorkOrganisation.organisation_id) \
#             .filter(YaisWorkOrganisation.work_id == w_db.work_id).first()
#         if w_db.work_id in [wp.work_id for wp in YaisWorkPayment.query.all()]:
#             payed = True
#             payment_id = YaisWorkPayment.query.filter(YaisWorkPayment.work_id == w_db.work_id) \
#                 .first().payment_id
#         else:
#             payed = False
#             payment_id = None
#         cat = YaisCategories.query \
#             .join(YaisWorkCategories, YaisCategories.cat_id == YaisWorkCategories.cat_id) \
#             .filter(YaisWorkCategories.work_id == w_db.work_id).first()
#         a['work'] = w_db.title
#         a['work_id'] = w_db.work_id
#         a['cat_id'] = cat.cat_id
#         a['cat_name'] = cat.cat_name
#         a['org_id'] = org.organisation_id
#         a['organisation'] = org.organisation_name
#         a['payed'] = payed
#         a['payment_id'] = payment_id
#         if a['supervisor_id'] in [au.supervisor_id for au in YaisArrival.query.all()]:
#             a['arrived'] = YaisArrival.query.filter(YaisArrival.supervisor_id == a['supervisor_id']).first().arrived
#         else:
#             a['arrived'] = True
#     return render_template('ya_issledovatel/yais_check_arrival.html', authors=authors, supervisors=supervisors)
#
#
# @app.route('/yais_save_arrival', methods=['POST'])
# def yais_save_arrival():
#     a = request.form.getlist('author')
#     s = request.form.getlist('supervisor')
#     authors = [int(au) for au in a]
#     supervisors = [int(su) for su in s]
#     auth_to_del = [a.author_id for a in YaisArrival.query.all() if a.author_id not in authors]
#     sup_to_del = [a.supervisor_id for a in YaisArrival.query.all() if a.supervisor_id not in supervisors]
#     for a in auth_to_del:
#         db.session.query(YaisArrival).filter(YaisArrival.author_id == a).update({YaisArrival.arrived: False})
#         db.session.commit()
#     for a in sup_to_del:
#         db.session.query(YaisArrival).filter(YaisArrival.supervisor_id == a).update({YaisArrival.arrived: False})
#         db.session.commit()
#     for a in authors:
#         arrived = YaisArrival(author_id=a, supervisor_id=None, arrived=True)
#         if arrived not in YaisArrival.query.all():
#             if a in [au.author_id for au in YaisArrival.query.all()]:
#                 db.session.query(YaisArrival).filter(YaisArrival.author_id == a).update({YaisArrival.arrived: True})
#                 db.session.commit()
#             else:
#                 db.session.add(arrived)
#                 db.session.commit()
#     for a in supervisors:
#         arrived = YaisArrival(author_id=None, supervisor_id=a, arrived=True)
#         if arrived not in YaisArrival.query.all():
#             if a in [su.supervisor_id for su in YaisArrival.query.all()]:
#                 db.session.query(YaisArrival).filter(YaisArrival.supervisor_id == a).update({YaisArrival.arrived: True})
#                 db.session.commit()
#             else:
#                 db.session.add(arrived)
#                 db.session.commit()
#     return redirect(url_for('.yais_check_arrival'))
#
#
# @app.route('/yais_stats')
# def yais_stats():
#     access = check_access(7)
#     if access is not True:
#         return access
#     ar = [w.work_id for w in YaisWorks.query
#     .join(YaisWorkAuthorSupervisor, YaisWorks.work_id == YaisWorkAuthorSupervisor.work_id)
#     .join(YaisArrival, YaisWorkAuthorSupervisor.author_id == YaisArrival.author_id)
#     .filter(YaisArrival.arrived == 1).all()]
#     arrived = len(ar)
#     pa = [w.work_id for w in YaisWorkPayment.query.all()]
#     p_payed = len(pa)
#     payed_not_arrived = len([w for w in pa if w not in ar])
#     un = [w for w in ar if w not in pa]
#     p = []
#     for part in un:
#         w_db = db.session.query(YaisWorks).filter(YaisWorks.work_id == part).first()
#         org = YaisOrganisations.query.join(YaisWorkOrganisation,
#                                            YaisOrganisations.organisation_id ==
#                                            YaisWorkOrganisation.organisation_id) \
#             .filter(YaisWorkOrganisation.work_id == w_db.work_id).first()
#         if w_db.work_id in [wp.work_id for wp in YaisWorkPayment.query.all()]:
#             payed = True
#             payment_id = YaisWorkPayment.query.filter(YaisWorkPayment.work_id == w.work_id) \
#                 .first().payment_id
#         else:
#             payed = False
#             payment_id = None
#         cat = YaisCategories.query \
#             .join(YaisWorkCategories, YaisCategories.cat_id == YaisWorkCategories.cat_id) \
#             .filter(YaisWorkCategories.work_id == w_db.work_id).first()
#         p.append({'work': w_db.title, 'work_id': w_db.work_id,
#                   'cat_id': cat.cat_id, 'cat_name': cat.cat_name,
#                   'org_id': org.organisation_id, 'organisation': org.organisation_name,
#                   'payed': payed, 'payment_id': payment_id})
#     for w in p:
#         w['authors'] = [{'author_id': a.author_id,
#                          'author_name': a.last_name + ' ' + a.first_name + ' ' + a.patronymic,
#                          'city': a.city}
#                         for a in YaisAuthors.query
#                         .join(YaisWorkAuthorSupervisor, YaisAuthors.author_id == YaisWorkAuthorSupervisor.author_id)
#                         .filter(YaisWorkAuthorSupervisor.work_id == w['work_id']).all()]
#         w['supervisors'] = [{'supervisor_id': a.supervisor_id,
#                              'supervisor_name': a.last_name + ' ' + a.first_name + ' ' + a.patronymic}
#                             for a in YaisSupervisors.query
#                             .join(YaisWorkAuthorSupervisor, YaisSupervisors.supervisor_id
#                                   == YaisWorkAuthorSupervisor.supervisor_id)
#                             .filter(YaisWorkAuthorSupervisor.work_id == w['work_id']).all()]
#         for a in w['authors']:
#             cl_db = YaisClasses.query.join(YaisAuthorClass, YaisClasses.class_id == YaisClasses.class_id) \
#                 .filter(YaisAuthorClass.author_id == a['author_id']).first()
#             a_class = str(cl_db.class_digit)
#             if cl_db.age:
#                 a_class += ' лет'
#             else:
#                 a_class += ' класс'
#             a['class'] = a_class
#     return render_template('ya_issledovatel/yais_stats.html', arrived=arrived, payed=p_payed,
#                            payed_not_arrived=payed_not_arrived, unpayed=p)


if __name__ == '__main__':
    app.run(debug=False)
